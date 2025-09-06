import logging
import requests
import json
import os
from typing import List, Dict, Any, Optional
from datetime import datetime
import hashlib
import re
import io

# Импорты для обработки документов
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Настройка логирования
logger = logging.getLogger(__name__)

class ChatDocumentService:
    """Сервис для обработки документов в чате с ИИ"""
    
    def __init__(self):
        # Конфигурация
        self.QDRANT_URL = "http://qdrant:6333"  # Qdrant в Docker
        self.RAG_SERVICE_URL = "http://ai-nk-rag-service-1:8003"  # RAG сервис
        self.DOCUMENT_PARSER_URL = "http://ai-nk-document-parser-1:8001"  # Document Parser
        self.CHAT_COLLECTION = "chat_documents"
        self.VECTOR_SIZE = 1024  # Размер эмбеддинга BGE-M3
        self.MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
        
        # Поддерживаемые форматы файлов
        self.SUPPORTED_FORMATS = [
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.ms-excel',
            'text/plain',
            'text/markdown'
        ]
        
        logger.info("🚀 [CHAT_DOCUMENT_SERVICE] Chat Document Service initialized")
    
    def validate_file(self, file_content: bytes, file_type: str, file_name: str) -> Dict[str, Any]:
        """Валидация загружаемого файла"""
        try:
            # Проверяем размер файла
            file_size = len(file_content)
            if file_size > self.MAX_FILE_SIZE:
                return {
                    "valid": False,
                    "error": f"Файл слишком большой. Максимальный размер: {self.MAX_FILE_SIZE // (1024*1024)}MB"
                }
            
            # Проверяем расширение файла
            file_extension = os.path.splitext(file_name)[1].lower()
            allowed_extensions = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt', '.md']
            if file_extension not in allowed_extensions:
                return {
                    "valid": False,
                    "error": f"Неподдерживаемое расширение файла: {file_extension}"
                }
            
            # Определяем правильный MIME тип по расширению файла
            extension_to_mime = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.doc': 'application/msword',
                '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                '.xls': 'application/vnd.ms-excel',
                '.txt': 'text/plain',
                '.md': 'text/markdown'
            }
            
            # Если MIME тип неопределен или неправильный, используем тип по расширению
            if file_type == 'application/octet-stream' or file_type not in self.SUPPORTED_FORMATS:
                file_type = extension_to_mime.get(file_extension, file_type)
                logger.info(f"🔍 [FILE_VALIDATION] Auto-detected MIME type for {file_name}: {file_type}")
            
            # Проверяем тип файла
            if file_type not in self.SUPPORTED_FORMATS:
                return {
                    "valid": False,
                    "error": f"Неподдерживаемый тип файла: {file_type}"
                }
            
            return {
                "valid": True,
                "file_size": file_size,
                "file_type": file_type,
                "file_extension": file_extension
            }
            
        except Exception as e:
            logger.error(f"❌ [FILE_VALIDATION] Error validating file: {e}")
            return {
                "valid": False,
                "error": f"Ошибка валидации файла: {str(e)}"
            }
    
    def process_document(self, file_content: bytes, file_type: str, file_name: str, user_message: str = "") -> Dict[str, Any]:
        """Обработка документа для чата"""
        try:
            logger.info(f"📄 [DOCUMENT_PROCESSING] Processing document: {file_name}")
            
            # Валидация файла
            validation = self.validate_file(file_content, file_type, file_name)
            if not validation["valid"]:
                return {
                    "success": False,
                    "error": validation["error"]
                }
            
            # Создаем уникальный ID для документа
            file_hash = hashlib.sha256(file_content).hexdigest()
            document_id = f"chat_{file_hash[:16]}"
            
            # Отправляем файл в Document Parser для извлечения текста
            parser_response = self._extract_text_from_document(file_content, file_type, file_name)
            if not parser_response["success"]:
                return {
                    "success": False,
                    "error": f"Ошибка извлечения текста: {parser_response['error']}"
                }
            
            extracted_text = parser_response["text"]
            chunks = parser_response.get("chunks", [])
            
            # Создаем эмбеддинги и индексируем в Qdrant
            indexing_result = self._index_document_in_qdrant(
                document_id, file_name, extracted_text, chunks, user_message
            )
            
            if not indexing_result["success"]:
                return {
                    "success": False,
                    "error": f"Ошибка индексации: {indexing_result['error']}"
                }
            
            # Формируем контекст для ИИ
            ai_context = self._prepare_ai_context(file_name, extracted_text, user_message)
            
            return {
                "success": True,
                "document_id": document_id,
                "file_name": file_name,
                "file_size": validation["file_size"],
                "extracted_text": extracted_text,
                "chunks_count": len(chunks),
                "ai_context": ai_context,
                "indexed": True
            }
            
        except Exception as e:
            logger.error(f"❌ [DOCUMENT_PROCESSING] Error processing document: {e}")
            return {
                "success": False,
                "error": f"Ошибка обработки документа: {str(e)}"
            }
    
    def _extract_text_from_document(self, file_content: bytes, file_type: str, file_name: str) -> Dict[str, Any]:
        """Извлечение текста из документа через прямую обработку"""
        try:
            file_extension = os.path.splitext(file_name)[1].lower()
            
            # Для TXT файлов
            if file_type == 'text/plain' or file_extension == '.txt':
                logger.info(f"🔍 [TEXT_EXTRACTION] Processing TXT file directly: {file_name}")
                return self._extract_text_from_txt(file_content)
            
            # Для DOCX файлов
            elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_extension == '.docx':
                logger.info(f"🔍 [TEXT_EXTRACTION] Processing DOCX file directly: {file_name}")
                return self._extract_text_from_docx(file_content)
            
            # Для PDF файлов
            elif file_type == 'application/pdf' or file_extension == '.pdf':
                logger.info(f"🔍 [TEXT_EXTRACTION] Processing PDF file directly: {file_name}")
                return self._extract_text_from_pdf(file_content)
            
            # Для Excel файлов
            elif file_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel'] or file_extension in ['.xlsx', '.xls']:
                logger.info(f"🔍 [TEXT_EXTRACTION] Processing Excel file directly: {file_name}")
                return self._extract_text_from_excel(file_content)
            
            # Для Markdown файлов
            elif file_type == 'text/markdown' or file_extension == '.md':
                logger.info(f"🔍 [TEXT_EXTRACTION] Processing Markdown file directly: {file_name}")
                return self._extract_text_from_txt(file_content)
            
            else:
                logger.error(f"❌ [TEXT_EXTRACTION] Unsupported file type: {file_type}")
                return {
                    "success": False,
                    "error": f"Unsupported file type: {file_type}"
                }
                
        except Exception as e:
            logger.error(f"❌ [TEXT_EXTRACTION] Error extracting text: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _extract_text_from_txt(self, file_content: bytes) -> Dict[str, Any]:
        """Извлечение текста из TXT файла"""
        try:
            # Декодируем содержимое файла
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text = file_content.decode('cp1251')  # Попытка с Windows-1251
                except UnicodeDecodeError:
                    text = file_content.decode('latin-1')  # Fallback
            
            # Создаем простые чанки
            chunk_size = 500
            chunks = []
            for i in range(0, len(text), chunk_size):
                chunk_text = text[i:i + chunk_size]
                chunks.append({
                    'content': chunk_text,
                    'page': 1,
                    'section': f'chunk_{i//chunk_size + 1}'
                })
            
            logger.info(f"✅ [TEXT_EXTRACTION] TXT file processed: {len(chunks)} chunks")
            return {
                "success": True,
                "text": text,
                "chunks": chunks,
                "pages": 1
            }
        except Exception as e:
            logger.error(f"❌ [TEXT_EXTRACTION] Error processing TXT: {e}")
            return {"success": False, "error": str(e)}
    
    def _extract_text_from_docx(self, file_content: bytes) -> Dict[str, Any]:
        """Извлечение текста из DOCX файла"""
        try:
            if not DOCX_AVAILABLE:
                return {"success": False, "error": "python-docx library not available"}
            
            # Создаем временный файл в памяти
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            
            # Извлекаем весь текст
            full_text = []
            chunks = []
            chunk_index = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
                    
                    # Создаем чанки из параграфов
                    if len(paragraph.text) > 100:
                        # Разбиваем длинные параграфы
                        words = paragraph.text.split()
                        for i in range(0, len(words), 50):  # 50 слов на чанк
                            chunk_text = ' '.join(words[i:i+50])
                            chunks.append({
                                'content': chunk_text,
                                'page': 1,
                                'section': f'paragraph_{chunk_index}'
                            })
                            chunk_index += 1
                    else:
                        chunks.append({
                            'content': paragraph.text,
                            'page': 1,
                            'section': f'paragraph_{chunk_index}'
                        })
                        chunk_index += 1
            
            # Обрабатываем таблицы
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))
                
                if table_text:
                    table_content = '\n'.join(table_text)
                    full_text.append(f"\n[ТАБЛИЦА]\n{table_content}\n")
                    chunks.append({
                        'content': table_content,
                        'page': 1,
                        'section': f'table_{chunk_index}'
                    })
                    chunk_index += 1
            
            extracted_text = '\n\n'.join(full_text)
            
            logger.info(f"✅ [TEXT_EXTRACTION] DOCX file processed: {len(chunks)} chunks")
            return {
                "success": True,
                "text": extracted_text,
                "chunks": chunks,
                "pages": 1
            }
        except Exception as e:
            logger.error(f"❌ [TEXT_EXTRACTION] Error processing DOCX: {e}")
            return {"success": False, "error": str(e)}
    
    def _extract_text_from_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """Извлечение текста из PDF файла"""
        try:
            if not PDF_AVAILABLE:
                return {"success": False, "error": "PyPDF2 library not available"}
            
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            full_text = []
            chunks = []
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    full_text.append(page_text)
                    
                    # Создаем чанки по страницам
                    if len(page_text) > 1000:
                        # Разбиваем длинные страницы
                        words = page_text.split()
                        for i in range(0, len(words), 100):  # 100 слов на чанк
                            chunk_text = ' '.join(words[i:i+100])
                            chunks.append({
                                'content': chunk_text,
                                'page': page_num,
                                'section': f'page_{page_num}_chunk_{i//100 + 1}'
                            })
                    else:
                        chunks.append({
                            'content': page_text,
                            'page': page_num,
                            'section': f'page_{page_num}'
                        })
            
            extracted_text = '\n\n'.join(full_text)
            
            logger.info(f"✅ [TEXT_EXTRACTION] PDF file processed: {len(chunks)} chunks")
            return {
                "success": True,
                "text": extracted_text,
                "chunks": chunks,
                "pages": len(pdf_reader.pages)
            }
        except Exception as e:
            logger.error(f"❌ [TEXT_EXTRACTION] Error processing PDF: {e}")
            return {"success": False, "error": str(e)}
    
    def _extract_text_from_excel(self, file_content: bytes) -> Dict[str, Any]:
        """Извлечение текста из Excel файла"""
        try:
            if not EXCEL_AVAILABLE:
                return {"success": False, "error": "openpyxl library not available"}
            
            excel_file = io.BytesIO(file_content)
            workbook = load_workbook(excel_file, data_only=True)
            
            full_text = []
            chunks = []
            chunk_index = 0
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = [f"[ЛИСТ: {sheet_name}]"]
                
                for row in sheet.iter_rows(values_only=True):
                    row_data = []
                    for cell in row:
                        if cell is not None and str(cell).strip():
                            row_data.append(str(cell).strip())
                    if row_data:
                        sheet_text.append(' | '.join(row_data))
                
                if len(sheet_text) > 1:  # Есть данные кроме заголовка
                    sheet_content = '\n'.join(sheet_text)
                    full_text.append(sheet_content)
                    
                    chunks.append({
                        'content': sheet_content,
                        'page': 1,
                        'section': f'sheet_{sheet_name}'
                    })
                    chunk_index += 1
            
            extracted_text = '\n\n'.join(full_text)
            
            logger.info(f"✅ [TEXT_EXTRACTION] Excel file processed: {len(chunks)} chunks")
            return {
                "success": True,
                "text": extracted_text,
                "chunks": chunks,
                "pages": 1
            }
        except Exception as e:
            logger.error(f"❌ [TEXT_EXTRACTION] Error processing Excel: {e}")
            return {"success": False, "error": str(e)}
    
    def _index_document_in_qdrant(self, document_id: str, file_name: str, text: str, chunks: List[Dict], user_message: str) -> Dict[str, Any]:
        """Индексация документа в Qdrant коллекцию chat_documents"""
        try:
            logger.info(f"🔍 [QDRANT_INDEXING] Starting indexing for document {document_id}")
            logger.info(f"🔍 [QDRANT_INDEXING] Document: {file_name}, chunks: {len(chunks)}")
            
            # Проверяем существование коллекции
            self._ensure_chat_collection_exists()
            
            # Создаем эмбеддинги для чанков
            logger.info(f"🔍 [QDRANT_INDEXING] Creating embeddings for {len(chunks)} chunks...")
            embeddings = self._create_chunk_embeddings(chunks)
            logger.info(f"🔍 [QDRANT_INDEXING] Created {len(embeddings)} embeddings")
            
            # Подготавливаем точки для Qdrant
            points = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                # Проверяем размер эмбеддинга
                if len(embedding) != self.VECTOR_SIZE:
                    logger.warning(f"⚠️ [QDRANT_INDEXING] Chunk {i} has wrong embedding size: {len(embedding)}, expected {self.VECTOR_SIZE}")
                    # Создаем правильный размер эмбеддинга
                    if len(embedding) > self.VECTOR_SIZE:
                        embedding = embedding[:self.VECTOR_SIZE]
                    else:
                        embedding.extend([0.0] * (self.VECTOR_SIZE - len(embedding)))
                
                point = {
                    'ids': [hash(f"{document_id}_{i}") % (2**63 - 1)],
                    'vectors': [embedding],
                    'payloads': [{
                        'document_id': document_id,
                        'file_name': file_name,
                        'chunk_index': i,
                        'content': chunk.get('content', ''),
                        'page': chunk.get('page', 1),
                        'section': chunk.get('section', ''),
                        'user_message': user_message,
                        'timestamp': datetime.now().isoformat(),
                        'chunk_type': 'chat_document'
                    }]
                }
                points.append(point)
                logger.debug(f"🔍 [QDRANT_INDEXING] Prepared point {i}: id={point['ids'][0]}, content_length={len(point['payloads'][0]['content'])}")
            
            logger.info(f"🔍 [QDRANT_INDEXING] Prepared {len(points)} points for Qdrant")
            
            # Добавляем точки в Qdrant
            logger.info(f"🔍 [QDRANT_INDEXING] Sending points to Qdrant...")
            
            # Отправляем каждую точку отдельно для избежания проблем с форматом
            for i, point in enumerate(points):
                try:
                    response = requests.post(
                        f"{self.QDRANT_URL}/collections/{self.CHAT_COLLECTION}/points",
                        json=point,
                        timeout=30
                    )
                    
                    if response.status_code != 200:
                        logger.error(f"❌ [QDRANT_INDEXING] Failed to add point {i}: {response.status_code} - {response.text}")
                        return {
                            "success": False,
                            "error": f"Qdrant indexing error for point {i}: {response.status_code} - {response.text}"
                        }
                    else:
                        logger.debug(f"✅ [QDRANT_INDEXING] Point {i} added successfully")
                        
                except Exception as e:
                    logger.error(f"❌ [QDRANT_INDEXING] Error adding point {i}: {e}")
                    return {
                        "success": False,
                        "error": f"Error adding point {i}: {str(e)}"
                    }
            
            logger.info(f"✅ [QDRANT_INDEXING] Indexed {len(points)} chunks for document {document_id}")
            return {"success": True, "chunks_indexed": len(points)}
                
        except Exception as e:
            logger.error(f"❌ [QDRANT_INDEXING] Error indexing in Qdrant: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e)
            }
    
    def _ensure_chat_collection_exists(self):
        """Создание коллекции chat_documents в Qdrant если не существует"""
        try:
            # Проверяем существование коллекции
            response = requests.get(f"{self.QDRANT_URL}/collections/{self.CHAT_COLLECTION}")
            
            if response.status_code == 404:
                # Создаем коллекцию
                collection_config = {
                    "vectors": {
                        "size": self.VECTOR_SIZE,
                        "distance": "Cosine"
                    },
                    "on_disk_payload": True
                }
                
                create_response = requests.put(
                    f"{self.QDRANT_URL}/collections/{self.CHAT_COLLECTION}",
                    json=collection_config
                )
                
                if create_response.status_code == 200:
                    logger.info(f"✅ [QDRANT_COLLECTION] Created collection {self.CHAT_COLLECTION}")
                else:
                    logger.error(f"❌ [QDRANT_COLLECTION] Failed to create collection: {create_response.status_code}")
            
        except Exception as e:
            logger.error(f"❌ [QDRANT_COLLECTION] Error ensuring collection exists: {e}")
    
    def _create_chunk_embeddings(self, chunks: List[Dict]) -> List[List[float]]:
        """Создание эмбеддингов для чанков через RAG сервис"""
        try:
            logger.info(f"🔍 [EMBEDDING] Starting embedding creation for {len(chunks)} chunks")
            embeddings = []
            
            for i, chunk in enumerate(chunks):
                content = chunk.get('content', '')
                logger.debug(f"🔍 [EMBEDDING] Processing chunk {i}: content_length={len(content)}")
                
                if not content.strip():
                    # Пустой чанк - создаем нулевой эмбеддинг
                    logger.warning(f"⚠️ [EMBEDDING] Chunk {i} is empty, creating zero embedding")
                    embeddings.append([0.0] * self.VECTOR_SIZE)
                    continue
                
                # Отправляем запрос на создание эмбеддинга
                logger.debug(f"🔍 [EMBEDDING] Sending embedding request for chunk {i}")
                response = requests.post(
                    f"{self.RAG_SERVICE_URL}/api/embeddings",
                    json={'text': content},
                    timeout=30
                )
                
                logger.debug(f"🔍 [EMBEDDING] Chunk {i} embedding response status: {response.status_code}")
                
                if response.status_code == 200:
                    result = response.json()
                    embedding = result.get('embedding', [])
                    logger.debug(f"🔍 [EMBEDDING] Chunk {i} embedding size: {len(embedding)}")
                    
                    if embedding:
                        # Проверяем, что эмбеддинг содержит числа
                        if all(isinstance(x, (int, float)) for x in embedding):
                            embeddings.append(embedding)
                            logger.debug(f"🔍 [EMBEDDING] Chunk {i} embedding added successfully")
                        else:
                            logger.warning(f"⚠️ [EMBEDDING] Chunk {i} embedding contains non-numeric values, using zero embedding")
                            embeddings.append([0.0] * self.VECTOR_SIZE)
                    else:
                        logger.warning(f"⚠️ [EMBEDDING] Chunk {i} empty embedding, using zero embedding")
                        embeddings.append([0.0] * self.VECTOR_SIZE)
                else:
                    logger.warning(f"⚠️ [EMBEDDING] Failed to create embedding for chunk {i}: {response.status_code} - {response.text}")
                    embeddings.append([0.0] * self.VECTOR_SIZE)
            
            logger.info(f"✅ [EMBEDDING] Created {len(embeddings)} embeddings successfully")
            return embeddings
            
        except Exception as e:
            logger.error(f"❌ [EMBEDDING] Error creating embeddings: {e}", exc_info=True)
            # Возвращаем нулевые эмбеддинги в случае ошибки
            logger.warning(f"⚠️ [EMBEDDING] Returning zero embeddings due to error")
            return [[0.0] * self.VECTOR_SIZE] * len(chunks)
    
    def _prepare_ai_context(self, file_name: str, text: str, user_message: str) -> str:
        """Подготовка контекста для ИИ"""
        try:
            # Ограничиваем размер текста для контекста
            max_context_length = 8000  # ~4000 токенов
            
            if len(text) > max_context_length:
                # Берем начало и конец документа
                half_length = max_context_length // 2
                text = text[:half_length] + "\n\n... [содержимое обрезано для экономии токенов] ...\n\n" + text[-half_length:]
            
            context = f"""📄 Файл: {file_name}

📋 Содержимое файла:
{text}

💬 Запрос пользователя: {user_message or 'Обработай этот файл'}

Пожалуйста, проанализируй содержимое файла и ответь на запрос пользователя. Если в файле есть таблицы, изображения или сложные структуры, опиши их подробно."""
            
            return context
            
        except Exception as e:
            logger.error(f"❌ [AI_CONTEXT] Error preparing AI context: {e}")
            return f"📄 Файл: {file_name}\n\n💬 Запрос: {user_message or 'Обработай этот файл'}"
    
    def search_chat_documents(self, query: str, limit: int = 5) -> List[Dict[str, Any]]:
        """Поиск по документам чата"""
        try:
            # Создаем эмбеддинг для запроса
            response = requests.post(
                f"{self.RAG_SERVICE_URL}/api/embeddings",
                json={'text': query},
                timeout=30
            )
            
            if response.status_code != 200:
                return []
            
            result = response.json()
            query_embedding = result.get('embedding', [])
            
            if not query_embedding:
                return []
            
            # Поиск в Qdrant
            search_response = requests.post(
                f"{self.QDRANT_URL}/collections/{self.CHAT_COLLECTION}/points/search",
                json={
                    'vector': query_embedding,
                    'limit': limit,
                    'with_payload': True
                },
                timeout=30
            )
            
            if search_response.status_code == 200:
                search_result = search_response.json()
                return search_result.get('result', [])
            else:
                return []
                
        except Exception as e:
            logger.error(f"❌ [CHAT_SEARCH] Error searching chat documents: {e}")
            return []
    
    def get_chat_documents_stats(self) -> Dict[str, Any]:
        """Получение статистики документов чата"""
        try:
            response = requests.get(f"{self.QDRANT_URL}/collections/{self.CHAT_COLLECTION}")
            
            if response.status_code == 200:
                result = response.json()
                return {
                    "total_documents": result['result']['points_count'],
                    "indexed_vectors": result['result']['indexed_vectors_count'],
                    "collection_status": result['result']['status'],
                    "timestamp": datetime.now().isoformat()
                }
            else:
                return {
                    "total_documents": 0,
                    "indexed_vectors": 0,
                    "collection_status": "unknown",
                    "timestamp": datetime.now().isoformat()
                }
                
        except Exception as e:
            logger.error(f"❌ [CHAT_STATS] Error getting chat documents stats: {e}")
            return {
                "total_documents": 0,
                "indexed_vectors": 0,
                "collection_status": "error",
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            }
