import json
import logging
from datetime import datetime
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)

class CalculationDOCXGenerator:
    """Генератор DOCX отчетов для инженерных расчетов"""
    
    def __init__(self):
        self.document = None
        self.default_font = 'Times New Roman'
        self.bold_font = 'Times New Roman'
        
    def generate_calculation_report(self, calculation_data: Dict[str, Any]) -> bytes:
        """Генерация DOCX отчета для расчета"""
        try:
            logger.info(f"📄 [DOCX_GENERATOR] Generating calculation report for: {calculation_data.get('name', 'Unknown')}")
            
            # Создаем новый документ
            self.document = Document()
            self._setup_styles()
            
            # Заголовок отчета
            self._create_header(calculation_data)
            
            # Информация о расчете
            self._create_calculation_info_section(calculation_data)
            
            # Параметры расчета
            self._create_parameters_section(calculation_data)
            
            # Результаты расчета
            if calculation_data.get('results'):
                self._create_results_section(calculation_data)
            
            # Заключение
            self._create_conclusion_section(calculation_data)
            
            # Сохраняем документ в буфер
            from io import BytesIO
            buffer = BytesIO()
            self.document.save(buffer)
            buffer.seek(0)
            docx_content = buffer.getvalue()
            buffer.close()
            
            logger.info(f"📄 [DOCX_GENERATOR] Calculation report generated successfully, size: {len(docx_content)} bytes")
            return docx_content
            
        except Exception as e:
            logger.error(f"❌ [DOCX_GENERATOR] Error generating calculation report: {e}")
            raise
    
    def _setup_styles(self):
        """Настройка стилей документа"""
        # Основные стили будут применяться при создании элементов
        pass
    
    def _create_header(self, calculation_data: Dict[str, Any]):
        """Создание заголовка отчета"""
        # Основной заголовок
        title = self.document.add_heading("ОТЧЕТ О ИНЖЕНЕРНОМ РАСЧЕТЕ", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о расчете
        info_table = self.document.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заполняем таблицу
        info_data = [
            ("Название расчета:", calculation_data.get('name', 'Не указано')),
            ("Тип расчета:", self._get_calculation_type_name(calculation_data.get('type', ''))),
            ("Категория:", self._get_calculation_category_name(calculation_data.get('category', ''))),
            ("Дата создания:", self._format_date(calculation_data.get('created_at'))),
            ("Автор:", calculation_data.get('author', 'Не указан')),
            ("Статус:", self._get_status_text(calculation_data.get('status', 'unknown')))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)
            
            # Стилизация
            info_table.cell(i, 0).paragraphs[0].runs[0].bold = True
            info_table.cell(i, 0).paragraphs[0].runs[0].font.name = self.bold_font
            info_table.cell(i, 1).paragraphs[0].runs[0].font.name = self.default_font
        
        self.document.add_paragraph()
    
    def _create_calculation_info_section(self, calculation_data: Dict[str, Any]):
        """Создание раздела с информацией о расчете"""
        self.document.add_heading("1. ИНФОРМАЦИЯ О РАСЧЕТЕ", 1)
        
        # Описание
        if calculation_data.get('description'):
            desc_para = self.document.add_paragraph()
            desc_para.add_run("Описание: ").bold = True
            desc_para.add_run(calculation_data['description'])
        
        # Применяемые нормы
        norms = self._get_applicable_norms(calculation_data.get('type', ''), calculation_data.get('category', ''))
        if norms:
            norms_para = self.document.add_paragraph()
            norms_para.add_run("Применяемые нормы: ").bold = True
            norms_para.add_run(", ".join(norms))
        
        # Время выполнения
        if calculation_data.get('execution_time'):
            time_para = self.document.add_paragraph()
            time_para.add_run("Время выполнения: ").bold = True
            time_para.add_run(f"{calculation_data['execution_time']:.3f} секунд")
        
        self.document.add_paragraph()
    
    def _create_parameters_section(self, calculation_data: Dict[str, Any]):
        """Создание раздела с параметрами расчета"""
        self.document.add_heading("2. ПАРАМЕТРЫ РАСЧЕТА", 1)
        
        parameters = calculation_data.get('parameters', {})
        if not parameters:
            self.document.add_paragraph("Параметры не указаны.")
            return
        
        # Создаем таблицу параметров
        params_table = self.document.add_table(rows=len(parameters) + 1, cols=3)
        params_table.style = 'Table Grid'
        
        # Заголовки таблицы
        header_cells = params_table.rows[0].cells
        header_cells[0].text = "Параметр"
        header_cells[1].text = "Значение"
        header_cells[2].text = "Единица измерения"
        
        # Стилизация заголовков
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.name = self.bold_font
        
        # Заполняем таблицу параметрами
        for i, (param_name, param_value) in enumerate(parameters.items(), 1):
            row = params_table.rows[i]
            row.cells[0].text = self._format_parameter_name(param_name)
            row.cells[1].text = str(param_value)
            row.cells[2].text = self._get_parameter_unit(param_name)
            
            # Стилизация
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = self.default_font
        
        self.document.add_paragraph()
    
    def _create_results_section(self, calculation_data: Dict[str, Any]):
        """Создание раздела с результатами расчета"""
        self.document.add_heading("3. РЕЗУЛЬТАТЫ РАСЧЕТА", 1)
        
        results = calculation_data.get('results', {})
        if not results:
            self.document.add_paragraph("Результаты расчета недоступны.")
            return
        
        # Создаем таблицу результатов
        results_table = self.document.add_table(rows=len(results) + 1, cols=3)
        results_table.style = 'Table Grid'
        
        # Заголовки таблицы
        header_cells = results_table.rows[0].cells
        header_cells[0].text = "Результат"
        header_cells[1].text = "Значение"
        header_cells[2].text = "Единица измерения"
        
        # Стилизация заголовков
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.name = self.bold_font
        
        # Заполняем таблицу результатами
        for i, (result_name, result_value) in enumerate(results.items(), 1):
            row = results_table.rows[i]
            row.cells[0].text = self._format_result_name(result_name)
            row.cells[1].text = self._format_result_value(result_value)
            row.cells[2].text = self._get_result_unit(result_name)
            
            # Стилизация
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = self.default_font
        
        self.document.add_paragraph()
    
    def _create_conclusion_section(self, calculation_data: Dict[str, Any]):
        """Создание раздела с заключением"""
        self.document.add_heading("4. ЗАКЛЮЧЕНИЕ", 1)
        
        # Статус расчета
        status = calculation_data.get('status', 'unknown')
        status_para = self.document.add_paragraph()
        status_para.add_run("Статус расчета: ").bold = True
        status_para.add_run(self._get_status_text(status))
        
        # Общие выводы
        conclusions_para = self.document.add_paragraph()
        conclusions_para.add_run("Общие выводы: ").bold = True
        
        conclusions = self._generate_conclusions(calculation_data)
        conclusions_para.add_run(conclusions)
        
        # Рекомендации
        recommendations = self._generate_recommendations(calculation_data)
        if recommendations:
            rec_para = self.document.add_paragraph()
            rec_para.add_run("Рекомендации: ").bold = True
            rec_para.add_run(recommendations)
        
        # Дата и подпись
        self.document.add_paragraph()
        self.document.add_paragraph()
        
        signature_para = self.document.add_paragraph()
        signature_para.add_run(f"Дата составления отчета: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def _get_calculation_type_name(self, calculation_type: str) -> str:
        """Получение названия типа расчета"""
        type_names = {
            'structural': 'Строительные конструкции',
            'foundation': 'Основания и фундаменты',
            'thermal': 'Теплотехнические расчеты',
            'ventilation': 'Вентиляция и кондиционирование',
            'electrical': 'Электротехнические расчеты',
            'water': 'Водоснабжение и водоотведение',
            'fire': 'Пожарная безопасность',
            'acoustic': 'Акустические расчеты',
            'lighting': 'Освещение и инсоляция',
            'geotechnical': 'Инженерно-геологические расчеты'
        }
        return type_names.get(calculation_type, calculation_type)
    
    def _get_calculation_category_name(self, category: str) -> str:
        """Получение названия категории расчета"""
        category_names = {
            'strength': 'Расчёт на прочность',
            'stability': 'Расчёт на устойчивость',
            'stiffness': 'Расчёт на жёсткость',
            'cracking': 'Расчёт на трещиностойкость',
            'dynamic': 'Динамический расчёт',
            'bearing_capacity': 'Расчёт несущей способности основания',
            'settlement': 'Расчёт осадок фундамента',
            'slope_stability': 'Расчёт устойчивости откосов',
            'pile_foundation': 'Расчёт свайных фундаментов',
            'retaining_wall': 'Расчёт подпорных стен',
            'heat_transfer': 'Расчёт теплопередачи через ограждения',
            'insulation': 'Расчёт теплоизоляции',
            'energy_efficiency': 'Расчёт энергоэффективности здания',
            'condensation': 'Расчёт конденсации влаги',
            'thermal_bridge': 'Расчёт тепловых мостов',
            'air_exchange': 'Расчёт воздухообмена',
            'duct_sizing': 'Расчёт воздуховодов',
            'fan_selection': 'Подбор вентиляторов',
            'cooling_load': 'Расчёт холодильной нагрузки',
            'humidity_control': 'Расчёт влажностного режима'
        }
        return category_names.get(category, category)
    
    def _get_applicable_norms(self, calculation_type: str, category: str) -> List[str]:
        """Получение списка применяемых норм"""
        norms_map = {
            'structural': {
                'strength': ['СП 63.13330', 'СП 16.13330', 'EN 1992', 'EN 1993'],
                'stability': ['СП 16.13330', 'СП 63.13330', 'EN 1993'],
                'stiffness': ['СП 63.13330', 'СП 64.13330', 'EN 1995'],
                'cracking': ['СП 63.13330', 'EN 1992'],
                'dynamic': ['СП 14.13330', 'EN 1998']
            },
            'foundation': {
                'bearing_capacity': ['СП 22.13330.2016', 'СП 24.13330.2011', 'СП 25.13330.2012'],
                'settlement': ['СП 22.13330.2016', 'СП 24.13330.2011'],
                'slope_stability': ['СП 22.13330.2016', 'СП 47.13330.2016'],
                'pile_foundation': ['СП 24.13330.2011', 'СП 25.13330.2012'],
                'retaining_wall': ['СП 22.13330.2016', 'СП 63.13330.2018']
            },
            'thermal': {
                'heat_transfer': ['СП 50.13330.2012', 'СП 23-101-2004', 'ГОСТ 30494-2011'],
                'insulation': ['СП 50.13330.2012', 'СП 23-101-2004'],
                'energy_efficiency': ['СП 50.13330.2012', 'СП 23-101-2004', 'ГОСТ 30494-2011'],
                'condensation': ['СП 50.13330.2012', 'СП 23-101-2004'],
                'thermal_bridge': ['СП 50.13330.2012', 'СП 23-101-2004']
            },
            'ventilation': {
                'air_exchange': ['СП 60.13330.2016', 'СП 7.13130.2013', 'СП 54.13330.2016'],
                'duct_sizing': ['СП 60.13330.2016', 'СП 7.13130.2013'],
                'fan_selection': ['СП 60.13330.2016', 'СП 7.13130.2013'],
                'cooling_load': ['СП 60.13330.2016', 'СП 7.13130.2013'],
                'humidity_control': ['СП 60.13330.2016', 'СП 7.13130.2013']
            }
        }
        
        return norms_map.get(calculation_type, {}).get(category, [])
    
    def _format_parameter_name(self, param_name: str) -> str:
        """Форматирование названия параметра"""
        param_names = {
            'load_value': 'Расчетная нагрузка',
            'section_area': 'Площадь сечения',
            'material_strength': 'Расчетное сопротивление материала',
            'safety_factor': 'Коэффициент надежности',
            'foundation_width': 'Ширина фундамента',
            'foundation_length': 'Длина фундамента',
            'foundation_depth': 'Глубина заложения',
            'soil_cohesion': 'Сцепление грунта',
            'soil_friction_angle': 'Угол внутреннего трения',
            'soil_density': 'Плотность грунта',
            'wall_thickness': 'Толщина стены',
            'wall_area': 'Площадь стены',
            'thermal_conductivity': 'Коэффициент теплопроводности',
            'indoor_temp': 'Внутренняя температура',
            'outdoor_temp': 'Наружная температура',
            'room_volume': 'Объем помещения',
            'room_area': 'Площадь помещения',
            'occupancy': 'Количество людей',
            'air_flow_rate': 'Расход воздуха',
            'duct_length': 'Длина воздуховода'
        }
        return param_names.get(param_name, param_name.replace('_', ' ').title())
    
    def _get_parameter_unit(self, param_name: str) -> str:
        """Получение единицы измерения параметра"""
        units = {
            'load_value': 'кН',
            'section_area': 'см²',
            'material_strength': 'МПа',
            'safety_factor': '',
            'foundation_width': 'м',
            'foundation_length': 'м',
            'foundation_depth': 'м',
            'soil_cohesion': 'кПа',
            'soil_friction_angle': 'град',
            'soil_density': 'т/м³',
            'wall_thickness': 'м',
            'wall_area': 'м²',
            'thermal_conductivity': 'Вт/(м·К)',
            'indoor_temp': '°C',
            'outdoor_temp': '°C',
            'room_volume': 'м³',
            'room_area': 'м²',
            'occupancy': 'чел',
            'air_flow_rate': 'м³/ч',
            'duct_length': 'м'
        }
        return units.get(param_name, '')
    
    def _format_result_name(self, result_name: str) -> str:
        """Форматирование названия результата"""
        result_names = {
            'bearing_capacity': 'Несущая способность',
            'safety_factor': 'Коэффициент запаса',
            'settlement': 'Осадка',
            'heat_transfer_coefficient': 'Коэффициент теплопередачи',
            'heat_loss': 'Тепловые потери',
            'thermal_resistance': 'Термическое сопротивление',
            'air_exchange_rate': 'Кратность воздухообмена',
            'pressure_loss': 'Потери давления',
            'fan_power': 'Мощность вентилятора',
            'cooling_capacity': 'Холодильная мощность'
        }
        return result_names.get(result_name, result_name.replace('_', ' ').title())
    
    def _format_result_value(self, result_value: Any) -> str:
        """Форматирование значения результата"""
        if isinstance(result_value, (int, float)):
            return f"{result_value:.3f}"
        return str(result_value)
    
    def _get_result_unit(self, result_name: str) -> str:
        """Получение единицы измерения результата"""
        units = {
            'bearing_capacity': 'кПа',
            'safety_factor': '',
            'settlement': 'мм',
            'heat_transfer_coefficient': 'Вт/(м²·К)',
            'heat_loss': 'Вт',
            'thermal_resistance': 'м²·К/Вт',
            'air_exchange_rate': '1/ч',
            'pressure_loss': 'Па',
            'fan_power': 'кВт',
            'cooling_capacity': 'кВт'
        }
        return units.get(result_name, '')
    
    def _get_status_text(self, status: str) -> str:
        """Получение текста статуса"""
        status_texts = {
            'created': 'Создан',
            'in_progress': 'Выполняется',
            'completed': 'Завершен',
            'error': 'Ошибка',
            'unknown': 'Неизвестно'
        }
        return status_texts.get(status, status)
    
    def _format_date(self, date_str: str) -> str:
        """Форматирование даты"""
        if not date_str:
            return 'Не указана'
        try:
            from datetime import datetime
            dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
            return dt.strftime('%d.%m.%Y %H:%M')
        except:
            return date_str
    
    def _generate_conclusions(self, calculation_data: Dict[str, Any]) -> str:
        """Генерация выводов по расчету"""
        status = calculation_data.get('status', 'unknown')
        results = calculation_data.get('results', {})
        
        if status == 'completed':
            if results:
                return "Расчет выполнен успешно. Все параметры рассчитаны в соответствии с действующими нормами."
            else:
                return "Расчет выполнен успешно."
        elif status == 'error':
            return "При выполнении расчета возникли ошибки. Требуется проверка входных данных."
        elif status == 'in_progress':
            return "Расчет выполняется в настоящее время."
        else:
            return "Расчет создан и ожидает выполнения."
    
    def _generate_recommendations(self, calculation_data: Dict[str, Any]) -> str:
        """Генерация рекомендаций по расчету"""
        status = calculation_data.get('status', 'unknown')
        results = calculation_data.get('results', {})
        
        if status == 'completed':
            return "Рекомендуется провести дополнительную проверку результатов расчета и при необходимости выполнить корректировку параметров."
        elif status == 'error':
            return "Необходимо проверить корректность входных данных и повторить расчет."
        else:
            return "Рекомендуется выполнить расчет для получения результатов."
