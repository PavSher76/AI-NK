"""
Модуль объединения документов по общему ШИФР проекта
"""

import logging
import os
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from pathlib import Path
import fitz  # PyMuPDF
import docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch

from .models import ArchiveDocument, DocumentSection, DocumentType
from .database_manager import ArchiveDatabaseManager

logger = logging.getLogger(__name__)

class DocumentMerger:
    """Класс для объединения документов по ШИФР проекта"""
    
    def __init__(self, db_manager: ArchiveDatabaseManager):
        self.db_manager = db_manager
    
    async def merge_project_documents(self, project_code: str, output_format: str = "pdf") -> Dict[str, Any]:
        """Объединение всех документов проекта в один файл"""
        try:
            logger.info(f"🔄 [MERGE_DOCUMENTS] Starting merge for project {project_code}")
            
            # Получаем документы проекта
            documents = self.db_manager.get_documents_by_project(project_code)
            if not documents:
                raise Exception(f"No documents found for project {project_code}")
            
            # Сортируем документы по типу и дате загрузки
            sorted_documents = self._sort_documents_for_merge(documents)
            
            # Создаем объединенный документ
            if output_format.lower() == "pdf":
                output_path = await self._merge_to_pdf(project_code, sorted_documents)
            elif output_format.lower() == "docx":
                output_path = await self._merge_to_docx(project_code, sorted_documents)
            else:
                raise Exception(f"Unsupported output format: {output_format}")
            
            # Создаем запись об объединенном документе
            merged_document = await self._create_merged_document_record(
                project_code, output_path, sorted_documents
            )
            
            logger.info(f"✅ [MERGE_DOCUMENTS] Merge completed for project {project_code}")
            
            return {
                "status": "success",
                "project_code": project_code,
                "merged_document_id": merged_document.id,
                "output_path": output_path,
                "total_documents": len(sorted_documents),
                "output_format": output_format
            }
            
        except Exception as e:
            logger.error(f"❌ [MERGE_DOCUMENTS] Error merging documents: {e}")
            raise
    
    def _sort_documents_for_merge(self, documents: List[ArchiveDocument]) -> List[ArchiveDocument]:
        """Сортировка документов для объединения"""
        # Порядок приоритета типов документов
        type_priority = {
            DocumentType.TEO: 1,
            DocumentType.PD: 2,
            DocumentType.RD: 3,
            DocumentType.SPECIFICATION: 4,
            DocumentType.CALCULATION: 5,
            DocumentType.REPORT: 6,
            DocumentType.DRAWING: 7,
            DocumentType.OTHER: 8
        }
        
        def sort_key(doc):
            priority = type_priority.get(doc.document_type, 9)
            return (priority, doc.upload_date or datetime.min)
        
        return sorted(documents, key=sort_key)
    
    async def _merge_to_pdf(self, project_code: str, documents: List[ArchiveDocument]) -> str:
        """Объединение документов в PDF"""
        try:
            # Создаем директорию для объединенных документов
            merge_dir = f"/app/uploads/merged/{project_code}"
            os.makedirs(merge_dir, exist_ok=True)
            
            # Создаем имя файла
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{project_code}_merged_{timestamp}.pdf"
            output_path = os.path.join(merge_dir, output_filename)
            
            # Создаем PDF документ
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            story = []
            
            # Стили для заголовков
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1  # Центрирование
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12,
                spaceBefore=20
            )
            
            # Добавляем титульную страницу
            story.append(Paragraph(f"Объединенная документация проекта", title_style))
            story.append(Paragraph(f"ШИФР проекта: {project_code}", styles['Heading2']))
            story.append(Paragraph(f"Дата объединения: {datetime.now().strftime('%d.%m.%Y %H:%M')}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Добавляем оглавление
            story.append(Paragraph("Содержание", heading_style))
            story.append(Spacer(1, 10))
            
            for i, doc in enumerate(documents, 1):
                doc_title = f"{i}. {doc.document_name} ({doc.document_type.value})"
                story.append(Paragraph(doc_title, styles['Normal']))
            
            story.append(PageBreak())
            
            # Добавляем содержимое каждого документа
            for i, doc in enumerate(documents, 1):
                logger.info(f"📄 [MERGE_PDF] Processing document {i}/{len(documents)}: {doc.document_name}")
                
                # Заголовок документа
                story.append(Paragraph(f"{i}. {doc.document_name}", heading_style))
                story.append(Paragraph(f"Тип: {doc.document_type.value}", styles['Normal']))
                if doc.document_number:
                    story.append(Paragraph(f"Номер: {doc.document_number}", styles['Normal']))
                if doc.author:
                    story.append(Paragraph(f"Автор: {doc.author}", styles['Normal']))
                if doc.version:
                    story.append(Paragraph(f"Версия: {doc.version}", styles['Normal']))
                story.append(Spacer(1, 20))
                
                # Добавляем содержимое документа
                if doc.file_path and os.path.exists(doc.file_path):
                    content = await self._extract_document_content(doc)
                    if content:
                        # Разбиваем на абзацы
                        paragraphs = content.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                story.append(Paragraph(para.strip(), styles['Normal']))
                                story.append(Spacer(1, 6))
                
                # Добавляем разделы документа
                sections = self.db_manager.get_document_sections(doc.id)
                if sections:
                    story.append(Paragraph("Разделы документа:", styles['Heading3']))
                    for section in sections:
                        if section.section_title:
                            story.append(Paragraph(f"• {section.section_title}", styles['Normal']))
                        if section.section_content:
                            # Ограничиваем длину содержимого раздела
                            content = section.section_content[:500] + "..." if len(section.section_content) > 500 else section.section_content
                            story.append(Paragraph(content, styles['Normal']))
                            story.append(Spacer(1, 6))
                
                story.append(PageBreak())
            
            # Создаем PDF
            doc.build(story)
            
            logger.info(f"✅ [MERGE_PDF] PDF created: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"❌ [MERGE_PDF] Error creating PDF: {e}")
            raise
    
    async def _merge_to_docx(self, project_code: str, documents: List[ArchiveDocument]) -> str:
        """Объединение документов в DOCX"""
        try:
            # Создаем директорию для объединенных документов
            merge_dir = f"/app/uploads/merged/{project_code}"
            os.makedirs(merge_dir, exist_ok=True)
            
            # Создаем имя файла
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{project_code}_merged_{timestamp}.docx"
            output_path = os.path.join(merge_dir, output_filename)
            
            # Создаем новый DOCX документ
            doc = docx.Document()
            
            # Добавляем титульную страницу
            title = doc.add_heading('Объединенная документация проекта', 0)
            doc.add_heading(f'ШИФР проекта: {project_code}', level=1)
            doc.add_paragraph(f'Дата объединения: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
            doc.add_page_break()
            
            # Добавляем оглавление
            doc.add_heading('Содержание', level=1)
            
            for i, document in enumerate(documents, 1):
                doc.add_paragraph(f'{i}. {document.document_name} ({document.document_type.value})')
            
            doc.add_page_break()
            
            # Добавляем содержимое каждого документа
            for i, document in enumerate(documents, 1):
                logger.info(f"📄 [MERGE_DOCX] Processing document {i}/{len(documents)}: {document.document_name}")
                
                # Заголовок документа
                doc.add_heading(f'{i}. {document.document_name}', level=1)
                doc.add_paragraph(f'Тип: {document.document_type.value}')
                if document.document_number:
                    doc.add_paragraph(f'Номер: {document.document_number}')
                if document.author:
                    doc.add_paragraph(f'Автор: {document.author}')
                if document.version:
                    doc.add_paragraph(f'Версия: {document.version}')
                
                # Добавляем содержимое документа
                if document.file_path and os.path.exists(document.file_path):
                    content = await self._extract_document_content(document)
                    if content:
                        # Разбиваем на абзацы
                        paragraphs = content.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                doc.add_paragraph(para.strip())
                
                # Добавляем разделы документа
                sections = self.db_manager.get_document_sections(document.id)
                if sections:
                    doc.add_heading('Разделы документа:', level=2)
                    for section in sections:
                        if section.section_title:
                            doc.add_paragraph(f'• {section.section_title}')
                        if section.section_content:
                            # Ограничиваем длину содержимого раздела
                            content = section.section_content[:500] + "..." if len(section.section_content) > 500 else section.section_content
                            doc.add_paragraph(content)
                
                doc.add_page_break()
            
            # Сохраняем документ
            doc.save(output_path)
            
            logger.info(f"✅ [MERGE_DOCX] DOCX created: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"❌ [MERGE_DOCX] Error creating DOCX: {e}")
            raise
    
    async def _extract_document_content(self, document: ArchiveDocument) -> str:
        """Извлечение содержимого документа"""
        try:
            if not document.file_path or not os.path.exists(document.file_path):
                return ""
            
            file_ext = Path(document.file_path).suffix.lower()
            
            if file_ext == '.pdf':
                return await self._extract_pdf_content(document.file_path)
            elif file_ext == '.docx':
                return await self._extract_docx_content(document.file_path)
            elif file_ext == '.txt':
                with open(document.file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                logger.warning(f"⚠️ [EXTRACT_CONTENT] Unsupported file type: {file_ext}")
                return ""
                
        except Exception as e:
            logger.error(f"❌ [EXTRACT_CONTENT] Error extracting content: {e}")
            return ""
    
    async def _extract_pdf_content(self, file_path: str) -> str:
        """Извлечение содержимого из PDF"""
        try:
            doc = fitz.open(file_path)
            text_content = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content += page.get_text()
            
            doc.close()
            return text_content
            
        except Exception as e:
            logger.error(f"❌ [EXTRACT_PDF] Error extracting PDF content: {e}")
            return ""
    
    async def _extract_docx_content(self, file_path: str) -> str:
        """Извлечение содержимого из DOCX"""
        try:
            doc = docx.Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            return text_content
            
        except Exception as e:
            logger.error(f"❌ [EXTRACT_DOCX] Error extracting DOCX content: {e}")
            return ""
    
    async def _create_merged_document_record(self, project_code: str, output_path: str, 
                                           source_documents: List[ArchiveDocument]) -> ArchiveDocument:
        """Создание записи об объединенном документе"""
        try:
            # Создаем метаданные объединенного документа
            metadata = {
                "merged_document": True,
                "source_documents": [doc.id for doc in source_documents],
                "merge_date": datetime.now().isoformat(),
                "total_source_documents": len(source_documents)
            }
            
            # Создаем объект объединенного документа
            merged_document = ArchiveDocument(
                project_code=project_code,
                document_type=DocumentType.REPORT,
                document_name=f"Объединенная документация проекта {project_code}",
                document_number=f"MERGED-{project_code}-{datetime.now().strftime('%Y%m%d')}",
                original_filename=os.path.basename(output_path),
                file_type=Path(output_path).suffix.lower(),
                file_size=os.path.getsize(output_path),
                file_path=output_path,
                document_hash=self._calculate_file_hash(output_path),
                token_count=0,  # Будет рассчитано при обработке
                upload_date=datetime.now(),
                processing_status=ProcessingStatus.PENDING,
                author="System",
                department="Archive Service",
                version="1.0",
                status="active",
                metadata=metadata
            )
            
            # Сохраняем в базу данных
            document_id = self.db_manager.save_document(merged_document)
            merged_document.id = document_id
            
            logger.info(f"✅ [CREATE_MERGED_RECORD] Merged document record created: {document_id}")
            return merged_document
            
        except Exception as e:
            logger.error(f"❌ [CREATE_MERGED_RECORD] Error creating merged document record: {e}")
            raise
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Вычисление хеша файла"""
        import hashlib
        try:
            hash_sha256 = hashlib.sha256()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
        except Exception as e:
            logger.error(f"❌ [CALCULATE_HASH] Error calculating file hash: {e}")
            return ""
    
    async def get_merge_options(self, project_code: str) -> Dict[str, Any]:
        """Получение опций объединения для проекта"""
        try:
            # Получаем документы проекта
            documents = self.db_manager.get_documents_by_project(project_code)
            
            if not documents:
                return {
                    "can_merge": False,
                    "reason": "No documents found for project"
                }
            
            # Анализируем типы документов
            document_types = {}
            total_size = 0
            
            for doc in documents:
                doc_type = doc.document_type.value
                if doc_type not in document_types:
                    document_types[doc_type] = 0
                document_types[doc_type] += 1
                total_size += doc.file_size or 0
            
            # Проверяем, можно ли объединить
            can_merge = len(documents) > 1
            
            return {
                "can_merge": can_merge,
                "project_code": project_code,
                "total_documents": len(documents),
                "document_types": document_types,
                "total_size": total_size,
                "supported_formats": ["pdf", "docx"],
                "estimated_merge_time": len(documents) * 2  # Примерная оценка в секундах
            }
            
        except Exception as e:
            logger.error(f"❌ [GET_MERGE_OPTIONS] Error getting merge options: {e}")
            return {
                "can_merge": False,
                "reason": f"Error: {str(e)}"
            }
