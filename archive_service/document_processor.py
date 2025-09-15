"""
Процессор документов для архива технической документации
"""

import logging
import hashlib
import os
import re
import tempfile
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import fitz  # PyMuPDF
import docx
import pandas as pd
from pathlib import Path

from .models import ArchiveDocument, DocumentSection, DocumentType
from .config import PROJECT_CODE_PATTERNS, CHUNK_SIZE, CHUNK_OVERLAP

logger = logging.getLogger(__name__)

class ArchiveDocumentProcessor:
    """Процессор документов для архива технической документации"""
    
    def __init__(self):
        self.project_code_patterns = [re.compile(pattern) for pattern in PROJECT_CODE_PATTERNS]
    
    def extract_project_code(self, filename: str, content: str = "") -> Optional[str]:
        """Извлечение ШИФР проекта из имени файла или содержимого"""
        try:
            # Сначала ищем в имени файла
            for pattern in self.project_code_patterns:
                match = pattern.search(filename)
                if match:
                    logger.info(f"🔍 [EXTRACT_PROJECT_CODE] Found project code in filename: {match.group()}")
                    return match.group()
            
            # Затем ищем в содержимом документа
            if content:
                for pattern in self.project_code_patterns:
                    match = pattern.search(content)
                    if match:
                        logger.info(f"🔍 [EXTRACT_PROJECT_CODE] Found project code in content: {match.group()}")
                        return match.group()
            
            logger.warning(f"⚠️ [EXTRACT_PROJECT_CODE] No project code found in {filename}")
            return None
            
        except Exception as e:
            logger.error(f"❌ [EXTRACT_PROJECT_CODE] Error extracting project code: {e}")
            return None
    
    def detect_document_type(self, filename: str, content: str = "") -> DocumentType:
        """Определение типа документа по имени файла и содержимому"""
        try:
            filename_lower = filename.lower()
            
            # Определяем по имени файла
            if any(keyword in filename_lower for keyword in ['пд', 'проект', 'project']):
                return DocumentType.PD
            elif any(keyword in filename_lower for keyword in ['рд', 'рабоч', 'working']):
                return DocumentType.RD
            elif any(keyword in filename_lower for keyword in ['тэо', 'teo', 'обоснование']):
                return DocumentType.TEO
            elif any(keyword in filename_lower for keyword in ['чертеж', 'drawing', 'dwg']):
                return DocumentType.DRAWING
            elif any(keyword in filename_lower for keyword in ['спецификация', 'specification', 'spec']):
                return DocumentType.SPECIFICATION
            elif any(keyword in filename_lower for keyword in ['расчет', 'calculation', 'calc']):
                return DocumentType.CALCULATION
            elif any(keyword in filename_lower for keyword in ['отчет', 'report']):
                return DocumentType.REPORT
            
            # Определяем по содержимому
            if content:
                content_lower = content.lower()
                if any(keyword in content_lower for keyword in ['проектная документация', 'project documentation']):
                    return DocumentType.PD
                elif any(keyword in content_lower for keyword in ['рабочая документация', 'working documentation']):
                    return DocumentType.RD
                elif any(keyword in content_lower for keyword in ['технико-экономическое обоснование', 'teo']):
                    return DocumentType.TEO
                elif any(keyword in content_lower for keyword in ['спецификация', 'specification']):
                    return DocumentType.SPECIFICATION
                elif any(keyword in content_lower for keyword in ['расчет', 'calculation']):
                    return DocumentType.CALCULATION
                elif any(keyword in content_lower for keyword in ['отчет', 'report']):
                    return DocumentType.REPORT
            
            logger.info(f"🔍 [DETECT_DOCUMENT_TYPE] Using default type for {filename}")
            return DocumentType.OTHER
            
        except Exception as e:
            logger.error(f"❌ [DETECT_DOCUMENT_TYPE] Error detecting document type: {e}")
            return DocumentType.OTHER
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Извлечение текста из PDF файла"""
        try:
            text_content = ""
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content += page.get_text()
            
            doc.close()
            logger.info(f"✅ [EXTRACT_PDF] Extracted text from PDF: {len(text_content)} characters")
            return text_content
            
        except Exception as e:
            logger.error(f"❌ [EXTRACT_PDF] Error extracting text from PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Извлечение текста из DOCX файла"""
        try:
            doc = docx.Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            logger.info(f"✅ [EXTRACT_DOCX] Extracted text from DOCX: {len(text_content)} characters")
            return text_content
            
        except Exception as e:
            logger.error(f"❌ [EXTRACT_DOCX] Error extracting text from DOCX: {e}")
            return ""
    
    def extract_text_from_xlsx(self, file_path: str) -> str:
        """Извлечение текста из XLSX файла"""
        try:
            text_content = ""
            df = pd.read_excel(file_path, sheet_name=None)
            
            for sheet_name, sheet_df in df.items():
                text_content += f"Лист: {sheet_name}\n"
                text_content += sheet_df.to_string(index=False) + "\n\n"
            
            logger.info(f"✅ [EXTRACT_XLSX] Extracted text from XLSX: {len(text_content)} characters")
            return text_content
            
        except Exception as e:
            logger.error(f"❌ [EXTRACT_XLSX] Error extracting text from XLSX: {e}")
            return ""
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Извлечение текста из файла по типу"""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                return self.extract_text_from_docx(file_path)
            elif file_extension == '.xlsx':
                return self.extract_text_from_xlsx(file_path)
            elif file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                logger.warning(f"⚠️ [EXTRACT_TEXT] Unsupported file type: {file_extension}")
                return ""
                
        except Exception as e:
            logger.error(f"❌ [EXTRACT_TEXT] Error extracting text from file: {e}")
            return ""
    
    def extract_sections_from_text(self, text: str, document_id: int) -> List[DocumentSection]:
        """Извлечение разделов из текста документа"""
        try:
            sections = []
            
            # Разбиваем текст на абзацы
            paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
            
            current_section = None
            section_number = 1
            page_number = 1
            
            for i, paragraph in enumerate(paragraphs):
                # Определяем, является ли абзац заголовком раздела
                is_header = self._is_section_header(paragraph)
                
                if is_header:
                    # Сохраняем предыдущий раздел
                    if current_section:
                        sections.append(current_section)
                    
                    # Создаем новый раздел
                    section_title = paragraph.strip()
                    section_number_str = self._extract_section_number(section_title)
                    
                    current_section = DocumentSection(
                        archive_document_id=document_id,
                        section_number=section_number_str or str(section_number),
                        section_title=section_title,
                        section_content="",
                        page_number=page_number,
                        section_type="header",
                        importance_level=self._determine_importance_level(section_title)
                    )
                    section_number += 1
                else:
                    # Добавляем содержимое к текущему разделу
                    if current_section:
                        if current_section.section_content:
                            current_section.section_content += "\n" + paragraph
                        else:
                            current_section.section_content = paragraph
                    else:
                        # Создаем раздел для содержимого без заголовка
                        current_section = DocumentSection(
                            archive_document_id=document_id,
                            section_number=str(section_number),
                            section_title="Введение",
                            section_content=paragraph,
                            page_number=page_number,
                            section_type="text",
                            importance_level=1
                        )
                        section_number += 1
                
                # Примерная оценка номера страницы (примерно 50 строк на страницу)
                if i % 50 == 0:
                    page_number += 1
            
            # Добавляем последний раздел
            if current_section:
                sections.append(current_section)
            
            logger.info(f"✅ [EXTRACT_SECTIONS] Extracted {len(sections)} sections from document {document_id}")
            return sections
            
        except Exception as e:
            logger.error(f"❌ [EXTRACT_SECTIONS] Error extracting sections: {e}")
            return []
    
    def _is_section_header(self, text: str) -> bool:
        """Определение, является ли текст заголовком раздела"""
        # Простые эвристики для определения заголовков
        text = text.strip()
        
        # Проверяем на наличие номера раздела
        if re.match(r'^\d+\.?\s+', text):
            return True
        
        # Проверяем на ключевые слова заголовков
        header_keywords = [
            'введение', 'введение', 'общие положения', 'технические требования',
            'методы испытаний', 'правила приемки', 'транспортирование', 'хранение',
            'гарантии изготовителя', 'приложение', 'литература', 'содержание'
        ]
        
        text_lower = text.lower()
        return any(keyword in text_lower for keyword in header_keywords)
    
    def _extract_section_number(self, text: str) -> Optional[str]:
        """Извлечение номера раздела из текста"""
        match = re.match(r'^(\d+(?:\.\d+)*)', text.strip())
        return match.group(1) if match else None
    
    def _determine_importance_level(self, text: str) -> int:
        """Определение уровня важности раздела"""
        text_lower = text.lower()
        
        # Критически важные разделы
        if any(keyword in text_lower for keyword in ['требования безопасности', 'критические параметры', 'аварийные ситуации']):
            return 5
        
        # Очень важные разделы
        if any(keyword in text_lower for keyword in ['технические требования', 'методы испытаний', 'гарантии']):
            return 4
        
        # Важные разделы
        if any(keyword in text_lower for keyword in ['общие положения', 'правила приемки', 'транспортирование']):
            return 3
        
        # Средней важности
        if any(keyword in text_lower for keyword in ['введение', 'приложение', 'литература']):
            return 2
        
        # Низкой важности
        return 1
    
    def create_document_chunks(self, text: str, document_id: int) -> List[Dict[str, Any]]:
        """Создание чанков из текста документа"""
        try:
            chunks = []
            
            # Разбиваем текст на предложения
            sentences = re.split(r'[.!?]+', text)
            sentences = [s.strip() for s in sentences if s.strip()]
            
            current_chunk = ""
            chunk_number = 1
            
            for sentence in sentences:
                # Если добавление предложения превысит размер чанка
                if len(current_chunk) + len(sentence) > CHUNK_SIZE:
                    if current_chunk:
                        chunks.append({
                            'chunk_id': f"{document_id}_{chunk_number}",
                            'content': current_chunk.strip(),
                            'document_id': document_id,
                            'chunk_number': chunk_number
                        })
                        chunk_number += 1
                    
                    current_chunk = sentence
                else:
                    if current_chunk:
                        current_chunk += ". " + sentence
                    else:
                        current_chunk = sentence
            
            # Добавляем последний чанк
            if current_chunk:
                chunks.append({
                    'chunk_id': f"{document_id}_{chunk_number}",
                    'content': current_chunk.strip(),
                    'document_id': document_id,
                    'chunk_number': chunk_number
                })
            
            logger.info(f"✅ [CREATE_CHUNKS] Created {len(chunks)} chunks for document {document_id}")
            return chunks
            
        except Exception as e:
            logger.error(f"❌ [CREATE_CHUNKS] Error creating chunks: {e}")
            return []
    
    def calculate_file_hash(self, file_path: str) -> str:
        """Вычисление хеша файла"""
        try:
            hash_sha256 = hashlib.sha256()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
        except Exception as e:
            logger.error(f"❌ [CALCULATE_HASH] Error calculating file hash: {e}")
            return ""
    
    def process_document(self, file_path: str, project_code: str = None) -> Tuple[ArchiveDocument, List[DocumentSection]]:
        """Полная обработка документа"""
        try:
            logger.info(f"🔄 [PROCESS_DOCUMENT] Processing document: {file_path}")
            
            # Извлекаем текст
            text_content = self.extract_text_from_file(file_path)
            if not text_content:
                raise Exception("Failed to extract text from document")
            
            # Определяем ШИФР проекта
            if not project_code:
                project_code = self.extract_project_code(os.path.basename(file_path), text_content)
                if not project_code:
                    raise Exception("Project code not found and not provided")
            
            # Определяем тип документа
            document_type = self.detect_document_type(os.path.basename(file_path), text_content)
            
            # Создаем объект документа
            document = ArchiveDocument(
                project_code=project_code,
                document_type=document_type,
                document_name=os.path.basename(file_path),
                original_filename=os.path.basename(file_path),
                file_type=Path(file_path).suffix.lower(),
                file_size=os.path.getsize(file_path),
                file_path=file_path,
                document_hash=self.calculate_file_hash(file_path),
                token_count=len(text_content.split()),
                upload_date=datetime.now(),
                processing_status=ProcessingStatus.PENDING
            )
            
            # Извлекаем разделы
            sections = self.extract_sections_from_text(text_content, 0)  # ID будет установлен при сохранении
            
            logger.info(f"✅ [PROCESS_DOCUMENT] Document processed successfully: {len(sections)} sections")
            return document, sections
            
        except Exception as e:
            logger.error(f"❌ [PROCESS_DOCUMENT] Error processing document: {e}")
            raise
