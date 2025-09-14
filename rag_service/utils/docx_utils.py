"""
Модуль для извлечения текста из DOCX документов
Содержит функции извлечения текста с поддержкой структуры документа
"""

import logging
import io
from typing import Dict, Any, List, Optional
from pathlib import Path

# Импорты для работы с DOCX
try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class DOCXTextExtractor:
    """Класс для извлечения текста из DOCX документов"""
    
    def __init__(self, extract_tables: bool = True, extract_headers: bool = True):
        """
        Инициализация экстрактора
        
        Args:
            extract_tables: Извлекать текст из таблиц
            extract_headers: Извлекать заголовки отдельно
        """
        self.extract_tables = extract_tables
        self.extract_headers = extract_headers
        
        if not DOCX_AVAILABLE:
            raise ImportError("Необходимо установить python-docx для работы с DOCX")
    
    def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """
        Извлечение текста из DOCX файла
        
        Args:
            file_path: Путь к DOCX файлу
            
        Returns:
            Словарь с извлеченным текстом и метаданными
        """
        try:
            doc = Document(file_path)
            return self._extract_from_document(doc, file_path)
                
        except Exception as e:
            logger.error(f"❌ [DOCX_EXTRACT] Ошибка извлечения текста из {file_path}: {e}")
            raise Exception(f"Ошибка извлечения текста из DOCX: {e}")
    
    def extract_text_from_bytes(self, file_content: bytes) -> Dict[str, Any]:
        """
        Извлечение текста из DOCX в виде байтов
        
        Args:
            file_content: Содержимое DOCX файла в байтах
            
        Returns:
            Словарь с извлеченным текстом и метаданными
        """
        try:
            doc = Document(io.BytesIO(file_content))
            return self._extract_from_document(doc, "bytes")
                
        except Exception as e:
            logger.error(f"❌ [DOCX_EXTRACT] Ошибка извлечения текста из байтов: {e}")
            raise Exception(f"Ошибка извлечения текста из DOCX байтов: {e}")
    
    def _extract_from_document(self, doc: DocumentType, source: str) -> Dict[str, Any]:
        """
        Извлечение текста из объекта Document
        
        Args:
            doc: Объект Document
            source: Источник документа (путь или "bytes")
            
        Returns:
            Словарь с извлеченным текстом и метаданными
        """
        try:
            paragraphs = []
            headers = []
            tables = []
            full_text = ""
            
            # Извлекаем параграфы
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                para_text = paragraph.text.strip()
                if not para_text:
                    continue
                
                # Определяем тип параграфа
                para_type = self._get_paragraph_type(paragraph)
                
                para_data = {
                    "paragraph_number": para_num,
                    "text": para_text,
                    "type": para_type,
                    "style": paragraph.style.name if paragraph.style else "Normal",
                    "char_count": len(para_text),
                    "word_count": len(para_text.split())
                }
                
                paragraphs.append(para_data)
                
                # Если это заголовок, добавляем в отдельный список
                if self.extract_headers and para_type in ["header", "title"]:
                    headers.append(para_data)
                
                full_text += para_text + "\n"
            
            # Извлекаем таблицы
            if self.extract_tables:
                for table_num, table in enumerate(doc.tables, 1):
                    table_data = self._extract_table_data(table, table_num)
                    tables.append(table_data)
                    
                    # Добавляем текст таблицы к общему тексту
                    table_text = self._table_to_text(table_data)
                    full_text += table_text + "\n"
            
            # Очищаем извлеченный текст
            cleaned_text = self._clean_extracted_text(full_text)
            
            logger.info(f"📄 [DOCX] Извлечено {len(paragraphs)} параграфов, {len(tables)} таблиц из DOCX")
            
            return {
                "success": True,
                "text": cleaned_text,
                "paragraphs": paragraphs,
                "headers": headers,
                "tables": tables,
                "method": "python-docx",
                "metadata": {
                    "total_paragraphs": len(paragraphs),
                    "total_headers": len(headers),
                    "total_tables": len(tables),
                    "total_chars": len(cleaned_text),
                    "total_words": len(cleaned_text.split()),
                    "source": source
                }
            }
            
        except Exception as e:
            logger.error(f"❌ [DOCX_EXTRACT] Ошибка обработки документа: {str(e)}")
            raise Exception(f"Ошибка обработки DOCX документа: {str(e)}")
    
    def _get_paragraph_type(self, paragraph: Paragraph) -> str:
        """
        Определение типа параграфа
        
        Args:
            paragraph: Объект параграфа
            
        Returns:
            Тип параграфа
        """
        if not paragraph.style:
            return "normal"
        
        style_name = paragraph.style.name.lower()
        
        if "heading" in style_name or "title" in style_name:
            return "header"
        elif "title" in style_name:
            return "title"
        elif "caption" in style_name:
            return "caption"
        elif "list" in style_name:
            return "list"
        else:
            return "normal"
    
    def _extract_table_data(self, table: Table, table_num: int) -> Dict[str, Any]:
        """
        Извлечение данных из таблицы
        
        Args:
            table: Объект таблицы
            table_num: Номер таблицы
            
        Returns:
            Словарь с данными таблицы
        """
        try:
            rows_data = []
            table_text = ""
            
            for row_num, row in enumerate(table.rows):
                row_data = []
                row_text = ""
                
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                    row_text += cell_text + "\t"
                
                rows_data.append({
                    "row_number": row_num + 1,
                    "cells": row_data,
                    "text": row_text.strip()
                })
                table_text += row_text.strip() + "\n"
            
            return {
                "table_number": table_num,
                "rows": rows_data,
                "text": table_text.strip(),
                "row_count": len(rows_data),
                "col_count": len(rows_data[0]["cells"]) if rows_data else 0
            }
            
        except Exception as e:
            logger.error(f"❌ [DOCX_TABLE] Ошибка извлечения таблицы {table_num}: {e}")
            return {
                "table_number": table_num,
                "rows": [],
                "text": "",
                "row_count": 0,
                "col_count": 0,
                "error": str(e)
            }
    
    def _table_to_text(self, table_data: Dict[str, Any]) -> str:
        """
        Преобразование данных таблицы в текст
        
        Args:
            table_data: Данные таблицы
            
        Returns:
            Текстовое представление таблицы
        """
        if not table_data.get("rows"):
            return ""
        
        text_lines = []
        for row in table_data["rows"]:
            text_lines.append(row["text"])
        
        return "\n".join(text_lines)
    
    def _clean_extracted_text(self, text: str) -> str:
        """
        Очистка извлеченного текста от лишних пробелов и символов
        
        Args:
            text: Исходный текст
            
        Returns:
            Очищенный текст
        """
        import re
        
        # Удаляем невидимые символы и специальные пробелы
        text = re.sub(r'[\u00A0\u2000-\u200F\u2028-\u202F\u205F\u3000]', ' ', text)
        
        # Удаляем множественные пробелы в строках, но сохраняем переносы строк
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Удаляем пробелы в начале и конце строк
        lines = text.split('\n')
        lines = [line.strip() for line in lines]
        text = '\n'.join(lines)
        
        # Удаляем лишние переносы строк (более 2 подряд)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Удаляем пробелы перед знаками препинания
        text = re.sub(r'\s+([.,!?;:])', r'\1', text)
        
        # Удаляем пробелы после открывающих скобок и перед закрывающими
        text = re.sub(r'\(\s+', '(', text)
        text = re.sub(r'\s+\)', ')', text)
        
        # Удаляем пробелы в кавычках
        text = re.sub(r'"\s+', '"', text)
        text = re.sub(r'\s+"', '"', text)
        
        return text.strip()
    
    def create_chunks(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[Dict[str, Any]]:
        """
        Разделение текста на чанки для обработки
        
        Args:
            text: Исходный текст
            chunk_size: Размер чанка в символах
            overlap: Перекрытие между чанками
            
        Returns:
            Список чанков с метаданными
        """
        chunks = []
        words = text.split()
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            chunks.append({
                'content': chunk_text,
                'start_word': i,
                'end_word': min(i + chunk_size, len(words)),
                'char_count': len(chunk_text),
                'word_count': len(chunk_words)
            })
        
        return chunks


# Функции для обратной совместимости
def extract_text_from_docx_file(file_path: str, extract_tables: bool = True, extract_headers: bool = True) -> Dict[str, Any]:
    """
    Извлечение текста из DOCX файла (функция для обратной совместимости)
    
    Args:
        file_path: Путь к DOCX файлу
        extract_tables: Извлекать текст из таблиц
        extract_headers: Извлекать заголовки отдельно
        
    Returns:
        Словарь с извлеченным текстом и метаданными
    """
    extractor = DOCXTextExtractor(extract_tables=extract_tables, extract_headers=extract_headers)
    return extractor.extract_text_from_file(file_path)


def extract_text_from_docx_bytes(file_content: bytes, extract_tables: bool = True, extract_headers: bool = True) -> Dict[str, Any]:
    """
    Извлечение текста из DOCX в виде байтов (функция для обратной совместимости)
    
    Args:
        file_content: Содержимое DOCX файла в байтах
        extract_tables: Извлекать текст из таблиц
        extract_headers: Извлекать заголовки отдельно
        
    Returns:
        Словарь с извлеченным текстом и метаданными
    """
    extractor = DOCXTextExtractor(extract_tables=extract_tables, extract_headers=extract_headers)
    return extractor.extract_text_from_bytes(file_content)
