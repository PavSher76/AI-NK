import json
import logging
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

logger = logging.getLogger(__name__)

class DOCXReportGenerator:
    def __init__(self):
        self.document = None
        self._setup_styles()
    
    def _setup_styles(self):
        """Настройка стилей для DOCX документа"""
        # Стили будут применяться при создании документа
        pass
    
    def _parse_json_string(self, json_str):
        """Парсинг JSON строки, словаря или Python repr строки"""
        try:
            if isinstance(json_str, str):
                # Пробуем сначала как JSON
                try:
                    return json.loads(json_str)
                except json.JSONDecodeError:
                    # Если не JSON, пробуем как Python repr (заменяем одинарные кавычки на двойные)
                    try:
                        # Заменяем одинарные кавычки на двойные для JSON совместимости
                        json_compatible = json_str.replace("'", '"')
                        return json.loads(json_compatible)
                    except json.JSONDecodeError:
                        logger.warning(f"Failed to parse string as JSON or Python repr: {json_str[:100]}...")
                        return {}
            elif isinstance(json_str, dict):
                return json_str
            else:
                logger.warning(f"Unexpected data type: {type(json_str)}")
                return {}
        except Exception as e:
            logger.warning(f"Failed to parse data: {e}")
            return {}
    
    def _add_heading(self, text, level=1):
        """Добавление заголовка"""
        heading = self.document.add_heading(text, level)
        if level == 1:
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Устанавливаем цвет заголовка
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 51, 102)  # Темно-синий
        return heading
    
    def _add_paragraph(self, text, style='Normal'):
        """Добавление параграфа"""
        paragraph = self.document.add_paragraph(text)
        if style == 'Normal':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return paragraph
    
    def _create_table(self, data, headers=None):
        """Создание таблицы"""
        if headers:
            table_data = [headers] + data
        else:
            table_data = data
        
        table = self.document.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Заполняем таблицу данными
        for i, row_data in enumerate(table_data):
            for j, cell_data in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = str(cell_data)
                
                # Стилизация заголовка таблицы
                if i == 0 and headers:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)  # Белый текст
                    # Устанавливаем фон заголовка
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:val'), 'clear')
                    shading_elm.set(qn('w:color'), 'auto')
                    shading_elm.set(qn('w:fill'), '4472C4')  # Синий фон
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                else:
                    # Стилизация обычных ячеек
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
        
        return table
    
    def _create_header(self, document_info):
        """Создание заголовка отчета"""
        # Основной заголовок
        self._add_heading("ОТЧЕТ О ПРОВЕРКЕ НОРМОКОНТРОЛЯ", 1)
        
        # Информация о документе
        self._add_heading("Информация о документе", 2)
        
        doc_data = [
            ['Название файла', document_info.get('original_filename', 'Не указано')],
            ['ID документа', str(document_info.get('id', 'Не указано'))],
            ['Статус обработки', document_info.get('processing_status', 'Не указано')],
            ['Дата создания отчета', datetime.now().strftime('%d.%m.%Y %H:%M:%S')]
        ]
        
        self._create_table(doc_data, ['Параметр', 'Значение'])
        
        # Добавляем отступ
        self.document.add_paragraph()
    
    def _create_project_info_section(self, project_info):
        """Создание раздела с информацией о проекте"""
        if not project_info:
            return
        
        self._add_heading("Информация о проекте", 2)
        
        project_data = self._parse_json_string(project_info)
        
        project_table_data = [
            ['Название проекта', project_data.get('project_name', 'Не указано')],
            ['Стадия проектирования', project_data.get('project_stage', 'Не указано')],
            ['Тип проекта', project_data.get('project_type', 'Не указано')],
            ['Марка комплекта', project_data.get('document_set', 'Не указано')],
            ['Уверенность анализа', f"{project_data.get('confidence', 0) * 100:.1f}%"]
        ]
        
        self._create_table(project_table_data, ['Параметр', 'Значение'])
        self.document.add_paragraph()
    
    def _create_compliance_summary_section(self, compliance_summary):
        """Создание раздела с сводкой по соответствию"""
        if not compliance_summary:
            return
        
        self._add_heading("Сводка по соответствию нормам", 2)
        
        compliance_data = self._parse_json_string(compliance_summary)
        
        # Добавляем информацию о том, на соответствие каким документам выполнялась проверка
        project_stage = compliance_data.get('project_stage', 'Неизвестная стадия')
        document_set = compliance_data.get('document_set', 'Неизвестный комплект')
        
        compliance_info = f"Документ проверен на соответствие: {project_stage}, {document_set}"
        self._add_paragraph(compliance_info)
        self.document.add_paragraph()
        
        # Добавляем перечень НТД, на соответствие которым выполнена проверка
        relevant_ntd = compliance_data.get('relevant_ntd', [])
        if relevant_ntd:
            self._add_heading("Перечень нормативных документов", 3)
            self._add_paragraph("Проверка выполнена на соответствие следующим нормативным документам согласно \"Перечень НТД для руководства внутри ОНК по маркам\":")
            self.document.add_paragraph()
            
            # Создаем таблицу с перечнем НТД
            ntd_data = []
            for i, ntd in enumerate(relevant_ntd, 1):
                ntd_data.append([str(i), ntd])
            
            self._create_table(ntd_data, ['№', 'Наименование нормативного документа'])
            self.document.add_paragraph()
        else:
            # Если перечень НТД не указан, добавляем информацию об этом
            self._add_heading("Перечень нормативных документов", 3)
            self._add_paragraph("Перечень нормативных документов для проверки не определен. Рекомендуется загрузить документ \"Перечень НТД для руководства внутри ОНК по маркам\" для автоматического определения релевантных НТД.")
            self.document.add_paragraph()
        
        # Общая статистика
        total_pages = compliance_data.get('total_pages', 0)
        a4_equivalent = self._calculate_a4_equivalent(total_pages, compliance_data.get('page_sizes', []))
        
        stats_data = [
            ['Страниц в документе', str(total_pages)],
            ['Листов формата А4', str(a4_equivalent)],
            ['Соответствующих страниц', str(compliance_data.get('compliant_pages', 0))],
            ['Процент соответствия', f"{compliance_data.get('compliance_percentage', 0):.1f}%"],
            ['Всего находок', str(compliance_data.get('total_findings', 0))],
            ['Критических находок', str(compliance_data.get('critical_findings', 0))],
            ['Предупреждений', str(compliance_data.get('warning_findings', 0))],
            ['Информационных находок', str(compliance_data.get('info_findings', 0))]
        ]
        
        self._create_table(stats_data, ['Параметр', 'Значение'])
        self.document.add_paragraph()
        
        # Детальные находки
        findings = compliance_data.get('findings', [])
        if findings:
            self._add_heading("Детальные находки", 3)
            
            findings_data = []
            for finding in findings[:10]:  # Ограничиваем первыми 10 находками
                findings_data.append([
                    finding.get('type', ''),
                    finding.get('title', '')[:50] + '...' if len(finding.get('title', '')) > 50 else finding.get('title', ''),
                    finding.get('description', '')[:60] + '...' if len(finding.get('description', '')) > 60 else finding.get('description', ''),
                    finding.get('recommendation', '')[:60] + '...' if len(finding.get('recommendation', '')) > 60 else finding.get('recommendation', ''),
                    str(finding.get('page_number', ''))
                ])
            
            self._create_table(findings_data, ['Тип', 'Заголовок', 'Описание', 'Рекомендация', 'Страница'])
            self.document.add_paragraph()
    
    def _create_sections_analysis_section(self, sections_analysis):
        """Создание раздела с анализом секций"""
        if not sections_analysis:
            return
        
        self._add_heading("Анализ секций документа", 2)
        
        sections_data = self._parse_json_string(sections_analysis)
        
        sections = sections_data.get('sections', [])
        if sections:
            sections_table_data = []
            for section in sections:
                # Переводим тип секции на русский язык для лучшего понимания
                section_type_ru = self._translate_section_type(section.get('section_type', ''))
                sections_table_data.append([
                    section_type_ru,
                    section.get('section_name', ''),
                    f"{section.get('start_page', '')}-{section.get('end_page', '')}",
                    section.get('check_priority', '')
                ])
            
            self._create_table(sections_table_data, ['Тип секции', 'Название', 'Страницы', 'Приоритет проверки'])
            self.document.add_paragraph()
    
    def _create_detailed_sections_report_section(self, sections_analysis):
        """Создание раздела с детальным отчетом проверки по секциям"""
        if not sections_analysis:
            return
        
        self._add_heading("Детальный отчет проверки по секциям", 2)
        
        # Парсим данные анализа секций
        sections_data = self._parse_json_string(sections_analysis)
        
        # Проверяем наличие детального анализа
        detailed_analysis = sections_data.get('detailed_sections_analysis')
        if not detailed_analysis:
            self.document.add_paragraph("Детальный анализ секций недоступен.")
            return
        
        sections_analysis_data = detailed_analysis.get('sections_analysis', [])
        
        for section_data in sections_analysis_data:
            section_name = section_data.get('section_name', 'Неизвестная секция')
            start_page = section_data.get('start_page', 0)
            end_page = section_data.get('end_page', 0)
            analysis = section_data.get('analysis', {})
            
            # Заголовок секции
            self._add_heading(f"{section_name} (страницы {start_page}-{end_page})", 3)
            
            # Статус соответствия
            compliance_status = analysis.get('compliance_status', 'unknown')
            status_text = self._translate_compliance_status(compliance_status)
            
            status_paragraph = self.document.add_paragraph()
            status_paragraph.add_run("Статус соответствия: ").bold = True
            status_paragraph.add_run(status_text)
            
            # Findings (находки)
            findings = analysis.get('findings', [])
            if findings:
                self.document.add_paragraph()
                findings_heading = self.document.add_paragraph()
                findings_heading.add_run("Найденные проблемы:").bold = True
                
                for finding in findings:
                    finding_para = self.document.add_paragraph()
                    finding_para.style = 'List Bullet'
                    finding_para.add_run(finding.get('title', 'Неизвестная проблема'))
                    
                    if finding.get('description'):
                        desc_para = self.document.add_paragraph()
                        desc_para.style = 'List Bullet 2'
                        desc_para.add_run(finding['description'])
                    
                    if finding.get('recommendation'):
                        rec_para = self.document.add_paragraph()
                        rec_para.style = 'List Bullet 2'
                        rec_para.add_run(f"Рекомендация: {finding['recommendation']}")
            else:
                self.document.add_paragraph()
                no_issues_para = self.document.add_paragraph()
                no_issues_para.add_run("Проблем не обнаружено.").italic = True
            
            # Рекомендации
            recommendations = analysis.get('recommendations', [])
            if recommendations:
                self.document.add_paragraph()
                rec_heading = self.document.add_paragraph()
                rec_heading.add_run("Рекомендации:").bold = True
                
                for recommendation in recommendations:
                    rec_para = self.document.add_paragraph()
                    rec_para.style = 'List Bullet'
                    rec_para.add_run(recommendation)
            
            self.document.add_paragraph()  # Пустая строка между секциями
    
    def _translate_compliance_status(self, status: str) -> str:
        """Перевод статуса соответствия на русский язык"""
        translations = {
            "compliant": "Соответствует требованиям",
            "partially_compliant": "Частично соответствует требованиям",
            "non_compliant": "Не соответствует требованиям",
            "unknown": "Статус неизвестен"
        }
        return translations.get(status, status)
    
    def _create_overall_status_section(self, overall_status, execution_time):
        """Создание раздела с общим статусом"""
        self._add_heading("Общий статус проверки", 2)
        
        # Определяем текст статуса
        status_texts = {
            'pass': 'ПРОЙДЕН',
            'warning': 'ТРЕБУЕТ ВНИМАНИЯ',
            'fail': 'НЕ ПРОЙДЕН'
        }
        
        status_text = status_texts.get(overall_status, 'НЕИЗВЕСТНО')
        
        status_data = [
            ['Общий статус', status_text],
            ['Время выполнения', f"{execution_time:.3f} секунд"]
        ]
        
        table = self._create_table(status_data, ['Параметр', 'Значение'])
        
        # Выделяем статус цветом
        if overall_status == 'pass':
            status_color = RGBColor(0, 128, 0)  # Зеленый
        elif overall_status == 'warning':
            status_color = RGBColor(255, 165, 0)  # Оранжевый
        elif overall_status == 'fail':
            status_color = RGBColor(255, 0, 0)  # Красный
        else:
            status_color = RGBColor(128, 128, 128)  # Серый
        
        # Применяем цвет к ячейке со статусом
        status_cell = table.cell(1, 1)
        for paragraph in status_cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = status_color
                run.bold = True
        
        self.document.add_paragraph()
    
    def generate_hierarchical_report_docx(self, report_data):
        """Генерация DOCX отчета для иерархической проверки"""
        try:
            # Создаем новый документ
            self.document = Document()
            
            # Получаем данные отчета
            hierarchical_result = report_data.get('hierarchical_result', {})
            document_info = report_data.get('document_info', {})
            
            # Создаем заголовок
            self._create_header(document_info)
            
            # Информация о проекте
            if hierarchical_result.get('project_info'):
                self._create_project_info_section(hierarchical_result['project_info'])
            
            # Сводка по соответствию
            if hierarchical_result.get('norm_compliance_summary'):
                self._create_compliance_summary_section(hierarchical_result['norm_compliance_summary'])
            
            # Анализ секций
            if hierarchical_result.get('sections_analysis'):
                self._create_sections_analysis_section(hierarchical_result['sections_analysis'])
            
            # Детальный отчет проверки по секциям
            if hierarchical_result.get('sections_analysis'):
                self._create_detailed_sections_report_section(hierarchical_result['sections_analysis'])
            
            # Общий статус
            overall_status = hierarchical_result.get('overall_status', 'unknown')
            execution_time = hierarchical_result.get('execution_time', 0)
            self._create_overall_status_section(overall_status, execution_time)
            
            # Сохраняем документ в буфер
            from io import BytesIO
            buffer = BytesIO()
            self.document.save(buffer)
            buffer.seek(0)
            docx_content = buffer.getvalue()
            buffer.close()
            
            logger.info(f"DOCX report generated successfully, size: {len(docx_content)} bytes")
            return docx_content
            
        except Exception as e:
            logger.error(f"Error generating DOCX report: {e}")
            raise
    
    def _calculate_a4_equivalent(self, total_pages: int, page_sizes: list) -> int:
        """Расчет эквивалента количества листов формата А4"""
        try:
            if not page_sizes or total_pages == 0:
                # Если нет информации о размерах страниц, возвращаем количество страниц
                return total_pages
            
            # Стандартные размеры форматов (в пунктах)
            A4_WIDTH = 595.28  # A4 width in points
            A4_HEIGHT = 841.89  # A4 height in points
            A4_AREA = A4_WIDTH * A4_HEIGHT
            
            total_a4_equivalent = 0
            
            for page_size in page_sizes:
                if isinstance(page_size, dict):
                    width = page_size.get('width', A4_WIDTH)
                    height = page_size.get('height', A4_HEIGHT)
                    
                    # Рассчитываем площадь страницы
                    page_area = width * height
                    
                    # Рассчитываем коэффициент для пересчета в А4
                    a4_ratio = page_area / A4_AREA
                    
                    # Добавляем к общему эквиваленту
                    total_a4_equivalent += a4_ratio
                else:
                    # Если формат данных неизвестен, считаем как А4
                    total_a4_equivalent += 1
            
            # Округляем до целого числа
            return max(1, round(total_a4_equivalent))
            
        except Exception as e:
            logger.warning(f"Error calculating A4 equivalent: {e}")
            # В случае ошибки возвращаем количество страниц
            return total_pages
    
    def _translate_section_type(self, section_type: str) -> str:
        """Перевод типа секции на русский язык"""
        translations = {
            "title": "Титул",
            "general_data": "Общие данные",
            "main_content": "Основное содержание",
            "specification": "Спецификация",
            "details": "Узлы и детали",
            "unknown": "Неизвестный раздел"
        }
        return translations.get(section_type, section_type)
    
    def generate_report_docx(self, report_data):
        """Генерация DOCX отчета (основной метод)"""
        # Проверяем тип отчета
        if report_data.get('hierarchical_result'):
            return self.generate_hierarchical_report_docx(report_data)
        else:
            # Для обычных отчетов пока возвращаем заглушку
            logger.warning("Regular norm control reports not implemented yet")
            raise NotImplementedError("Regular norm control reports not implemented yet")
