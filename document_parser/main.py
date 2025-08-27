import os
import hashlib
import json
import logging
import io
import psutil
import gc
import time
import signal
import sys
from typing import List, Dict, Any, Optional
from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks, Form
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import psycopg2
from psycopg2.extras import RealDictCursor
import qdrant_client
from qdrant_client.models import Distance, VectorParams
import magic
import PyPDF2
from docx import Document
import re
import httpx
import asyncio
import tiktoken

# Импорт конфигурации проекта
import sys
sys.path.append('/app')
try:
    from config import MAX_CHECKABLE_DOCUMENT_SIZE, MAX_NORMATIVE_DOCUMENT_SIZE, LLM_REQUEST_TIMEOUT
except ImportError:
    # Fallback значения из переменных окружения
    MAX_CHECKABLE_DOCUMENT_SIZE = int(os.getenv("MAX_CHECKABLE_DOCUMENT_SIZE", "104857600"))  # 100 МБ
    MAX_NORMATIVE_DOCUMENT_SIZE = int(os.getenv("MAX_NORMATIVE_DOCUMENT_SIZE", "209715200"))  # 200 МБ
    LLM_REQUEST_TIMEOUT = int(os.getenv("LLM_REQUEST_TIMEOUT", "120"))

# Глобальные переменные для graceful shutdown
shutdown_event = asyncio.Event()
is_shutting_down = False
startup_time = None  # Будет установлено после импорта datetime

def signal_handler(signum, frame):
    """Обработчик сигналов для graceful shutdown"""
    global is_shutting_down
    signal_name = {
        signal.SIGTERM: "SIGTERM",
        signal.SIGINT: "SIGINT",
        signal.SIGHUP: "SIGHUP",
        signal.SIGUSR1: "SIGUSR1",
        signal.SIGUSR2: "SIGUSR2"
    }.get(signum, f"Signal {signum}")
    
    logger.info(f"🔍 [SHUTDOWN] Received {signal_name} at {datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')[:-3]}, initiating graceful shutdown...")
    logger.info(f"🔍 [SHUTDOWN] Process ID: {os.getpid()}, Parent PID: {os.getppid()}")
    
    # Логируем информацию о памяти перед shutdown
    try:
        memory_info = get_memory_usage()
        if "error" not in memory_info:
            logger.info(f"🔍 [SHUTDOWN] Memory usage before shutdown: RSS: {memory_info['rss_mb']:.1f}MB, VMS: {memory_info['vms_mb']:.1f}MB, Percent: {memory_info['percent']:.1f}%")
    except Exception as e:
        logger.warning(f"🔍 [SHUTDOWN] Could not get memory info: {e}")
    
    is_shutting_down = True
    shutdown_event.set()

# Регистрируем обработчики сигналов
signal.signal(signal.SIGTERM, signal_handler)
signal.signal(signal.SIGINT, signal_handler)

def get_memory_usage() -> Dict[str, float]:
    """Получение информации об использовании памяти"""
    try:
        process = psutil.Process()
        memory_info = process.memory_info()
        memory_percent = process.memory_percent()
        
        return {
            "rss_mb": memory_info.rss / 1024 / 1024,  # RSS в МБ
            "vms_mb": memory_info.vms / 1024 / 1024,  # VMS в МБ
            "percent": memory_percent,  # Процент использования
            "available_mb": psutil.virtual_memory().available / 1024 / 1024  # Доступная память в МБ
        }
    except Exception as e:
        logger.error(f"Error getting memory usage: {e}")
        return {"error": str(e)}

def get_available_memory() -> float:
    """Получение доступной памяти в МБ"""
    try:
        return psutil.virtual_memory().available / 1024 / 1024
    except Exception as e:
        logger.error(f"Error getting available memory: {e}")
        return 0.0

def log_memory_usage(context: str = ""):
    """Логирование использования памяти"""
    try:
        memory_info = get_memory_usage()
        if "error" not in memory_info:
            logger.info(f"🔍 [DEBUG] DocumentParser: Memory usage {context}: "
                       f"RSS: {memory_info['rss_mb']:.1f}MB, "
                       f"VMS: {memory_info['vms_mb']:.1f}MB, "
                       f"Percent: {memory_info['percent']:.1f}%, "
                       f"Available: {memory_info['available_mb']:.1f}MB")
        else:
            logger.warning(f"🔍 [DEBUG] DocumentParser: Could not get memory usage {context}: {memory_info['error']}")
    except Exception as e:
        logger.error(f"Error in log_memory_usage: {e}")

def cleanup_memory():
    """Очистка памяти"""
    try:
        gc.collect()
        logger.info(f"🔍 [DEBUG] DocumentParser: Memory cleanup completed")
    except Exception as e:
        logger.error(f"Error in cleanup_memory: {e}")

def check_memory_pressure() -> bool:
    """Проверка давления на память"""
    try:
        memory_info = get_memory_usage()
        if "error" in memory_info:
            return False
        
        # Проверяем, если используется больше 80% памяти или доступно меньше 500MB
        if memory_info['percent'] > 80 or memory_info['available_mb'] < 500:
            logger.warning(f"🔍 [MEMORY] High memory pressure detected: "
                          f"Usage: {memory_info['percent']:.1f}%, "
                          f"Available: {memory_info['available_mb']:.1f}MB")
            return True
        return False
    except Exception as e:
        logger.error(f"Error checking memory pressure: {e}")
        return False

# OCR imports
import pytesseract
from PIL import Image
import cv2
import numpy as np
from pdf2image import convert_from_bytes
import tempfile
import math

from datetime import datetime, timedelta

# Инициализируем startup_time после импорта datetime
startup_time = datetime.now()

# PDF generation imports
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# структура данных о результате проверки документа и проверяемом документе
class DocumentInspectionResult:
    """Структура данных для результата проверки документа"""
    
    def __init__(self):
        # Основная информация о документе
        self.document_id = None # ID документа
        self.document_name = None # Наименование документа. Проверить на соответствие "A9.5.MTH.04"
        self.document_type = None # pdf, docx, dwg, txt
        self.document_engineering_stage = None # ПД, РД, ТЭО
        self.document_mark = None # Марка комплекта документации: КЖ, КМ, АС, ТХ, КС, КП, КР, КРС, КРС-1, КРС-2, КРС-3, КРС-4, КРС-5, КРС-6, КРС-7, КРС-8, КРС-9, КРС-10
        self.document_number = None # Номер комплекта документации
        self.document_date = None # Дата комплекта документации
        self.document_author = None # РАЗРАБОТАЛ комплект документации
        self.document_reviewer = None # ПРОВЕРИЛ комплект документации
        self.document_approver = None # ГИП комплекта документации
        self.document_status = None # Статус комплекта документации: IFR, IFD, IFC, IFS, IFT, IFT-1, IFT-2, IFT-3, IFT-4, IFT-5, IFT-6, IFT-7, IFT-8, IFT-9, IFT-10
        self.document_size = None # Размер файла в байтах
        
        # Информация о страницах
        self.document_pages = None # Количество страниц в документе
        self.document_pages_vector = None # Количество векторных страниц в документе
        self.document_pages_scanned = None # Количество сканированных страниц в документе
        self.document_pages_unknown = None # Количество неизвестных страниц в документе
        self.document_pages_total = None # Общее количество страниц в документе
        self.document_pages_total_a4_sheets_equivalent = None # Общее количество листов А4 эквивалентных в документе
        
        # Статистика по страницам
        self.document_pages_with_violations = None # Количество страниц с отклонениями
        self.document_pages_clean = None # Количество страниц без отклонений
        
        # Общие показатели отклонений
        self.document_total_violations = None # Общее количество выявленных отклонений
        self.document_critical_violations = None # Критические отклонения
        self.document_major_violations = None # Значительные отклонения
        self.document_minor_violations = None # Мелкие отклонения
        self.document_info_violations = None # Информационные замечания
        
        # Процентные показатели
        self.document_compliance_percentage = None # Процент соответствия (0-100)
        self.document_violations_per_page_avg = None # Среднее количество отклонений на страницу
        
        # Устаревшие поля (для обратной совместимости)
        self.document_total_errors_count = None # Общее количество ошибок в документе
        self.document_total_warnings_count = None # Общее количество предупреждений в документе
        self.document_total_info_count = None # Общее количество информационных сообщений в документе
        self.document_total_suggestions_count = None # Общее количество предложений в документе
        self.document_total_corrections_count = None # Общее количество исправлений в документе
        
        # Результаты проверки листов документа
        self.document_pages_results = [] # Результаты проверки листов документа

class DocumentPageInspectionResult:
    """Структура данных для результата проверки листа документа"""
    
    def __init__(self):
        # Основная информация о странице
        self.page_number = None # Номер страницы
        self.page_type = None # vector, scanned, unknown
        self.page_format = None # A0, A1, A2, A3, A4, A5, A6, A7, A8, A9, A10, B0, B1, B2, B3, B4, B5, B6, B7, B8, B9, B10
        self.page_width_mm = None # Ширина страницы в мм
        self.page_height_mm = None # Высота страницы в мм
        self.page_orientation = None # portrait, landscape
        self.page_a4_sheets_equivalent = None # Эквивалент листов А4
        
        # Информация о тексте страницы
        self.page_text = None # Текст листа документа
        self.page_text_confidence = None # Уверенность в тексте листа документа
        self.page_text_method = None # Метод извлечения текста листа документа
        self.page_text_length = None # Длина текста в символах
        
        # Результаты OCR (для сканированных страниц)
        self.page_ocr_confidence = None # Уверенность OCR
        self.page_ocr_method = None # Метод OCR
        self.page_ocr_all_results = [] # Все результаты OCR
        
        # Показатели отклонений на странице
        self.page_violations_count = None # Общее количество отклонений на странице
        self.page_critical_violations = None # Критические отклонения на странице
        self.page_major_violations = None # Значительные отклонения на странице
        self.page_minor_violations = None # Мелкие отклонения на странице
        self.page_info_violations = None # Информационные замечания на странице
        self.page_compliance_percentage = None # Процент соответствия страницы
        
        # Детали отклонений на странице
        self.page_violations_details = [] # Детали отклонений на странице


class DocumentViolationDetail:
    """Структура данных для деталей отклонения"""
    
    def __init__(self):
        self.violation_type = None # critical, major, minor, info
        self.violation_category = None # general_requirements, text_part, graphical_part, specifications, assembly_drawings, detail_drawings, schemes
        self.violation_description = None # Описание отклонения
        self.violation_clause = None # Пункт нормы (например: ГОСТ Р 21.1101-2013 п.5.2)
        self.violation_severity = None # Серьезность (1-5)
        self.violation_recommendation = None # Рекомендация по исправлению
        self.violation_location = None # Местоположение на странице
        self.violation_confidence = None # Уверенность в отклонении (0-1)




# Настройка логирования с детальным трейслогом
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class TraceLogger:
    """Класс для детального трейслога всех операций"""
    
    @staticmethod
    def log_step(step: str, details: str = "", document_id: int = None):
        """Логирование шага процесса"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
        doc_info = f"[DOC:{document_id}]" if document_id else ""
        logger.info(f"🔍 [TRACE] {timestamp} {doc_info} STEP: {step} - {details}")
    
    @staticmethod
    def log_upload_start(filename: str, file_size: int):
        """Логирование начала загрузки документа"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
        size_mb = file_size / (1024 * 1024)
        logger.info(f"🚀 [UPLOAD] {timestamp} START: Загрузка документа '{filename}' ({size_mb:.2f}MB)")
    
    @staticmethod
    def log_upload_success(filename: str, document_id: int, processing_time: float):
        """Логирование успешной загрузки документа"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
        logger.info(f"✅ [UPLOAD] {timestamp} [DOC:{document_id}] SUCCESS: Документ '{filename}' загружен за {processing_time:.2f}с")
    
    @staticmethod
    def log_indexing_start(document_id: int, pages_count: int):
        """Логирование начала индексации"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
        logger.info(f"📚 [INDEX] {timestamp} [DOC:{document_id}] START: Индексация {pages_count} страниц")
    
    @staticmethod
    def log_indexing_success(document_id: int, processing_time: float, tokens_count: int):
        """Логирование успешной индексации"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
        logger.info(f"✅ [INDEX] {timestamp} [DOC:{document_id}] SUCCESS: Индексация завершена за {processing_time:.2f}с, {tokens_count} токенов")
    
    @staticmethod
    def log_normcontrol_start(document_id: int, pages_count: int):
        """Логирование начала нормоконтроля"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
        logger.info(f"🔍 [NORMCONTROL] {timestamp} [DOC:{document_id}] START: Проверка нормоконтроля {pages_count} страниц")
    
    @staticmethod
    def log_llm_request(document_id: int, page_number: int, prompt_length: int):
        """Логирование запроса к LLM"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
        logger.info(f"🤖 [LLM] {timestamp} [DOC:{document_id}] REQUEST: Страница {page_number}, промпт {prompt_length} символов")
    
    @staticmethod
    def log_llm_response(document_id: int, page_number: int, response_time: float, response_length: int):
        """Логирование ответа от LLM"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
        logger.info(f"🤖 [LLM] {timestamp} [DOC:{document_id}] RESPONSE: Страница {page_number}, время {response_time:.2f}с, ответ {response_length} символов")
    
    @staticmethod
    def log_report_generation(document_id: int, findings_count: int):
        """Логирование генерации отчета"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
        logger.info(f"📄 [REPORT] {timestamp} [DOC:{document_id}] GENERATION: Отчет с {findings_count} находками")
    
    @staticmethod
    def log_error(step: str, error: str, document_id: int = None):
        """Логирование ошибок"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
        doc_info = f"[DOC:{document_id}]" if document_id else ""
        logger.error(f"❌ [ERROR] {timestamp} {doc_info} {step}: {error}")

# Создаем глобальный экземпляр трейслоггера
trace_logger = TraceLogger()

app = FastAPI(title="Document Parser Service", version="1.0.0")

# Настройки для увеличения лимита размера файлов
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Увеличиваем лимит размера файлов до 100MB
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request
from starlette.responses import Response

class LargeFileMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        # Увеличиваем лимит для загрузки файлов из конфигурации
        if request.url.path == "/upload":
            request.scope["max_content_size"] = MAX_NORMATIVE_DOCUMENT_SIZE
        elif request.url.path == "/upload/checkable":
            request.scope["max_content_size"] = MAX_CHECKABLE_DOCUMENT_SIZE
        return await call_next(request)

class ErrorHandlingMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        start_time = time.time()
        request_id = f"req_{int(start_time * 1000)}"
        
        try:
            # Логируем начало запроса
            logger.info(f"🔍 [REQUEST] {request_id}: {request.method} {request.url.path} started")
            
            # Проверяем состояние shutdown
            if is_shutting_down and request.url.path not in ["/health", "/metrics"]:
                logger.warning(f"🔍 [REQUEST] {request_id}: Service is shutting down, rejecting request")
                return JSONResponse(
                    status_code=503,
                    content={"error": "Service is shutting down", "request_id": request_id}
                )
            
            # Проверяем давление на память
            if check_memory_pressure():
                logger.warning(f"🔍 [REQUEST] {request_id}: High memory pressure detected")
                cleanup_memory()
            
            response = await call_next(request)
            
            # Логируем успешное завершение
            duration = time.time() - start_time
            logger.info(f"🔍 [REQUEST] {request_id}: {request.method} {request.url.path} completed in {duration:.3f}s (status: {response.status_code})")
            
            return response
            
        except Exception as e:
            # Логируем ошибку
            duration = time.time() - start_time
            logger.error(f"🔍 [REQUEST] {request_id}: {request.method} {request.url.path} failed after {duration:.3f}s: {e}")
            
            # Возвращаем ошибку 500
            return JSONResponse(
                status_code=500,
                content={
                    "error": "Internal server error",
                    "request_id": request_id,
                    "timestamp": datetime.now().isoformat()
                }
            )

app.add_middleware(LargeFileMiddleware)
app.add_middleware(ErrorHandlingMiddleware)

# CORS middleware уже добавлен выше

# Конфигурация
POSTGRES_HOST = os.getenv("POSTGRES_HOST", "norms-db")
POSTGRES_DB = os.getenv("POSTGRES_DB", "norms_db")
POSTGRES_USER = os.getenv("POSTGRES_USER", "norms_user")
POSTGRES_PASSWORD = os.getenv("POSTGRES_PASSWORD", "norms_password")
QDRANT_HOST = os.getenv("QDRANT_HOST", "qdrant")
QDRANT_PORT = int(os.getenv("QDRANT_PORT", "6333"))
RAG_SERVICE_URL = os.getenv("RAG_SERVICE_URL", "http://rag-service:8003")

class TransactionContext:
    """Контекстный менеджер для управления транзакциями PostgreSQL"""
    
    def __init__(self, connection):
        self.connection = connection
        self.cursor = None
        self.transaction_id = f"tx_{int(time.time() * 1000)}"
    
    def __enter__(self):
        """Начало транзакции"""
        try:
            self.cursor = self.connection.cursor()
            logger.debug(f"🔍 [TRANSACTION] {self.transaction_id}: Transaction started")
            return self.connection
        except Exception as e:
            logger.error(f"🔍 [TRANSACTION] {self.transaction_id}: Error starting transaction: {e}")
            raise
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        """Завершение транзакции с автоматическим commit/rollback"""
        try:
            if exc_type is None:
                # Нет исключений - коммитим транзакцию
                self.connection.commit()
                logger.debug(f"🔍 [TRANSACTION] {self.transaction_id}: Transaction committed successfully")
            else:
                # Есть исключения - откатываем транзакцию
                self.connection.rollback()
                logger.error(f"🔍 [TRANSACTION] {self.transaction_id}: Transaction rolled back due to error: {exc_type.__name__}: {exc_val}")
        except Exception as e:
            logger.error(f"🔍 [TRANSACTION] {self.transaction_id}: Error during transaction cleanup: {e}")
            # Пытаемся откатить транзакцию при ошибке очистки
            try:
                if not self.connection.closed:
                    self.connection.rollback()
                    logger.debug(f"🔍 [TRANSACTION] {self.transaction_id}: Emergency rollback completed")
            except Exception as rollback_error:
                logger.error(f"🔍 [TRANSACTION] {self.transaction_id}: Emergency rollback failed: {rollback_error}")
        finally:
            # Закрываем курсор
            if self.cursor:
                try:
                    self.cursor.close()
                    logger.debug(f"🔍 [TRANSACTION] {self.transaction_id}: Cursor closed")
                except Exception as cursor_error:
                    logger.error(f"🔍 [TRANSACTION] {self.transaction_id}: Error closing cursor: {cursor_error}")


class ReadOnlyTransactionContext:
    """Контекстный менеджер для операций только для чтения PostgreSQL"""
    
    def __init__(self, connection):
        self.connection = connection
        self.cursor = None
        self.transaction_id = f"read_tx_{int(time.time() * 1000)}"
    
    def __enter__(self):
        """Начало транзакции только для чтения"""
        try:
            self.cursor = self.connection.cursor()
            # Устанавливаем транзакцию только для чтения
            self.cursor.execute("SET TRANSACTION READ ONLY")
            logger.debug(f"🔍 [READ_TRANSACTION] {self.transaction_id}: Read-only transaction started")
            return self.connection
        except Exception as e:
            logger.error(f"🔍 [READ_TRANSACTION] {self.transaction_id}: Error starting read-only transaction: {e}")
            raise
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        """Завершение транзакции только для чтения"""
        try:
            if exc_type is None:
                # Нет исключений - коммитим транзакцию (для чтения это безопасно)
                self.connection.commit()
                logger.debug(f"🔍 [READ_TRANSACTION] {self.transaction_id}: Read-only transaction committed successfully")
            else:
                # Есть исключения - откатываем транзакцию
                self.connection.rollback()
                logger.error(f"🔍 [READ_TRANSACTION] {self.transaction_id}: Read-only transaction rolled back due to error: {exc_type.__name__}: {exc_val}")
        except Exception as e:
            logger.error(f"🔍 [READ_TRANSACTION] {self.transaction_id}: Error during read-only transaction cleanup: {e}")
            # Пытаемся откатить транзакцию при ошибке очистки
            try:
                if not self.connection.closed:
                    self.connection.rollback()
                    logger.debug(f"🔍 [READ_TRANSACTION] {self.transaction_id}: Emergency rollback completed")
            except Exception as rollback_error:
                logger.error(f"🔍 [READ_TRANSACTION] {self.transaction_id}: Emergency rollback failed: {rollback_error}")
        finally:
            # Закрываем курсор
            if self.cursor:
                try:
                    self.cursor.close()
                    logger.debug(f"🔍 [READ_TRANSACTION] {self.transaction_id}: Cursor closed")
                except Exception as cursor_error:
                    logger.error(f"🔍 [READ_TRANSACTION] {self.transaction_id}: Error closing cursor: {cursor_error}")


class DocumentParser:
    def __init__(self):
        self.db_conn = None
        self.qdrant_client = None
        self.connection_retry_count = 0
        self.max_retries = 3
        self.retry_delay = 5  # секунды
        self.connect_databases()
    
    def connect_databases(self):
        """Подключение к базам данных с повторными попытками"""
        connection_start_time = datetime.now()
        logger.info(f"🔍 [CONNECTION] Starting database connections at {connection_start_time.strftime('%Y-%m-%d %H:%M:%S.%f')[:-3]}")
        
        for attempt in range(self.max_retries):
            try:
                logger.info(f"🔍 [CONNECTION] Attempt {attempt + 1}/{self.max_retries} to connect to databases")
                
                # PostgreSQL
                postgres_start_time = datetime.now()
                self.db_conn = psycopg2.connect(
                    host=POSTGRES_HOST,
                    database=POSTGRES_DB,
                    user=POSTGRES_USER,
                    password=POSTGRES_PASSWORD,
                    connect_timeout=10,
                    application_name="document_parser"
                )
                postgres_end_time = datetime.now()
                postgres_duration = (postgres_end_time - postgres_start_time).total_seconds()
                logger.info(f"🔍 [CONNECTION] Connected to PostgreSQL in {postgres_duration:.3f}s")
                
                # Qdrant
                qdrant_start_time = datetime.now()
                self.qdrant_client = qdrant_client.QdrantClient(
                    host=QDRANT_HOST,
                    port=QDRANT_PORT,
                    timeout=10
                )
                qdrant_end_time = datetime.now()
                qdrant_duration = (qdrant_end_time - qdrant_start_time).total_seconds()
                logger.info(f"🔍 [CONNECTION] Connected to Qdrant in {qdrant_duration:.3f}s")
                
                # Сброс счетчика попыток при успешном подключении
                self.connection_retry_count = 0
                connection_end_time = datetime.now()
                total_duration = (connection_end_time - connection_start_time).total_seconds()
                logger.info(f"🔍 [CONNECTION] All database connections established successfully in {total_duration:.3f}s")
                return
                
            except Exception as e:
                self.connection_retry_count += 1
                logger.error(f"🔍 [CONNECTION] Database connection error (attempt {attempt + 1}/{self.max_retries}): {e}")
                logger.error(f"🔍 [CONNECTION] Error type: {type(e).__name__}")
                
                if attempt < self.max_retries - 1:
                    logger.info(f"🔍 [CONNECTION] Retrying in {self.retry_delay} seconds...")
                    time.sleep(self.retry_delay)
                else:
                    connection_end_time = datetime.now()
                    total_duration = (connection_end_time - connection_start_time).total_seconds()
                    logger.error(f"🔍 [CONNECTION] Failed to connect to databases after {self.max_retries} attempts in {total_duration:.3f}s")
                    raise
    
    def get_db_connection(self):
        """Безопасное получение соединения с базой данных"""
        try:
            # Проверяем, что соединение существует и активно
            if self.db_conn is None or self.db_conn.closed:
                logger.info("🔍 [CONNECTION] Reconnecting to PostgreSQL...")
                self.db_conn = psycopg2.connect(
                    host=POSTGRES_HOST,
                    database=POSTGRES_DB,
                    user=POSTGRES_USER,
                    password=POSTGRES_PASSWORD,
                    connect_timeout=10,
                    application_name="document_parser"
                )
                logger.info("🔍 [CONNECTION] Reconnected to PostgreSQL")
            
            # Проверяем, что соединение работает
            try:
                with self.db_conn.cursor() as cursor:
                    cursor.execute("SELECT 1")
                    cursor.fetchone()
            except Exception as test_error:
                logger.warning(f"🔍 [CONNECTION] Database connection test failed: {test_error}")
                # Переподключаемся при ошибке теста
                if self.db_conn and not self.db_conn.closed:
                    try:
                        self.db_conn.close()
                    except:
                        pass
                
                self.db_conn = psycopg2.connect(
                    host=POSTGRES_HOST,
                    database=POSTGRES_DB,
                    user=POSTGRES_USER,
                    password=POSTGRES_PASSWORD,
                    connect_timeout=10,
                    application_name="document_parser"
                )
                logger.info("🔍 [CONNECTION] Reconnected to PostgreSQL after test failure")
            
            return self.db_conn
            
        except Exception as e:
            logger.error(f"🔍 [CONNECTION] Database connection check failed: {e}")
            # Пытаемся переподключиться
            try:
                if self.db_conn and not self.db_conn.closed:
                    self.db_conn.close()
            except Exception as close_error:
                logger.error(f"🔍 [CONNECTION] Error closing connection: {close_error}")
            
            try:
                self.db_conn = psycopg2.connect(
                    host=POSTGRES_HOST,
                    database=POSTGRES_DB,
                    user=POSTGRES_USER,
                    password=POSTGRES_PASSWORD,
                    connect_timeout=10,
                    application_name="document_parser"
                )
                logger.info("🔍 [CONNECTION] Reconnected to PostgreSQL after error")
                return self.db_conn
            except Exception as reconnect_error:
                logger.error(f"🔍 [CONNECTION] Failed to reconnect: {reconnect_error}")
                raise
    
    def transaction_context(self):
        """Контекстный менеджер для управления транзакциями"""
        return TransactionContext(self.get_db_connection())
    
    def read_only_transaction_context(self):
        """Контекстный менеджер для операций только для чтения"""
        return ReadOnlyTransactionContext(self.get_db_connection())
    
    def execute_in_transaction(self, operation, *args, **kwargs):
        """Выполнение операции в транзакции с автоматическим управлением"""
        with self.transaction_context() as conn:
            return operation(conn, *args, **kwargs)
    
    def execute_in_read_only_transaction(self, operation, *args, **kwargs):
        """Выполнение операции в транзакции только для чтения"""
        with self.read_only_transaction_context() as conn:
            return operation(conn, *args, **kwargs)
    def calculate_file_hash(self, file_content: bytes) -> str:
        """Вычисление SHA-256 хеша файла"""
        return hashlib.sha256(file_content).hexdigest()

    def extract_text_from_image(self, image: Image.Image, page_number: int) -> Dict[str, Any]:
        """Извлечение текста из изображения с помощью OCR"""
        try:
            # Конвертируем PIL Image в OpenCV формат
            opencv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            
            # Предобработка изображения для улучшения OCR
            # Конвертируем в оттенки серого
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            
            # Применяем различные методы улучшения
            # 1. Адаптивная пороговая обработка
            thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            # 2. Морфологические операции для удаления шума
            kernel = np.ones((1, 1), np.uint8)
            thresh = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
            
            # 3. Увеличение контрастности
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            enhanced = clahe.apply(gray)
            
            # Настройки OCR для русского языка
            custom_config = r'--oem 3 --psm 6 -l rus+eng'
            
            # Пробуем разные варианты обработки
            results = []
            
            # 1. Оригинальное изображение
            try:
                text1 = pytesseract.image_to_string(image, config=custom_config)
                if text1.strip():
                    results.append(("original", text1.strip(), 0.7))
            except Exception as e:
                logger.warning(f"OCR failed on original image for page {page_number}: {e}")
            
            # 2. Обработанное изображение (адаптивная пороговая обработка)
            try:
                text2 = pytesseract.image_to_string(thresh, config=custom_config)
                if text2.strip():
                    results.append(("threshold", text2.strip(), 0.8))
            except Exception as e:
                logger.warning(f"OCR failed on threshold image for page {page_number}: {e}")
            
            # 3. Улучшенное изображение (CLAHE)
            try:
                text3 = pytesseract.image_to_string(enhanced, config=custom_config)
                if text3.strip():
                    results.append(("enhanced", text3.strip(), 0.9))
            except Exception as e:
                logger.warning(f"OCR failed on enhanced image for page {page_number}: {e}")
            
            # Выбираем лучший результат
            if results:
                # Сортируем по длине текста и уверенности
                results.sort(key=lambda x: (len(x[1]), x[2]), reverse=True)
                best_method, best_text, confidence = results[0]
                
                logger.info(f"OCR completed for page {page_number} using {best_method} method, confidence: {confidence}")
                
                return {
                    "text": best_text,
                    "confidence": confidence,
                    "method": best_method,
                    "all_results": results
                }
            else:
                logger.warning(f"No OCR results for page {page_number}")
                return {
                    "text": "",
                    "confidence": 0.0,
                    "method": "none",
                    "all_results": []
                }
                
        except Exception as e:
            logger.error(f"OCR processing error for page {page_number}: {e}")
            return {
                "text": "",
                "confidence": 0.0,
                "method": "error",
                "all_results": []
            }

    def get_page_format_info(self, page) -> Dict[str, Any]:
        """Получение информации о формате страницы PDF"""
        try:
            # Получаем размеры страницы в точках (1/72 дюйма)
            media_box = page.mediabox
            width_pt = float(media_box.width)
            height_pt = float(media_box.height)
            
            # Конвертируем в миллиметры (1 дюйм = 25.4 мм, 1 точка = 1/72 дюйма)
            width_mm = width_pt * 25.4 / 72
            height_mm = height_pt * 25.4 / 72
            
            # Определяем стандартный формат
            format_name = self.determine_page_format(width_mm, height_mm)
            
            # Вычисляем количество листов А4
            a4_sheets = self.calculate_a4_sheets(width_mm, height_mm)
            
            return {
                "width_pt": width_pt,
                "height_pt": height_pt,
                "width_mm": round(width_mm, 2),
                "height_mm": round(height_mm, 2),
                "format_name": format_name,
                "a4_sheets": a4_sheets,
                "orientation": "portrait" if height_mm > width_mm else "landscape"
            }
            
        except Exception as e:
            logger.error(f"Error getting page format info: {e}")
            return {
                "width_pt": 0,
                "height_pt": 0,
                "width_mm": 0,
                "height_mm": 0,
                "format_name": "unknown",
                "a4_sheets": 0,
                "orientation": "unknown"
            }

    def determine_page_format(self, width_mm: float, height_mm: float) -> str:
        """Определение стандартного формата страницы"""
        # Стандартные форматы в мм (ширина x высота)
        formats = {
            "A0": (841, 1189),
            "A1": (594, 841),
            "A2": (420, 594),
            "A3": (297, 420),
            "A4": (210, 297),
            "A5": (148, 210),
            "A6": (105, 148),
            "A7": (74, 105),
            "A8": (52, 74),
            "A9": (37, 52),
            "A10": (26, 37),
            "B0": (1000, 1414),
            "B1": (707, 1000),
            "B2": (500, 707),
            "B3": (353, 500),
            "B4": (250, 353),
            "B5": (176, 250),
            "B6": (125, 176),
            "B7": (88, 125),
            "B8": (62, 88),
            "B9": (44, 62),
            "B10": (31, 44),
        }
        
        # Допустимое отклонение в мм
        tolerance = 5
        
        for format_name, (std_width, std_height) in formats.items():
            # Проверяем в обеих ориентациях
            if (abs(width_mm - std_width) <= tolerance and abs(height_mm - std_height) <= tolerance) or \
               (abs(width_mm - std_height) <= tolerance and abs(height_mm - std_width) <= tolerance):
                return format_name
        
        # Если не найден стандартный формат, возвращаем кастомный
        return f"custom_{int(width_mm)}x{int(height_mm)}"

    def calculate_a4_sheets(self, width_mm: float, height_mm: float) -> float:
        """Вычисление количества листов А4 (с округлением)"""
        try:
            # Площадь А4 в мм²
            a4_area = 210 * 297  # 62,370 мм²
            
            # Площадь текущей страницы в мм²
            page_area = width_mm * height_mm
            
            # Количество листов А4
            a4_sheets = page_area / a4_area
            
            # Округляем до 2 знаков после запятой
            return round(a4_sheets, 2)
            
        except Exception as e:
            logger.error(f"Error calculating A4 sheets: {e}")
            return 0.0

    def detect_page_type(self, page) -> str:
        """Определение типа страницы (векторная или сканированная)"""
        try:
            # Извлекаем текст
            text = page.extract_text()
            
            # Если текст извлекается легко и его достаточно, считаем страницу векторной
            if text and len(text.strip()) > 50:
                return "vector"
            
            # Проверяем наличие изображений на странице
            if '/XObject' in page['/Resources']:
                xObject = page['/Resources']['/XObject'].get_object()
                for obj in xObject:
                    if xObject[obj]['/Subtype'] == '/Image':
                        return "scanned"
            
            # Если текст короткий или отсутствует, считаем сканированной
            if not text or len(text.strip()) < 10:
                return "scanned"
            
            return "vector"
            
        except Exception as e:
            logger.warning(f"Error detecting page type: {e}")
            return "unknown"

    def count_tokens(self, text: str) -> int:
        """Подсчет количества токенов в тексте"""
        try:
            # Используем cl100k_base кодировку (GPT-4, GPT-3.5-turbo)
            encoding = tiktoken.get_encoding("cl100k_base")
            tokens = encoding.encode(text)
            return len(tokens)
        except Exception as e:
            logger.error(f"Token counting error: {e}")
            # Fallback: приблизительный подсчет по словам
            return len(text.split())

    def calculate_document_tokens(self, elements: List[Dict[str, Any]]) -> int:
        """Подсчет общего количества токенов в документе"""
        total_tokens = 0
        for element in elements:
            content = element.get("element_content", "")
            if content:
                total_tokens += self.count_tokens(content)
        return total_tokens
    
    def detect_file_type(self, file_content: bytes) -> str:
        """Определение типа файла по содержимому"""
        print(f"🔍 [DEBUG] DocumentParser: detect_file_type called with content length: {len(file_content)}")
        print(f"🔍 [DEBUG] DocumentParser: Content preview: {file_content[:100]}...")
        
        mime_type = magic.from_buffer(file_content, mime=True)
        print(f"🔍 [DEBUG] DocumentParser: Magic detected MIME type: {mime_type}")
        
        if mime_type == "application/pdf":
            print(f"🔍 [DEBUG] DocumentParser: Detected as PDF")
            return "pdf"
        else:
            print(f"🔍 [DEBUG] DocumentParser: Unsupported MIME type: {mime_type}")
            raise ValueError(f"Only PDF files are supported. Detected MIME type: {mime_type}")

    def parse_file(self, file_content: bytes, file_type: str) -> List[Dict[str, Any]]:
        """Парсинг файла в зависимости от типа"""
        if file_type == "pdf":
            return self.parse_pdf(file_content)
        elif file_type == "docx":
            return self.parse_docx(file_content)
        elif file_type == "dwg":
            return self.parse_dwg(file_content)
        elif file_type == "txt":
            return self.parse_text(file_content)
        else:
            raise ValueError(f"Unsupported file type for parsing: {file_type}")

    def parse_text(self, file_content: bytes) -> List[Dict[str, Any]]:
        """Парсинг текстового файла"""
        elements = []
        
        try:
            # Декодируем содержимое файла
            text_content = file_content.decode('utf-8')
            
            # Разбиваем на строки
            lines = text_content.split('\n')
            
            for i, line in enumerate(lines, 1):
                line = line.strip()
                if line:  # Пропускаем пустые строки
                    elements.append({
                        "element_type": "text",
                        "element_content": line,
                        "page_number": 1,
                        "confidence_score": 1.0
                    })
            
            return elements
            
        except Exception as e:
            logger.error(f"Text parsing error: {e}")
            raise
    
    def parse_pdf(self, file_content: bytes) -> DocumentInspectionResult:
        """Парсинг PDF документа с поддержкой векторных и сканированных страниц"""
        logger.info(f"🔍 [DEBUG] DocumentParser: Starting PDF parsing for {len(file_content)} bytes")
        log_memory_usage("before PDF parsing")
        result = DocumentInspectionResult()
        document_format_stats = {
            "total_pages": 0,
            "total_a4_sheets": 0.0,
            "formats": {},
            "orientations": {"portrait": 0, "landscape": 0},
            "page_types": {"vector": 0, "scanned": 0, "unknown": 0}
        }
        
        try:
            logger.info(f"🔍 [DEBUG] DocumentParser: Creating PDF reader from {len(file_content)} bytes")
            # Создаем file-like объект из bytes
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            logger.info(f"🔍 [DEBUG] DocumentParser: PDF reader created successfully")
            
            # Получаем информацию о документе
            document_info = self.get_document_info(pdf_reader)
            result.document_name = document_info["name"]
            result.document_type = document_info["type"]
            result.document_engineering_stage = document_info["engineering_stage"]
            result.document_mark = document_info["mark"]
            result.document_number = document_info["number"]
            result.document_date = document_info["date"]
            result.document_author = document_info["author"]
            result.document_reviewer = document_info["reviewer"]
            result.document_approver = document_info["approver"]
            result.document_status = document_info["status"]
            result.document_size = document_info["size"]
            # покажем в лог информацию о документе
            logger.info(f"Document info: {document_info}")
            
            
            document_format_stats["total_pages"] = len(pdf_reader.pages)
            logger.info(f"Processing PDF with {len(pdf_reader.pages)} pages")
            
            # Конвертируем PDF в изображения для OCR
            try:
                images = convert_from_bytes(file_content, dpi=300)
                logger.info(f"Converted PDF to {len(images)} images")
            except Exception as e:
                logger.warning(f"Failed to convert PDF to images: {e}")
                images = []
            
            logger.info(f"🔍 [DEBUG] DocumentParser: Starting to process {len(pdf_reader.pages)} pages")
            
            # Определение параметров страниц перед обработкой
            logger.info(f"🔍 [DEBUG] DocumentParser: PDF document parameters:")
            logger.info(f"🔍 [DEBUG] DocumentParser: - Total pages: {len(pdf_reader.pages)}")
            logger.info(f"🔍 [DEBUG] DocumentParser: - File size: {len(file_content)} bytes")
            logger.info(f"🔍 [DEBUG] DocumentParser: - Average page size: {len(file_content) // len(pdf_reader.pages)} bytes")
            
            # Настройки батчевой обработки
            BATCH_SIZE = 5  # Обрабатываем по 5 страниц за раз
            total_batches = (len(pdf_reader.pages) + BATCH_SIZE - 1) // BATCH_SIZE
            logger.info(f"🔍 [DEBUG] DocumentParser: Batch processing settings:")
            logger.info(f"🔍 [DEBUG] DocumentParser: - Batch size: {BATCH_SIZE} pages")
            logger.info(f"🔍 [DEBUG] DocumentParser: - Total batches: {total_batches}")
            
            # Обработка страниц батчами
            for batch_num in range(total_batches):
                start_page = batch_num * BATCH_SIZE
                end_page = min((batch_num + 1) * BATCH_SIZE, len(pdf_reader.pages))
                
                logger.info(f"🔍 [DEBUG] DocumentParser: Starting batch {batch_num + 1}/{total_batches} (pages {start_page + 1}-{end_page})")
                log_memory_usage(f"before batch {batch_num + 1}")
                
                # Обработка страниц в текущем батче
                for page_num in range(start_page, end_page):
                    page = pdf_reader.pages[page_num]
                    page_number = page_num + 1
                    
                    logger.info(f"🔍 [DEBUG] DocumentParser: Processing page {page_number}/{len(pdf_reader.pages)} in batch {batch_num + 1}")
                    
                    # Создаем объект результата страницы
                    page_result = DocumentPageInspectionResult()
                    page_result.page_number = page_number
                    
                    # Анализ параметров страницы
                    try:
                        page_text = page.extract_text()
                        page_size = len(page_text) if page_text else 0
                        logger.info(f"🔍 [DEBUG] DocumentParser: Page {page_number} parameters:")
                        logger.info(f"🔍 [DEBUG] DocumentParser: - Text size: {page_size} characters")
                        logger.info(f"🔍 [DEBUG] DocumentParser: - Has text: {bool(page_text)}")
                        logger.info(f"🔍 [DEBUG] DocumentParser: - Text preview: {page_text[:100] if page_text else 'No text'}...")
                    except Exception as e:
                        logger.warning(f"🔍 [DEBUG] DocumentParser: Failed to extract text from page {page_number}: {e}")
                        page_size = 0
                    
                    # Получаем информацию о формате страницы
                    format_info = self.get_page_format_info(page)
                    document_format_stats["total_a4_sheets"] += format_info["a4_sheets"]
                    
                    # Заполняем информацию о формате страницы
                    page_result.page_format = format_info["format_name"]
                    page_result.page_width_mm = format_info["width_mm"]
                    page_result.page_height_mm = format_info["height_mm"]
                    page_result.page_orientation = format_info["orientation"]
                    page_result.page_a4_sheets_equivalent = format_info["a4_sheets"]
                    
                    # Обновляем статистику форматов
                    format_name = format_info["format_name"]
                    if format_name not in document_format_stats["formats"]:
                        document_format_stats["formats"][format_name] = 0
                    document_format_stats["formats"][format_name] += 1
                    
                    # Обновляем статистику ориентаций
                    orientation = format_info["orientation"]
                    document_format_stats["orientations"][orientation] += 1
                    
                    # Логируем информацию о формате страницы
                    logger.info(f"Page {page_number} format: {format_name} "
                              f"({format_info['width_mm']}x{format_info['height_mm']} mm, "
                              f"{orientation}, {format_info['a4_sheets']} A4 sheets)")
                    
                    # Определяем тип страницы
                    page_type = self.detect_page_type(page)
                    page_result.page_type = page_type
                    document_format_stats["page_types"][page_type] += 1
                    logger.info(f"Page {page_number} type: {page_type}")
                    
                    if page_type == "vector":
                        # Обработка векторной страницы
                        text = page.extract_text()
                        if text.strip():
                            page_result.page_text = text
                            page_result.page_text_confidence = 0.9
                            page_result.page_text_method = "direct_extraction"
                            page_result.page_text_length = len(text)
                            logger.info(f"Page {page_number}: Extracted {len(text)} characters (vector)")
                        else:
                            logger.warning(f"Page {page_number}: No text extracted from vector page")
                            page_result.page_text = ""
                            page_result.page_text_confidence = 0.0
                            page_result.page_text_method = "failed"
                            page_result.page_text_length = 0
                    
                    elif page_type == "scanned" and page_num < len(images):
                        # Обработка сканированной страницы с OCR
                        logger.info(f"🔍 [DEBUG] DocumentParser: Page {page_number}: Processing as scanned page with OCR")
                    
                        try:
                            logger.info(f"🔍 [DEBUG] DocumentParser: Page {page_number}: Starting OCR processing")
                            # Извлекаем текст с помощью OCR
                            ocr_result = self.extract_text_from_image(images[page_num], page_number)
                            logger.info(f"🔍 [DEBUG] DocumentParser: Page {page_number}: OCR processing completed")
                            
                            if ocr_result["text"]:
                                page_result.page_text = ocr_result["text"]
                                page_result.page_text_confidence = ocr_result["confidence"]
                                page_result.page_text_method = f"ocr_{ocr_result['method']}"
                                page_result.page_text_length = len(ocr_result["text"])
                                page_result.page_ocr_confidence = ocr_result["confidence"]
                                page_result.page_ocr_method = ocr_result["method"]
                                page_result.page_ocr_all_results = ocr_result.get("all_results", [])
                                logger.info(f"Page {page_number}: OCR extracted {len(ocr_result['text'])} characters "
                                          f"(confidence: {ocr_result['confidence']})")
                            else:
                                logger.warning(f"Page {page_number}: OCR failed to extract text")
                                page_result.page_text = f"[OCR не смог извлечь текст со страницы {page_number}]"
                                page_result.page_text_confidence = 0.1
                                page_result.page_text_method = "ocr_failed"
                                page_result.page_text_length = 0
                        
                        except Exception as e:
                            logger.error(f"OCR processing failed for page {page_number}: {e}")
                            page_result.page_text = f"[Ошибка OCR обработки страницы {page_number}: {str(e)}]"
                            page_result.page_text_confidence = 0.0
                            page_result.page_text_method = "ocr_error"
                            page_result.page_text_length = 0
                    
                    else:
                        # Неизвестный тип страницы или нет изображения
                        logger.warning(f"Page {page_number}: Unknown type or no image available")
                        text = page.extract_text()
                        if text.strip():
                            page_result.page_text = text
                            page_result.page_text_confidence = 0.5
                            page_result.page_text_method = "fallback_extraction"
                            page_result.page_text_length = len(text)
                        else:
                            page_result.page_text = f"[Не удалось обработать страницу {page_number}]"
                            page_result.page_text_confidence = 0.0
                            page_result.page_text_method = "failed"
                            page_result.page_text_length = 0
                
                    # Добавляем результат страницы в общий результат
                    logger.info(f"🔍 [DEBUG] DocumentParser: Adding page {page_result.page_number} to results. Total results before: {len(result.document_pages_results)}")
                    result.document_pages_results.append(page_result)
                    logger.info(f"🔍 [DEBUG] DocumentParser: Added page {page_result.page_number} to results. Total results after: {len(result.document_pages_results)}")
                    
                    # Сохраняем только страницы с содержимым
                    if page_result.page_text and len(page_result.page_text.strip()) > 0:
                        logger.info(f"🔍 [DEBUG] DocumentParser: Page {page_result.page_number} has content ({len(page_result.page_text)} chars)")
                    else:
                        logger.warning(f"🔍 [DEBUG] DocumentParser: Page {page_result.page_number} has no content!")
                    
                    # TODO: Добавить извлечение таблиц, изображений, штампов
                    # с помощью OpenCV и дополнительного анализа
                    
                    # Мониторинг памяти для каждой страницы
                    log_memory_usage(f"processing page {page_number}")
                    
                    # Очистка памяти каждые 3 страницы в батче
                    if (page_num - start_page) % 3 == 0:
                        cleanup_memory()
                
                # Завершение обработки батча
                logger.info(f"🔍 [DEBUG] DocumentParser: Completed batch {batch_num + 1}/{total_batches}")
                log_memory_usage(f"after batch {batch_num + 1}")
                cleanup_memory()
            
            # Заполняем общую статистику документа
            result.document_pages = document_format_stats['total_pages']
            result.document_pages_total = document_format_stats['total_pages']
            result.document_pages_total_a4_sheets_equivalent = document_format_stats['total_a4_sheets']
            result.document_pages_vector = document_format_stats['page_types'].get('vector', 0)
            result.document_pages_scanned = document_format_stats['page_types'].get('scanned', 0)
            result.document_pages_unknown = document_format_stats['page_types'].get('unknown', 0)
            
            # Логируем итоговую статистику документа
            logger.info("=== PDF DOCUMENT FORMAT STATISTICS ===")
            logger.info(f"Total pages: {document_format_stats['total_pages']}")
            logger.info(f"Total A4 sheets equivalent: {document_format_stats['total_a4_sheets']:.2f}")
            logger.info(f"Page formats: {document_format_stats['formats']}")
            logger.info(f"Orientations: {document_format_stats['orientations']}")
            logger.info(f"Page types: {document_format_stats['page_types']}")
            logger.info("=====================================")
                
        except Exception as e:
            logger.error(f"🔍 [DEBUG] DocumentParser: PDF parsing error: {e}")
            import traceback
            logger.error(f"🔍 [DEBUG] DocumentParser: PDF parsing traceback: {traceback.format_exc()}")
            raise
        
        log_memory_usage("after PDF parsing")
        cleanup_memory()
        logger.info(f"🔍 [DEBUG] DocumentParser: PDF parsing completed successfully. Total pages: {len(result.document_pages_results)}")
        return result
    
    def parse_docx(self, file_content: bytes) -> List[Dict[str, Any]]:
        """Парсинг DOCX документа"""
        elements = []
        
        try:
            # Сохраняем временный файл
            temp_path = "/app/temp/temp.docx"
            with open(temp_path, "wb") as f:
                f.write(file_content)
            
            doc = Document(temp_path)
            
            for para in doc.paragraphs:
                if para.text.strip():
                    elements.append({
                        "element_type": "text",
                        "element_content": para.text,
                        "page_number": 1,  # DOCX не имеет страниц
                        "confidence_score": 0.95
                    })
            
            # Извлечение таблиц
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    elements.append({
                        "element_type": "table",
                        "element_content": json.dumps(table_data),
                        "page_number": 1,
                        "confidence_score": 0.9
                    })
            
            # Удаляем временный файл
            os.remove(temp_path)
            
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            raise
        
        return elements
    
    def parse_dwg(self, file_content: bytes) -> List[Dict[str, Any]]:
        """Парсинг DWG документа"""
        elements = []
        
        try:
            # Простое извлечение текста из DWG файла
            # В реальной реализации здесь будет ezdxf
            elements.append({
                "element_type": "text",
                "element_content": "DWG файл загружен (парсинг не реализован)",
                "page_number": 1,
                "confidence_score": 0.5
            })
            
        except Exception as e:
            logger.error(f"DWG parsing error: {e}")
            raise
        
        return elements
    
    def parse_ifc(self, file_content: bytes) -> List[Dict[str, Any]]:
        """Парсинг IFC документа с продвинутым текстовым анализом"""
        elements = []
        
        try:
            # Декодируем содержимое файла
            text_content = file_content.decode('utf-8', errors='ignore')
            
            # Счетчики для статистики
            entity_counts = {}
            project_info = None
            buildings = []
            storeys = []
            walls = []
            windows = []
            doors = []
            materials = []
            property_sets = []
            
            # Регулярные выражения для поиска IFC элементов
            entity_pattern = r'#(\d+)\s*=\s*([A-Z_]+)\s*\('
            name_pattern = r'Name\s*=\s*[\'"]([^\'"]*)[\'"]'
            description_pattern = r'Description\s*=\s*[\'"]([^\'"]*)[\'"]'
            object_type_pattern = r'ObjectType\s*=\s*[\'"]([^\'"]*)[\'"]'
            elevation_pattern = r'Elevation\s*=\s*([-\d.]+)'
            
            # Обрабатываем весь текст как единое целое для поиска IFC сущностей
            entity_matches = re.finditer(entity_pattern, text_content)
            
            for entity_match in entity_matches:
                entity_id = entity_match.group(1)
                entity_type = entity_match.group(2)
                
                # Находим конец записи (закрывающую скобку)
                start_pos = entity_match.end()
                bracket_count = 0
                end_pos = start_pos
                
                for i, char in enumerate(text_content[start_pos:], start_pos):
                    if char == '(':
                        bracket_count += 1
                    elif char == ')':
                        bracket_count -= 1
                        if bracket_count < 0:
                            end_pos = i + 1
                            break
                
                entity_params = text_content[start_pos:end_pos]
                
                # Подсчитываем типы сущностей
                entity_counts[entity_type] = entity_counts.get(entity_type, 0) + 1
                
                # Извлекаем информацию в зависимости от типа сущности
                if entity_type == 'IFCPROJECT':
                    name_match = re.search(name_pattern, entity_params)
                    desc_match = re.search(description_pattern, entity_params)
                    project_name = name_match.group(1) if name_match else 'Без названия'
                    project_desc = desc_match.group(1) if desc_match else 'Без описания'
                    project_info = f"Проект: {project_name} - {project_desc}"
                    
                elif entity_type == 'IFCBUILDING':
                    name_match = re.search(name_pattern, entity_params)
                    building_name = name_match.group(1) if name_match else 'Без названия'
                    buildings.append(f"Здание: {building_name}")
                    
                elif entity_type == 'IFCBUILDINGSTOREY':
                    name_match = re.search(name_pattern, entity_params)
                    elevation_match = re.search(elevation_pattern, entity_params)
                    storey_name = name_match.group(1) if name_match else 'Без названия'
                    elevation = elevation_match.group(1) if elevation_match else 'Не указана'
                    storeys.append(f"Этаж: {storey_name} - Высота: {elevation}")
                    
                elif entity_type == 'IFCWALL':
                    name_match = re.search(name_pattern, entity_params)
                    type_match = re.search(object_type_pattern, entity_params)
                    wall_name = name_match.group(1) if name_match else 'Без названия'
                    wall_type = type_match.group(1) if type_match else 'Не указан'
                    walls.append(f"Стена: {wall_name} - Тип: {wall_type}")
                    
                elif entity_type == 'IFCWINDOW':
                    name_match = re.search(name_pattern, entity_params)
                    type_match = re.search(object_type_pattern, entity_params)
                    window_name = name_match.group(1) if name_match else 'Без названия'
                    window_type = type_match.group(1) if type_match else 'Не указан'
                    windows.append(f"Окно: {window_name} - Тип: {window_type}")
                    
                elif entity_type == 'IFCDOOR':
                    name_match = re.search(name_pattern, entity_params)
                    type_match = re.search(object_type_pattern, entity_params)
                    door_name = name_match.group(1) if name_match else 'Без названия'
                    door_type = type_match.group(1) if type_match else 'Не указан'
                    doors.append(f"Дверь: {door_name} - Тип: {door_type}")
                    
                elif entity_type == 'IFCMATERIAL':
                    name_match = re.search(name_pattern, entity_params)
                    desc_match = re.search(description_pattern, entity_params)
                    material_name = name_match.group(1) if name_match else 'Без названия'
                    material_desc = desc_match.group(1) if desc_match else 'Без описания'
                    materials.append(f"Материал: {material_name} - {material_desc}")
                        
                elif entity_type == 'IFCPROPERTYSET':
                    name_match = re.search(name_pattern, entity_params)
                    desc_match = re.search(description_pattern, entity_params)
                    pset_name = name_match.group(1) if name_match else 'Без названия'
                    pset_desc = desc_match.group(1) if desc_match else 'Без описания'
                    property_sets.append(f"Набор свойств: {pset_name} - {pset_desc}")
            
            # Добавляем извлеченную информацию в элементы
            if project_info:
                elements.append({
                    "element_type": "project_info",
                    "element_content": project_info,
                    "page_number": 1,
                    "confidence_score": 0.95
                })
            
            for building in buildings[:10]:  # Ограничиваем количество
                elements.append({
                    "element_type": "building",
                    "element_content": building,
                    "page_number": 1,
                    "confidence_score": 0.9
                })
            
            for storey in storeys[:10]:
                elements.append({
                    "element_type": "storey",
                    "element_content": storey,
                    "page_number": 1,
                    "confidence_score": 0.9
                })
            
            for wall in walls[:20]:
                elements.append({
                    "element_type": "wall",
                    "element_content": wall,
                    "page_number": 1,
                    "confidence_score": 0.85
                })
            
            for window in windows[:20]:
                elements.append({
                    "element_type": "window",
                    "element_content": window,
                    "page_number": 1,
                    "confidence_score": 0.85
                })
            
            for door in doors[:20]:
                elements.append({
                    "element_type": "door",
                    "element_content": door,
                    "page_number": 1,
                    "confidence_score": 0.85
                })
            
            for material in materials[:10]:
                elements.append({
                    "element_type": "material",
                    "element_content": material,
                    "page_number": 1,
                    "confidence_score": 0.8
                })
            
            for pset in property_sets[:10]:
                elements.append({
                    "element_type": "property_set",
                    "element_content": pset,
                    "page_number": 1,
                    "confidence_score": 0.8
                })
            
            # Статистика модели
            total_entities = sum(entity_counts.values())
            top_entities = sorted(entity_counts.items(), key=lambda x: x[1], reverse=True)[:10]
            stats_text = f"Всего сущностей: {total_entities}. Топ типов: " + ", ".join([f"{entity}({count})" for entity, count in top_entities])
            
            elements.append({
                "element_type": "statistics",
                "element_content": stats_text,
                "page_number": 1,
                "confidence_score": 0.95
            })
            
            # Если не нашли структурированных элементов, сохраняем базовую информацию
            if len(elements.document_pages_results) <= 1:  # Только статистика
                elements.append({
                    "element_type": "text",
                    "element_content": f"IFC файл содержит {total_entities} сущностей различных типов",
                    "page_number": 1,
                    "confidence_score": 0.7
                })
            
        except Exception as e:
            logger.error(f"IFC parsing error: {e}")
            # Fallback к простому текстовому парсингу
            try:
                text_content = file_content.decode('utf-8', errors='ignore')
                elements.append({
                    "element_type": "text",
                    "element_content": f"IFC файл (продвинутый парсинг не удался): {text_content[:500]}",
                    "page_number": 1,
                    "confidence_score": 0.3
                })
            except:
                elements.append({
                    "element_type": "error",
                    "element_content": f"Ошибка парсинга IFC файла: {str(e)}",
                    "page_number": 1,
                    "confidence_score": 0.1
                })
        
        return elements
    
    def save_to_database(self, filename: str, original_filename: str, file_type: str,
                         file_size: int, document_hash: str, inspection_result: DocumentInspectionResult,
                         category: str = "other", document_type: str = "normative") -> int:
        """Сохранение документа и элементов в базу данных"""
        logger.info(f"🔍 [DEBUG] DocumentParser: save_to_database called")
        logger.info(f"🔍 [DEBUG] DocumentParser: document_type: {document_type}")
        logger.info(f"🔍 [DEBUG] DocumentParser: pages to save: {len(inspection_result.document_pages_results)}")
        
        try:
            with self.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
                if document_type == "checkable":
                    # Сохраняем как проверяемый документ
                    logger.info(f"🔍 [DEBUG] DocumentParser: Saving as checkable document...")
                    review_deadline = datetime.now() + timedelta(days=2)
                    logger.info(f"🔍 [DEBUG] DocumentParser: Review deadline: {review_deadline}")
                    
                    cursor.execute("""
                        INSERT INTO checkable_documents 
                        (filename, original_filename, file_type, file_size, document_hash, 
                         processing_status, category, review_deadline)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                        RETURNING id
                    """, (filename, original_filename, file_type, file_size, document_hash, 
                          "completed", category, review_deadline))
                    document_id = cursor.fetchone()["id"]
                    logger.info(f"🔍 [DEBUG] DocumentParser: Checkable document saved with ID: {document_id}")
                else:
                    # Сохраняем как нормативный документ
                    # Подсчитываем количество токенов
                    token_count = self.calculate_document_tokens(inspection_result)
                    
                    cursor.execute("""
                        INSERT INTO uploaded_documents 
                        (filename, original_filename, file_type, file_size, document_hash, 
                         processing_status, category, token_count)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                        RETURNING id
                    """, (filename, original_filename, file_type, file_size, document_hash, 
                          "completed", category, token_count))
                    document_id = cursor.fetchone()["id"]
                
                # Сохранение элементов
                logger.info(f"🔍 [DEBUG] DocumentParser: Starting to save {len(inspection_result.document_pages_results)} elements...")
                for i, page_result in enumerate(inspection_result.document_pages_results):
                    logger.debug(f"🔍 [DEBUG] DocumentParser: Saving element {i+1}/{len(inspection_result.document_pages_results)}, page {page_result.page_number}")
                    
                    # Подготавливаем дополнительные данные
                    metadata = {
                        "page_type": page_result.page_type,
                        "processing_method": page_result.page_text_method,
                        "format_info": {
                            "format_name": page_result.page_format,
                            "width_mm": page_result.page_width_mm,
                            "height_mm": page_result.page_height_mm,
                            "orientation": page_result.page_orientation,
                            "a4_sheets": page_result.page_a4_sheets_equivalent
                        }
                    }
                    
                    if page_result.page_ocr_confidence is not None:
                        metadata["ocr_confidence"] = page_result.page_ocr_confidence
                    if page_result.page_ocr_method is not None:
                        metadata["ocr_method"] = page_result.page_ocr_method
                    
                    if document_type == "checkable":
                        logger.debug(f"🔍 [DEBUG] DocumentParser: Inserting into checkable_elements, text length: {len(page_result.page_text)}")
                        cursor.execute("""
                            INSERT INTO checkable_elements 
                            (checkable_document_id, element_type, element_content, page_number, confidence_score, element_metadata)
                            VALUES (%s, %s, %s, %s, %s, %s)
                        """, (
                            document_id,
                            "text",
                            page_result.page_text,
                            page_result.page_number,
                            page_result.page_text_confidence,
                            json.dumps(metadata)
                        ))
                    else:
                        logger.debug(f"🔍 [DEBUG] DocumentParser: Inserting into extracted_elements, text length: {len(page_result.page_text)}")
                        cursor.execute("""
                            INSERT INTO extracted_elements 
                            (uploaded_document_id, element_type, element_content, page_number, confidence_score, metadata)
                            VALUES (%s, %s, %s, %s, %s, %s)
                        """, (
                            document_id,
                            "text",
                            page_result.page_text,
                            page_result.page_number,
                            page_result.page_text_confidence,
                            json.dumps(metadata)
                        ))
                
                self.db_conn.commit()
                logger.info(f"🔍 [DEBUG] DocumentParser: Database commit successful")
                
                # Автоматическая проверка нормоконтроля для проверяемых документов
                if document_type == "checkable":
                    logger.info(f"🔍 [DEBUG] DocumentParser: Starting automatic norm control check...")
                    try:
                        # Объединяем содержимое для проверки
                        document_content = "\n\n".join([page_result.page_text for page_result in inspection_result.document_pages_results])
                        logger.debug(f"🔍 [DEBUG] DocumentParser: Document content length for norm control: {len(document_content)} characters")
                        
                        # Запускаем проверку в фоновом режиме
                        asyncio.create_task(self.perform_norm_control_check(document_id, document_content))
                        logger.info(f"🔍 [DEBUG] DocumentParser: Started automatic norm control check for document {document_id}")
                    except Exception as e:
                        logger.error(f"🔍 [DEBUG] DocumentParser: Failed to start norm control check for document {document_id}: {e}")
                
                logger.info(f"🔍 [DEBUG] DocumentParser: save_to_database completed successfully, returning document_id: {document_id}")
                return document_id
                
        except Exception as e:
            self.db_conn.rollback()
            logger.error(f"Database save error: {e}")
            raise
    
    def create_initial_document_record(self, filename: str, file_type: str, file_size: int, document_hash: str, file_path: str, category: str = "other") -> int:
        """Создание начальной записи документа в базе данных"""
        def _create_record(conn):
            with conn.cursor(cursor_factory=RealDictCursor) as cursor:
                cursor.execute("""
                    INSERT INTO uploaded_documents 
                    (filename, original_filename, file_type, file_size, document_hash, processing_status, file_path, category)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    RETURNING id
                """, (filename, filename, file_type, file_size, document_hash, "uploaded", file_path, category))
                
                document_id = cursor.fetchone()["id"]
                logger.info(f"Created initial document record with ID: {document_id}")
                return document_id
        
        try:
            return self.execute_in_transaction(_create_record)
        except Exception as e:
            logger.error(f"Error creating initial document record: {e}")
            raise
    
    def update_document_status(self, document_id: int, status: str):
        """Обновление статуса документа"""
        def _update_status(conn):
            with conn.cursor() as cursor:
                cursor.execute("""
                    UPDATE uploaded_documents 
                    SET processing_status = %s
                    WHERE id = %s
                """, (status, document_id))
                logger.info(f"Updated document {document_id} status to: {status}")
        
        try:
            self.execute_in_transaction(_update_status)
        except Exception as e:
            logger.error(f"Error updating document status: {e}")
            raise
    
    def update_checkable_document_status(self, document_id: int, status: str):
        """Обновление статуса проверяемого документа"""
        def _update_status(conn):
            with conn.cursor() as cursor:
                cursor.execute("""
                    UPDATE checkable_documents 
                    SET processing_status = %s
                    WHERE id = %s
                """, (status, document_id))
                logger.info(f"Updated checkable document {document_id} status to: {status}")
        
        try:
            self.execute_in_transaction(_update_status)
        except Exception as e:
            logger.error(f"Error updating checkable document status: {e}")
            raise
    
    def save_elements_to_database(self, document_id: int, inspection_result: DocumentInspectionResult):
        """Сохранение элементов документа в базу данных"""
        def _save_elements(conn):
            with conn.cursor() as cursor:
                # Очищаем старые элементы для этого документа
                cursor.execute("DELETE FROM extracted_elements WHERE uploaded_document_id = %s", (document_id,))
                logger.info(f"Cleared old elements for document {document_id}")
                
                # Сохранение элементов
                logger.info(f"Saving new document information {document_id}")
                logger.info(f"Saving new document information {inspection_result.document_pages_results}")
                # Сохранение элементов
                logger.info(f"🔍 [DEBUG] DocumentParser: Starting to save {len(inspection_result.document_pages_results)} page results")
                for i, page_result in enumerate(inspection_result.document_pages_results):
                    logger.info(f"🔍 [DEBUG] DocumentParser: Saving page {page_result.page_number} (index {i}), text length: {len(page_result.page_text or '')}")
                    # Подготавливаем дополнительные данные
                    metadata = {
                        "page_type": page_result.page_type,
                        "processing_method": page_result.page_text_method,
                        "format_info": {
                            "format_name": page_result.page_format,
                            "width_mm": page_result.page_width_mm,
                            "height_mm": page_result.page_height_mm,
                            "orientation": page_result.page_orientation,
                            "a4_sheets": page_result.page_a4_sheets_equivalent
                        }
                    }
                    
                    if page_result.page_ocr_confidence is not None:
                        metadata["ocr_confidence"] = page_result.page_ocr_confidence
                    if page_result.page_ocr_method is not None:
                        metadata["ocr_method"] = page_result.page_ocr_method
                    
                    cursor.execute("""
                        INSERT INTO extracted_elements 
                        (uploaded_document_id, element_type, element_content, page_number, confidence_score, metadata)
                        VALUES (%s, %s, %s, %s, %s, %s)
                    """, (
                        document_id,
                        "text",
                        page_result.page_text,
                        page_result.page_number,
                        page_result.page_text_confidence,
                        json.dumps(metadata)
                    ))
                
                logger.info(f"Saved {len(inspection_result.document_pages_results)} elements for document {document_id}")
        
        try:
            self.execute_in_transaction(_save_elements)
        except Exception as e:
            logger.error(f"Error saving elements to database: {e}")
            raise
    
    def update_document_completion(self, document_id: int, elements_count: int, token_count: int):
        """Обновление записи документа после завершения обработки"""
        def _update_completion(conn):
            with conn.cursor() as cursor:
                cursor.execute("""
                    UPDATE uploaded_documents 
                    SET processing_status = %s, token_count = %s
                    WHERE id = %s
                """, ("completed", token_count, document_id))
                logger.info(f"Updated document {document_id} completion: {elements_count} elements, {token_count} tokens")
        
        try:
            self.execute_in_transaction(_update_completion)
        except Exception as e:
            logger.error(f"Error updating document completion: {e}")
            raise
    
    async def index_to_rag_service_async(self, document_id: int, document_title: str, inspection_result: DocumentInspectionResult):
        """Асинхронная индексация в RAG-сервис"""
        try:
            logger.info(f"Starting async RAG indexing for document {document_id}")
            
            # Подготавливаем данные для индексации
            elements = []
            for page_result in inspection_result.document_pages_results:
                elements.append({
                    "element_type": "text",
                    "element_content": page_result.page_text,
                    "page_number": page_result.page_number,
                    "confidence_score": page_result.page_text_confidence
                })
            
            # Объединяем текст всех страниц
            combined_text = ""
            for page_result in inspection_result.document_pages_results:
                if page_result.page_text:
                    combined_text += page_result.page_text + "\n\n"
            
            # Шаг 4: Индексация документа в RAG-сервис
            trace_logger.log_indexing_start(document_id, len(inspection_result.document_pages_results))
            trace_logger.log_step("INDEXING_START", f"Отправка в RAG-сервис", document_id)
            
            # Отправляем данные в RAG-сервис
            async with httpx.AsyncClient(timeout=300.0) as client:
                response = await client.post(
                    f"{RAG_SERVICE_URL}/index",
                    data={
                        "document_id": document_id,
                        "document_title": document_title,
                        "content": combined_text,
                        "chapter": "",
                        "section": "",
                        "page_number": 1
                    }
                )
                
                if response.status_code == 200:
                    # Подсчитываем токены
                    tokens_count = self.calculate_document_tokens(inspection_result)
                    trace_logger.log_indexing_success(document_id, 0.0, tokens_count)  # Время будет добавлено позже
                    trace_logger.log_step("INDEXING_SUCCESS", f"Индексация завершена, {tokens_count} токенов", document_id)
                    logger.info(f"Successfully indexed document {document_id} to RAG service")
                else:
                    trace_logger.log_error("INDEXING", f"Ошибка RAG-сервиса: {response.status_code}", document_id)
                    logger.error(f"Failed to index document {document_id} to RAG service: {response.status_code}")
                    
        except Exception as e:
            logger.error(f"Error in async RAG indexing for document {document_id}: {e}")
    
    def calculate_document_tokens(self, inspection_result: DocumentInspectionResult) -> int:
        """Подсчет токенов в документе"""
        try:
            total_tokens = 0
            encoding = tiktoken.get_encoding("cl100k_base")  # GPT-4 encoding
            
            for page_result in inspection_result.document_pages_results:
                if page_result.page_text:
                    tokens = encoding.encode(page_result.page_text)
                    total_tokens += len(tokens)
            
            logger.info(f"Calculated {total_tokens} tokens for document")
            return total_tokens
            
        except Exception as e:
            logger.error(f"Error calculating document tokens: {e}")
            return 0
            logger.error(f"Database save error: {e}")
            raise

    def save_checkable_document(self, filename: str, original_filename: str, file_type: str, 
                               file_size: int, document_hash: str, inspection_result: DocumentInspectionResult, 
                               category: str = "other") -> int:
        """Сохранение проверяемого документа"""
        logger.info(f"🔍 [DEBUG] DocumentParser: save_checkable_document called")
        logger.info(f"🔍 [DEBUG] DocumentParser: filename: {filename}")
        logger.info(f"🔍 [DEBUG] DocumentParser: original_filename: {original_filename}")
        logger.info(f"🔍 [DEBUG] DocumentParser: file_type: {file_type}")
        logger.info(f"🔍 [DEBUG] DocumentParser: file_size: {file_size}")
        logger.info(f"🔍 [DEBUG] DocumentParser: document_hash: {document_hash}")
        logger.info(f"🔍 [DEBUG] DocumentParser: category: {category}")
        logger.info(f"🔍 [DEBUG] DocumentParser: inspection_result pages: {len(inspection_result.document_pages_results)}")
        
        document_id = self.save_to_database(filename, original_filename, file_type, file_size, 
                                   document_hash, inspection_result, category, "checkable")
        
        logger.info(f"🔍 [DEBUG] DocumentParser: save_checkable_document completed, document_id: {document_id}")
        return document_id

    def cleanup_expired_documents(self) -> int:
        """Очистка просроченных проверяемых документов"""
        def _cleanup(conn):
            with conn.cursor() as cursor:
                cursor.execute("SELECT cleanup_expired_checkable_documents()")
                result = cursor.fetchone()
                return result[0] if result else 0
        
        try:
            return self.execute_in_transaction(_cleanup)
        except Exception as e:
            logger.error(f"Cleanup error: {e}")
            return 0

    def get_checkable_documents(self) -> List[Dict[str, Any]]:
        """Получение списка документов подлежащих нормоконтролю"""
        def _get_documents(conn):
            logger.debug(f"🔍 [DATABASE] _get_documents called")
            try:
                with conn.cursor(cursor_factory=RealDictCursor) as cursor:
                    cursor.execute("""
                        SELECT id, original_filename, file_type, file_size, upload_date, 
                               processing_status, category, review_deadline, review_status, 
                               assigned_reviewer
                        FROM checkable_documents 
                        ORDER BY upload_date DESC
                    """)
                    documents = cursor.fetchall()
                    logger.debug(f"🔍 [DATABASE] Retrieved {len(documents)} checkable documents")
                    return [dict(doc) for doc in documents]
            except Exception as db_error:
                logger.error(f"🔍 [DATABASE] Error in _get_documents: {db_error}")
                raise
        
        try:
            # Проверяем давление на память перед выполнением
            if check_memory_pressure():
                logger.warning("🔍 [MEMORY] High memory pressure detected, performing cleanup")
                cleanup_memory()
            
            # Используем транзакцию только для чтения для оптимизации производительности
            logger.info(f"🔍 [DATABASE] Starting read-only transaction for get_checkable_documents")
            result = self.execute_in_read_only_transaction(_get_documents)
            logger.info(f"🔍 [DATABASE] Successfully retrieved {len(result)} checkable documents using read-only transaction")
            return result
        except Exception as e:
            logger.error(f"🔍 [DATABASE] Get checkable documents error: {e}")
            # Возвращаем пустой список вместо падения
            return []

    def get_checkable_document(self, document_id: int) -> Optional[Dict[str, Any]]:
        """Получение информации о проверяемом документе"""
        def _get_document(conn):
            try:
                with conn.cursor(cursor_factory=RealDictCursor) as cursor:
                    cursor.execute("""
                        SELECT id, original_filename, file_type, file_size, upload_date, 
                               processing_status, category, review_deadline, review_status, 
                               assigned_reviewer
                        FROM checkable_documents 
                        WHERE id = %s
                    """, (document_id,))
                    document = cursor.fetchone()
                    return dict(document) if document else None
            except Exception as db_error:
                logger.error(f"🔍 [DATABASE] Error in _get_document: {db_error}")
                raise
        
        try:
            logger.debug(f"🔍 [DATABASE] Starting read-only transaction for get_checkable_document {document_id}")
            result = self.execute_in_read_only_transaction(_get_document)
            logger.debug(f"🔍 [DATABASE] Successfully retrieved checkable document {document_id} using read-only transaction")
            return result
        except Exception as e:
            logger.error(f"Get checkable document error: {e}")
            return None

    def get_norm_control_result_by_document_id(self, document_id: int) -> Optional[Dict[str, Any]]:
        """Получение результатов нормоконтроля по ID документа"""
        def _get_norm_control_result(conn):
            try:
                with conn.cursor(cursor_factory=RealDictCursor) as cursor:
                    cursor.execute("""
                        SELECT id, analysis_status, total_findings, critical_findings, warning_findings, 
                               info_findings, analysis_date
                        FROM norm_control_results
                        WHERE checkable_document_id = %s
                        ORDER BY analysis_date DESC
                        LIMIT 1
                    """, (document_id,))
                    result = cursor.fetchone()
                    return dict(result) if result else None
            except Exception as db_error:
                logger.error(f"🔍 [DATABASE] Error in _get_norm_control_result: {db_error}")
                raise
        
        try:
            logger.debug(f"🔍 [DATABASE] Starting read-only transaction for get_norm_control_result_by_document_id {document_id}")
            result = self.execute_in_read_only_transaction(_get_norm_control_result)
            logger.debug(f"🔍 [DATABASE] Successfully retrieved norm control result for document {document_id} using read-only transaction")
            return result
        except Exception as e:
            logger.error(f"Get norm control result error: {e}")
            return None

    def get_page_results_by_document_id(self, document_id: int) -> List[Dict[str, Any]]:
        """Получение результатов по страницам документа"""
        def _get_page_results(conn):
            try:
                with conn.cursor(cursor_factory=RealDictCursor) as cursor:
                    cursor.execute("""
                        SELECT 
                            f.id,
                            f.finding_type,
                            f.severity_level,
                            f.category,
                            f.title,
                            f.description,
                            f.recommendation,
                            f.confidence_score,
                            f.created_at
                        FROM findings f
                        JOIN norm_control_results ncr ON f.norm_control_result_id = ncr.id
                        WHERE ncr.checkable_document_id = %s
                        ORDER BY f.severity_level DESC, f.created_at
                    """, (document_id,))
                    results = cursor.fetchall()
                    return [dict(result) for result in results]
            except Exception as db_error:
                logger.error(f"🔍 [DATABASE] Error in _get_page_results: {db_error}")
                raise
        
        try:
            logger.debug(f"🔍 [DATABASE] Starting read-only transaction for get_page_results_by_document_id {document_id}")
            result = self.execute_in_read_only_transaction(_get_page_results)
            logger.debug(f"🔍 [DATABASE] Successfully retrieved page results for document {document_id} using read-only transaction")
            return result
        except Exception as e:
            logger.error(f"Get page results error: {e}")
            return []

    def get_review_report_by_norm_control_id(self, norm_control_id: int) -> Optional[Dict[str, Any]]:
        """Получение отчета рецензента по ID результата нормоконтроля"""
        try:
            with self.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
                cursor.execute("""
                    SELECT id, reviewer_name, conclusion, report_date
                    FROM review_reports
                    WHERE norm_control_result_id = %s
                    ORDER BY report_date DESC
                    LIMIT 1
                """, (norm_control_id,))
                result = cursor.fetchone()
                return dict(result) if result else None
        except Exception as e:
            logger.error(f"Get review report error: {e}")
            return None

    def get_findings_by_norm_control_id(self, norm_control_result_id: int) -> List[Dict[str, Any]]:
        """Получение детальных findings по ID результата нормоконтроля"""
        try:
            with self.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
                cursor.execute("""
                    SELECT f.*, 
                           dc.clause_title as normative_document_title,
                           dc.clause_number as normative_clause_number
                    FROM findings f
                    LEFT JOIN document_clauses dc ON f.related_clause_id = dc.id
                    WHERE f.norm_control_result_id = %s
                    ORDER BY f.severity_level DESC, f.id ASC
                """, (norm_control_result_id,))
                results = cursor.fetchall()
                return [dict(result) for result in results]
        except Exception as e:
            logger.error(f"Get findings error: {e}")
            return []

    def get_findings_by_document_id(self, document_id: int) -> List[Dict[str, Any]]:
        """Получение всех findings для документа"""
        try:
            with self.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
                cursor.execute("""
                    SELECT f.*, 
                           dc.clause_title as normative_document_title,
                           dc.clause_number as normative_clause_number,
                           ncr.analysis_date
                    FROM findings f
                    JOIN norm_control_results ncr ON f.norm_control_result_id = ncr.id
                    LEFT JOIN document_clauses dc ON f.related_clause_id = dc.id
                    WHERE ncr.checkable_document_id = %s
                    ORDER BY f.severity_level DESC, f.id ASC
                """, (document_id,))
                results = cursor.fetchall()
                return [dict(result) for result in results]
        except Exception as e:
            logger.error(f"Get findings by document error: {e}")
            return []

    def get_document_info(self, pdf_reader) -> Dict[str, Any]:
        """Извлечение информации о документе из PDF метаданных"""
        try:
            info = pdf_reader.metadata
            if info:
                # Извлекаем информацию из метаданных PDF
                title = info.get('/Title', 'Неизвестный документ')
                author = info.get('/Author', 'Неизвестный автор')
                subject = info.get('/Subject', '')
                creator = info.get('/Creator', '')
                producer = info.get('/Producer', '')
                creation_date = info.get('/CreationDate', '')
                mod_date = info.get('/ModDate', '')
                
                # Попытка извлечь номер документа из названия
                document_number = ""
                if title:
                    # Ищем паттерны типа "A9.5.MTH.04" или "КЖ-01" и т.д.
                    import re
                    number_match = re.search(r'[A-ZА-Я]{1,3}[-\d\.]+[A-ZА-Я]*\d*', title)
                    if number_match:
                        document_number = number_match.group(0)
                
                # Определяем тип документа по названию
                document_type = "pdf"
                engineering_stage = "ПД"  # По умолчанию
                document_mark = ""
                document_status = "IFR"  # По умолчанию
                
                # Попытка определить марку по названию
                if title:
                    title_upper = title.upper()
                    if any(mark in title_upper for mark in ['КЖ', 'KZH']):
                        document_mark = "КЖ"
                    elif any(mark in title_upper for mark in ['КМ', 'KM']):
                        document_mark = "КМ"
                    elif any(mark in title_upper for mark in ['АС', 'AS']):
                        document_mark = "АС"
                    elif any(mark in title_upper for mark in ['ТХ', 'TX']):
                        document_mark = "ТХ"
                    elif any(mark in title_upper for mark in ['КС', 'KS']):
                        document_mark = "КС"
                    elif any(mark in title_upper for mark in ['КП', 'KP']):
                        document_mark = "КП"
                    elif any(mark in title_upper for mark in ['КР', 'KR']):
                        document_mark = "КР"
                
                return {
                    "name": title,
                    "type": document_type,
                    "engineering_stage": engineering_stage,
                    "mark": document_mark,
                    "number": document_number,
                    "date": creation_date or mod_date,
                    "author": author,
                    "reviewer": "",  # Не извлекается из PDF
                    "approver": "",  # Не извлекается из PDF
                    "status": document_status,
                    "size": 0  # Будет заполнено позже
                }
            else:
                # Если метаданные отсутствуют, возвращаем значения по умолчанию
                return {
                    "name": "Документ без названия",
                    "type": "pdf",
                    "engineering_stage": "ПД",
                    "mark": "",
                    "number": "",
                    "date": "",
                    "author": "Неизвестный автор",
                    "reviewer": "",
                    "approver": "",
                    "status": "IFR",
                    "size": 0
                }
        except Exception as e:
            logger.error(f"Error extracting document info: {e}")
            return {
                "name": "Документ с ошибкой метаданных",
                "type": "pdf",
                "engineering_stage": "ПД",
                "mark": "",
                "number": "",
                "date": "",
                "author": "Неизвестный автор",
                "reviewer": "",
                "approver": "",
                "status": "IFR",
                "size": 0
            }

    def get_document_format_statistics(self, document_id: int) -> Dict[str, Any]:
        """Получение статистики форматов документа"""
        try:
            with self.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
                cursor.execute("""
                    SELECT metadata->>'format_info' as format_info
                    FROM checkable_elements
                    WHERE checkable_document_id = %s
                    ORDER BY page_number
                """, (document_id,))
                elements = cursor.fetchall()
                
                if not elements:
                    return {"error": "Document not found or no elements"}
                
                stats = {
                    "total_pages": len(elements),
                    "total_a4_sheets": 0.0,
                    "formats": {},
                    "orientations": {"portrait": 0, "landscape": 0},
                    "page_types": {"vector": 0, "scanned": 0, "unknown": 0},
                    "pages": []
                }
                
                for element in elements:
                    try:
                        format_info = json.loads(element["format_info"]) if element["format_info"] else {}
                        
                        # Добавляем информацию о странице
                        page_info = {
                            "page_number": format_info.get("page_number", 0),
                            "format": format_info.get("format_name", "unknown"),
                            "width_mm": format_info.get("width_mm", 0),
                            "height_mm": format_info.get("height_mm", 0),
                            "orientation": format_info.get("orientation", "unknown"),
                            "a4_sheets": format_info.get("a4_sheets", 0)
                        }
                        stats["pages"].append(page_info)
                        
                        # Обновляем общую статистику
                        a4_sheets = format_info.get("a4_sheets", 0)
                        if a4_sheets is not None:
                            stats["total_a4_sheets"] += a4_sheets
                        
                        format_name = format_info.get("format_name", "unknown")
                        if format_name not in stats["formats"]:
                            stats["formats"][format_name] = 0
                        stats["formats"][format_name] += 1
                        
                        orientation = format_info.get("orientation", "unknown")
                        if orientation in stats["orientations"]:
                            stats["orientations"][orientation] += 1
                        
                    except Exception as e:
                        logger.warning(f"Error parsing format info: {e}")
                        # Добавляем пустую информацию о странице
                        stats["pages"].append({
                            "page_number": 0,
                            "format": "unknown",
                            "width_mm": 0,
                            "height_mm": 0,
                            "orientation": "unknown",
                            "a4_sheets": 0
                        })
                        continue
                
                # Округляем общее количество листов А4
                stats["total_a4_sheets"] = round(stats["total_a4_sheets"], 2)
                
                return stats
                
        except Exception as e:
            logger.error(f"Get document format statistics error: {e}")
            return {"error": str(e)}

    def determine_document_category(self, filename: str, content: str) -> str:
        """Определение категории документа по имени файла и содержимому"""
        filename_lower = filename.lower()
        content_lower = content.lower()
        
        # Проверка по ключевым словам в названии и содержимом
        if any(keyword in filename_lower for keyword in ['gost', 'гост']):
            return 'gost'
        elif any(keyword in filename_lower for keyword in ['sp', 'сн']):
            return 'sp'
        elif any(keyword in filename_lower for keyword in ['snip', 'снип']):
            return 'snip'
        elif any(keyword in filename_lower for keyword in ['tr', 'тр']):
            return 'tr'
        elif any(keyword in content_lower for keyword in ['корпоративный', 'корп', 'corporate']):
            return 'corporate'
        else:
            return 'other'

    def update_review_status(self, document_id: int, status: str, reviewer: str = None, notes: str = None) -> bool:
        """Обновление статуса проверки документа"""
        def _update_status(conn):
            with conn.cursor() as cursor:
                cursor.execute("""
                    UPDATE checkable_documents 
                    SET review_status = %s, assigned_reviewer = %s, updated_at = CURRENT_TIMESTAMP
                    WHERE id = %s
                """, (status, reviewer, document_id))
                return cursor.rowcount > 0
        
        try:
            return self.execute_in_transaction(_update_status)
        except Exception as e:
            logger.error(f"Update review status error: {e}")
            return False

    async def delete_normative_document(self, document_id: int) -> bool:
        """Удаление нормативного документа и связанных данных"""
        def _delete_document(conn):
            with conn.cursor() as cursor:
                # Получаем информацию о документе для логирования
                cursor.execute("""
                    SELECT original_filename, document_hash 
                    FROM uploaded_documents 
                    WHERE id = %s
                """, (document_id,))
                doc_info = cursor.fetchone()
                
                if not doc_info:
                    logger.warning(f"Document {document_id} not found")
                    return False, None
                
                # Удаляем элементы документа
                cursor.execute("DELETE FROM extracted_elements WHERE uploaded_document_id = %s", (document_id,))
                elements_deleted = cursor.rowcount
                
                # Удаляем результаты нормоконтроля
                cursor.execute("DELETE FROM norm_control_results WHERE uploaded_document_id = %s", (document_id,))
                results_deleted = cursor.rowcount
                
                # Удаляем сам документ
                cursor.execute("DELETE FROM uploaded_documents WHERE id = %s", (document_id,))
                document_deleted = cursor.rowcount
                
                logger.info(f"Deleted normative document {document_id} ({doc_info[0]}): "
                          f"{elements_deleted} elements, {results_deleted} results, {document_deleted} document")
                
                return document_deleted > 0, doc_info[0]
        
        try:
            success, filename = self.execute_in_transaction(_delete_document)
            
            if success:
                # Удаляем индексы из RAG сервиса
                try:
                    async with httpx.AsyncClient() as client:
                        response = await client.delete(f"{RAG_SERVICE_URL}/index-documentes/document/{document_id}")
                        if response.status_code == 200:
                            logger.info(f"Successfully deleted indexes for document {document_id}")
                        else:
                            logger.warning(f"Failed to delete indexes for document {document_id}: {response.status_code}")
                except Exception as e:
                    logger.error(f"Error deleting indexes for document {document_id}: {e}")
            
            return success
                
        except Exception as e:
            logger.error(f"Delete normative document error: {e}")
            return False

    def delete_checkable_document(self, document_id: int) -> bool:
        """Удаление проверяемого документа и связанных данных"""
        def _delete_document(conn):
            with conn.cursor() as cursor:
                # Получаем информацию о документе для логирования
                cursor.execute("""
                    SELECT original_filename, document_hash 
                    FROM checkable_documents 
                    WHERE id = %s
                """, (document_id,))
                doc_info = cursor.fetchone()
                
                if not doc_info:
                    logger.warning(f"Checkable document {document_id} not found")
                    return False
                
                # Удаляем элементы документа
                cursor.execute("DELETE FROM checkable_elements WHERE checkable_document_id = %s", (document_id,))
                elements_deleted = cursor.rowcount
                
                # Удаляем отчеты о проверке
                cursor.execute("DELETE FROM review_reports WHERE checkable_document_id = %s", (document_id,))
                reports_deleted = cursor.rowcount
                
                # Удаляем результаты нормоконтроля
                cursor.execute("DELETE FROM norm_control_results WHERE checkable_document_id = %s", (document_id,))
                results_deleted = cursor.rowcount
                
                # Удаляем сам документ
                cursor.execute("DELETE FROM checkable_documents WHERE id = %s", (document_id,))
                document_deleted = cursor.rowcount
                
                logger.info(f"Deleted checkable document {document_id} ({doc_info[0]}): "
                          f"{elements_deleted} elements, {reports_deleted} reports, "
                          f"{results_deleted} results, {document_deleted} document")
                
                return document_deleted > 0
        
        try:
            return self.execute_in_transaction(_delete_document)
        except Exception as e:
            logger.error(f"Delete checkable document error: {e}")
            return False
    
    async def index_to_rag_service(self, document_id: int, document_title: str, elements: List[Dict[str, Any]]):
        """Индексация документа в RAG-сервис"""
        try:
            # Объединяем все элементы в один текст
            content = "\n\n".join([elem["element_content"] for elem in elements])
            
            # Определяем главу и раздел из элементов
            chapter = ""
            section = ""
            for elem in elements:
                if elem["element_type"] in ["project_info", "building", "storey"]:
                    chapter = elem["element_content"][:100]
                    break
            
            # Отправляем в RAG-сервис
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    f"{RAG_SERVICE_URL}/index",
                    data={
                        "document_id": document_id,
                        "document_title": document_title,
                        "content": content,
                        "chapter": chapter,
                        "section": section,
                        "page_number": 1
                    },
                    timeout=30.0
                )
                
                if response.status_code == 200:
                    result = response.json()
                    logger.info(f"Successfully indexed document {document_id} in RAG service: {result}")
                else:
                    logger.error(f"Failed to index document {document_id} in RAG service: {response.text}")
                    
        except Exception as e:
            logger.error(f"RAG indexing error for document {document_id}: {e}")
            # Не прерываем основной процесс, если RAG индексация не удалась

    def get_system_settings(self) -> List[Dict[str, Any]]:
        """Получение всех настроек системы"""
        try:
            with self.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
                cursor.execute("""
                    SELECT setting_key, setting_value, setting_description, setting_type, is_public, updated_at
                    FROM system_settings
                    WHERE is_public = true
                    ORDER BY setting_key
                """)
                settings = cursor.fetchall()
                return [dict(setting) for setting in settings]
        except Exception as e:
            logger.error(f"Get system settings error: {e}")
            return []

    def get_system_setting(self, setting_key: str) -> Optional[str]:
        """Получение конкретной настройки системы"""
        try:
            with self.db_conn.cursor() as cursor:
                cursor.execute("""
                    SELECT setting_value
                    FROM system_settings
                    WHERE setting_key = %s AND is_public = true
                """, (setting_key,))
                result = cursor.fetchone()
                return result[0] if result else None
        except Exception as e:
            logger.error(f"Get system setting error: {e}")
            return None

    def update_system_setting(self, setting_key: str, setting_value: str) -> bool:
        """Обновление настройки системы"""
        def _update_setting(conn):
            with conn.cursor() as cursor:
                cursor.execute("""
                    UPDATE system_settings
                    SET setting_value = %s, updated_at = CURRENT_TIMESTAMP
                    WHERE setting_key = %s AND is_public = true
                """, (setting_value, setting_key))
                return cursor.rowcount > 0
        
        try:
            return self.execute_in_transaction(_update_setting)
        except Exception as e:
            logger.error(f"Update system setting error: {e}")
            return False

    def create_system_setting(self, setting_key: str, setting_value: str, 
                            setting_description: str, setting_type: str = "text") -> bool:
        """Создание новой настройки системы"""
        def _create_setting(conn):
            with conn.cursor() as cursor:
                cursor.execute("""
                    INSERT INTO system_settings (setting_key, setting_value, setting_description, setting_type)
                    VALUES (%s, %s, %s, %s)
                    ON CONFLICT (setting_key) DO UPDATE SET
                    setting_value = EXCLUDED.setting_value,
                    setting_description = EXCLUDED.setting_description,
                    setting_type = EXCLUDED.setting_type,
                    updated_at = CURRENT_TIMESTAMP
                """, (setting_key, setting_value, setting_description, setting_type))
                return True
        
        try:
            return self.execute_in_transaction(_create_setting)
        except Exception as e:
            logger.error(f"Create system setting error: {e}")
            return False

    def delete_system_setting(self, setting_key: str) -> bool:
        """Удаление настройки системы"""
        def _delete_setting(conn):
            with conn.cursor() as cursor:
                cursor.execute("""
                    DELETE FROM system_settings
                    WHERE setting_key = %s AND is_public = true
                """, (setting_key,))
                return cursor.rowcount > 0
        
        try:
            return self.execute_in_transaction(_delete_setting)
        except Exception as e:
            logger.error(f"Delete system setting error: {e}")
            return False

    def get_normcontrol_prompt_template(self) -> str:
        """Получение полного шаблона промпта для нормоконтроля из системы настроек"""
        try:
            # Получаем основной промпт для нормоконтроля
            normcontrol_prompt = self.get_system_setting("normcontrol_prompt")
            if not normcontrol_prompt:
                logger.warning("normcontrol_prompt не найден в настройках, используем базовый промпт")
                normcontrol_prompt = "Ты - эксперт по нормоконтролю в строительстве и проектировании. Проведи проверку документа на соответствие нормативным требованиям."
            
            # Получаем пользовательский шаблон промпта (если есть)
            prompt_template = self.get_system_setting("normcontrol_prompt_template")
            if prompt_template:
                # Используем пользовательский шаблон
                prompt_template = prompt_template.replace("{normcontrol_prompt}", normcontrol_prompt)
                return prompt_template
            
            # Если пользовательский шаблон не задан, используем основной промпт напрямую
            # Заменяем плейсхолдеры в основном промпте на плейсхолдеры для страниц
            processed_prompt = normcontrol_prompt.replace("{document_content}", "{page_content}")
            processed_prompt = processed_prompt.replace("{normative_docs}", "нормативным требованиям")
            
            # Экранируем все остальные плейсхолдеры в промпте, чтобы избежать конфликтов с форматированием
            processed_prompt = processed_prompt.replace("{", "{{").replace("}", "}}")
            
            # Добавляем инструкции для JSON ответа
            template = f"""
{processed_prompt}

СОДЕРЖИМОЕ СТРАНИЦЫ:
{{page_content}}

ИНСТРУКЦИИ:
1. Проанализируйте содержимое страницы на соответствие нормативным требованиям
2. ВАЖНО: Ответьте ТОЛЬКО валидным JSON без дополнительного текста
3. Не добавляйте никаких комментариев или пояснений вне JSON

ТРЕБУЕМЫЙ ФОРМАТ JSON:
{{
  "page_number": НОМЕР_СТРАНИЦЫ,
  "overall_status": "pass|fail|uncertain",
  "confidence": 0.0-1.0,
  "total_findings": число,
  "critical_findings": число,
  "warning_findings": число,
  "info_findings": число,
  "compliance_percentage": 0-100,
  "findings": [],
  "summary": "общий_вывод_по_странице",
  "recommendations": "общие_рекомендации_по_улучшению"
}}

ПРИМЕР ОТВЕТА:
{{
  "page_number": 1,
  "overall_status": "pass",
  "confidence": 0.85,
  "total_findings": 0,
  "critical_findings": 0,
  "warning_findings": 0,
  "info_findings": 0,
  "compliance_percentage": 95,
  "findings": [],
  "summary": "Страница соответствует нормативным требованиям",
  "recommendations": "Документ оформлен корректно"
}}
"""
            
            return template
            
        except Exception as e:
            logger.error(f"Get normcontrol prompt template error: {e}")
            # Возвращаем базовый промпт в случае ошибки
            return "Ты - эксперт по нормоконтролю в строительстве и проектировании. Проведи проверку документа на соответствие нормативным требованиям."

    def combine_page_results(self, document_id: int, page_results: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Объединение результатов проверки всех страниц в общий результат"""
        try:
            logger.info(f"🔍 [DEBUG] DocumentParser: Combining results from {len(page_results)} pages for document {document_id}")
            
            # Инициализируем общие счетчики
            total_findings = 0
            critical_findings = 0
            warning_findings = 0
            info_findings = 0
            all_findings = []
            successful_pages = 0
            failed_pages = 0
            
            # Обрабатываем результаты каждой страницы
            for page_result in page_results:
                page_number = page_result.get("page_number", 0)
                
                if page_result.get("status") == "success":
                    successful_pages += 1
                    result_data = page_result.get("result", {})
                    
                    # Собираем замечания со страницы
                    page_findings = result_data.get("findings", [])
                    for finding in page_findings:
                        finding["page_number"] = page_number
                        all_findings.append(finding)
                        
                        # Подсчитываем по типам
                        severity = finding.get("severity", "info")
                        if severity == "critical":
                            critical_findings += 1
                        elif severity == "warning":
                            warning_findings += 1
                        else:
                            info_findings += 1
                    
                    total_findings += len(page_findings)
                    logger.info(f"🔍 [DEBUG] DocumentParser: Page {page_number}: {len(page_findings)} findings")
                    
                else:
                    failed_pages += 1
                    logger.warning(f"🔍 [DEBUG] DocumentParser: Page {page_number} failed: {page_result.get('error', 'Unknown error')}")
            
            # Создаем общий результат
            combined_result = {
                "document_id": document_id,
                "total_pages": len(page_results),
                "successful_pages": successful_pages,
                "failed_pages": failed_pages,
                "total_findings": total_findings,
                "critical_findings": critical_findings,
                "warning_findings": warning_findings,
                "info_findings": info_findings,
                "findings": all_findings,
                "status": "completed" if failed_pages == 0 else "completed_with_errors"
            }
            
            logger.info(f"🔍 [DEBUG] DocumentParser: Combined result: {total_findings} total findings, "
                       f"{critical_findings} critical, {warning_findings} warnings, {info_findings} info")
            
            return combined_result
            
        except Exception as e:
            logger.error(f"🔍 [DEBUG] DocumentParser: Error combining page results: {e}")
            return {
                "document_id": document_id,
                "status": "error",
                "error": str(e)
            }

    def split_document_into_pages(self, document_id: int) -> List[Dict[str, Any]]:
        """Разбиение документа на страницы в соответствии с реальной структурой PDF"""
        logger.info(f"🔍 [DEBUG] DocumentParser: Starting split_document_into_pages for document {document_id}")
        try:
            pages = []
            
            # Получаем элементы документа из базы данных, сгруппированные по страницам
            with self.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
                cursor.execute("""
                    SELECT page_number, element_content, element_type, confidence_score
                    FROM checkable_elements
                    WHERE checkable_document_id = %s
                    ORDER BY page_number, id
                """, (document_id,))
                elements = cursor.fetchall()
            
            if not elements:
                logger.warning(f"No elements found for document {document_id}")
                return []
            
            # Группируем элементы по страницам
            current_page = None
            current_page_content = []
            current_page_elements = []
            
            for element in elements:
                page_number = element["page_number"]
                
                if current_page is None:
                    current_page = page_number
                    current_page_content = []
                    current_page_elements = []
                
                if page_number != current_page:
                    # Сохраняем предыдущую страницу
                    pages.append({
                        "page_number": current_page,
                        "content": "\n\n".join(current_page_content),
                        "char_count": len("\n\n".join(current_page_content)),
                        "element_count": len(current_page_elements),
                        "elements": current_page_elements
                    })
                    
                    # Начинаем новую страницу
                    current_page = page_number
                    current_page_content = []
                    current_page_elements = []
                
                # Добавляем элемент к текущей странице
                current_page_content.append(element["element_content"])
                current_page_elements.append({
                    "type": element["element_type"],
                    "content": element["element_content"],
                    "confidence": element["confidence_score"]
                })
            
            # Добавляем последнюю страницу
            if current_page is not None:
                pages.append({
                    "page_number": current_page,
                    "content": "\n\n".join(current_page_content),
                    "char_count": len("\n\n".join(current_page_content)),
                    "element_count": len(current_page_elements),
                    "elements": current_page_elements
                })
            
            logger.info(f"🔍 [DEBUG] DocumentParser: Split document {document_id} into {len(pages)} pages based on PDF structure")
            logger.info(f"🔍 [DEBUG] DocumentParser: Page details for document {document_id}:")
            for page in pages:
                logger.info(f"🔍 [DEBUG] DocumentParser: - Page {page['page_number']}: {page['char_count']} chars, {page['element_count']} elements")
                logger.info(f"🔍 [DEBUG] DocumentParser: - Page {page['page_number']} content preview: {page['content'][:100]}...")
            
            logger.info(f"🔍 [DEBUG] DocumentParser: split_document_into_pages completed for document {document_id}")
            
            return pages
            
        except Exception as e:
            logger.error(f"Error splitting document into pages: {e}")
            # В случае ошибки возвращаем пустой список
            return []

    async def perform_norm_control_check_for_page(self, document_id: int, page_data: Dict[str, Any]) -> Dict[str, Any]:
        """Выполнение проверки нормоконтроля для одной страницы документа"""
        start_time = time.time()
        page_number = page_data["page_number"]
        
        # Шаг 5: Начало проверки нормоконтроля страницы
        trace_logger.log_normcontrol_start(document_id, 1)  # 1 страница
        trace_logger.log_step("NORMCONTROL_PAGE_START", f"Страница {page_number}, контент: {len(page_data['content'])} символов", document_id)
        
        logger.info(f"🔍 [DEBUG] DocumentParser: Starting norm control check for document {document_id}, page {page_data['page_number']}")
        logger.info(f"🔍 [DEBUG] DocumentParser: Page parameters:")
        logger.info(f"🔍 [DEBUG] DocumentParser: - Page number: {page_data['page_number']}")
        logger.info(f"🔍 [DEBUG] DocumentParser: - Content length: {len(page_data['content'])} characters")
        logger.info(f"🔍 [DEBUG] DocumentParser: - Element count: {page_data['element_count']}")
        logger.info(f"🔍 [DEBUG] DocumentParser: - Content preview: {page_data['content'][:100]}...")
        
        try:
            page_number = page_data["page_number"]
            page_content = page_data["content"]
            
            logger.info(f"🔍 [DEBUG] DocumentParser: Starting norm control check for document {document_id}, page {page_number}")
            logger.info(f"🔍 [DEBUG] DocumentParser: Page content length: {len(page_content)} characters")
            
            # ===== ПОЛУЧЕНИЕ ПРОМПТА ДЛЯ LLM ИЗ СИСТЕМЫ НАСТРОЕК =====
            # Получаем полный шаблон промпта для нормоконтроля из системы настроек
            prompt_template = self.get_normcontrol_prompt_template()
            
            # Формируем запрос к LLM для конкретной страницы с использованием шаблона
            # Заменяем плейсхолдеры вручную, чтобы избежать конфликтов с форматированием
            prompt = prompt_template.replace("{page_content}", page_content)
            
            # ===== ОТПРАВКА ЗАПРОСА К LLM ДЛЯ ПРОВЕРКИ СТРАНИЦЫ =====
            # Отправляем запрос к LLM через gateway для выполнения нормоконтроля
            llm_start_time = time.time()
            trace_logger.log_llm_request(document_id, page_number, len(prompt))
            trace_logger.log_step("LLM_REQUEST", f"Отправка запроса к LLM для страницы {page_number}", document_id)
            
            logger.info(f"Sending request to LLM for page {page_number}...")
            logger.info(f"Prompt length: {len(prompt)} characters")
            
            # Устанавливаем таймаут 10 минут для асинхронного нормоконтроля
            timeout = httpx.Timeout(600.0, connect=30.0)
            async with httpx.AsyncClient(verify=False, timeout=timeout) as client:
                response = await client.post(
                    "http://ollama:11434/api/generate",
                    headers={
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": "llama3.1-optimized-v2:latest",
                        "prompt": f"Ты — эксперт по нормоконтролю проектной документации.\n\n{prompt}",
                        "stream": False
                    }
                )
                
                # ===== ПОЛУЧЕНИЕ РЕЗУЛЬТАТА ПРОВЕРКИ ОТ LLM =====
                if response.status_code == 200:
                    result = response.json()
                    content = result["response"]
                    
                    # Логирование ответа от LLM
                    llm_response_time = time.time() - llm_start_time
                    trace_logger.log_llm_response(document_id, page_number, llm_response_time, len(content))
                    trace_logger.log_step("LLM_RESPONSE", f"Получен ответ от LLM за {llm_response_time:.2f}с", document_id)
                    
                    # Парсим JSON ответ от LLM
                    try:
                        import json
                        import re
                        
                        logger.info(f"🔍 [DEBUG] DocumentParser: Raw LLM response length: {len(content)}")
                        logger.info(f"🔍 [DEBUG] DocumentParser: Raw LLM response preview: {content[:200]}...")
                        
                        # Очищаем ответ от лишних символов
                        cleaned_content = content.strip()
                        
                        # Ищем JSON в ответе (между фигурными скобками)
                        json_match = re.search(r'\{.*\}', cleaned_content, re.DOTALL)
                        if json_match:
                            json_str = json_match.group(0)
                            logger.info(f"🔍 [DEBUG] DocumentParser: Found JSON match: {json_str[:100]}...")
                            check_result = json.loads(json_str)
                        else:
                            # Если JSON не найден, пробуем парсить весь ответ
                            logger.info(f"🔍 [DEBUG] DocumentParser: No JSON match found, trying to parse entire response")
                            check_result = json.loads(cleaned_content)
                        
                        # Проверяем, что результат содержит необходимые поля
                        required_fields = ["page_number", "overall_status", "confidence", "total_findings"]
                        missing_fields = [field for field in required_fields if field not in check_result]
                        
                        if missing_fields:
                            logger.warning(f"🔍 [DEBUG] DocumentParser: Missing required fields: {missing_fields}")
                            # Создаем дефолтный результат
                            check_result = {
                                "page_number": page_number,
                                "overall_status": "uncertain",
                                "confidence": 0.5,
                                "total_findings": 0,
                                "critical_findings": 0,
                                "warning_findings": 0,
                                "info_findings": 0,
                                "compliance_percentage": 50,
                                "findings": [],
                                "summary": "Обработка завершена с предупреждениями",
                                "recommendations": "Требуется дополнительная проверка"
                            }
                        
                        logger.info(f"🔍 [DEBUG] DocumentParser: Successfully parsed JSON result")
                        return {
                            "status": "success",
                            "result": check_result,
                            "raw_response": content
                        }
                    except json.JSONDecodeError as e:
                        logger.error(f"🔍 [DEBUG] DocumentParser: JSON parsing error: {e}")
                        logger.error(f"🔍 [DEBUG] DocumentParser: Raw response: {content}")
                        
                        # Создаем дефолтный результат в случае ошибки парсинга
                        default_result = {
                            "page_number": page_number,
                            "overall_status": "uncertain",
                            "confidence": 0.3,
                            "total_findings": 0,
                            "critical_findings": 0,
                            "warning_findings": 0,
                            "info_findings": 0,
                            "compliance_percentage": 30,
                            "findings": [],
                            "summary": "Ошибка обработки ответа от LLM",
                            "recommendations": "Требуется повторная проверка"
                        }
                        
                        return {
                            "status": "success",
                            "result": default_result,
                            "raw_response": content
                        }
                else:
                    logger.error(f"LLM request failed: {response.status_code} - {response.text}")
                    return {
                        "status": "error",
                        "error": f"LLM request failed: {response.status_code}"
                    }
                    
        except Exception as e:
            logger.error(f"Norm control check error: {e}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return {
                "status": "error",
                "error": str(e)
            }

    async def perform_norm_control_check(self, document_id: int, document_content: str) -> Dict[str, Any]:
        """Выполнение проверки нормоконтроля для документа по страницам с применением LLM"""
        try:
            logger.info(f"Starting norm control check for document {document_id}")
            logger.info(f"Document content length: {len(document_content)} characters")
            
            # ===== ОСНОВНАЯ ФУНКЦИЯ ПРОВЕРКИ НОРМОКОНТРОЛЯ С LLM =====
            
            # Разбиваем документ на страницы в соответствии с реальной структурой PDF
            pages = self.split_document_into_pages(document_id)
            logger.info(f"Document split into {len(pages)} pages based on PDF structure")
            
            # Проверяем каждую страницу отдельно
            page_results = []
            total_findings = 0
            total_critical_findings = 0
            total_warning_findings = 0
            total_info_findings = 0
            all_findings = []
            
            for page_data in pages:
                logger.info(f"Processing page {page_data['page_number']} of {len(pages)}")
                
                # Проверяем страницу
                # ===== ВЫЗОВ ПРОВЕРКИ СТРАНИЦЫ С ПРИМЕНЕНИЕМ LLM =====
                # Выполняем проверку для каждой страницы с использованием LLM
                page_result = await self.perform_norm_control_check_for_page(document_id, page_data)
                
                if page_result["status"] == "success":
                    result_data = page_result["result"]
                    page_results.append(result_data)
                    
                    # Собираем статистику
                    total_findings += result_data.get("total_findings", 0)
                    total_critical_findings += result_data.get("critical_findings", 0)
                    total_warning_findings += result_data.get("warning_findings", 0)
                    total_info_findings += result_data.get("info_findings", 0)
                    
                    # Собираем все findings
                    page_findings = result_data.get("findings", [])
                    for finding in page_findings:
                        finding["page_number"] = page_data["page_number"]
                    all_findings.extend(page_findings)
                    
                    logger.info(f"Page {page_data['page_number']} processed successfully")
                else:
                    logger.error(f"Failed to process page {page_data['page_number']}: {page_result.get('error', 'Unknown error')}")
                    # Добавляем пустой результат для страницы
                    page_results.append({
                        "page_number": page_data["page_number"],
                        "overall_status": "error",
                        "confidence": 0.0,
                        "total_findings": 0,
                        "critical_findings": 0,
                        "warning_findings": 0,
                        "info_findings": 0,
                        "findings": [],
                        "summary": f"Ошибка обработки страницы: {page_result.get('error', 'Unknown error')}",
                        "compliance_percentage": 0
                    })
            
            # Формируем общий результат
            overall_status = "pass"
            if total_critical_findings > 0:
                overall_status = "fail"
            elif total_warning_findings > 0:
                overall_status = "uncertain"
            
            # Вычисляем общий процент соответствия
            total_pages = len(pages)
            successful_pages = len([r for r in page_results if r.get("overall_status") == "pass"])
            compliance_percentage = (successful_pages / total_pages * 100) if total_pages > 0 else 0
            
            # Формируем общий отчет
            combined_result = {
                "overall_status": overall_status,
                "confidence": 0.8,  # Высокая уверенность при постраничной проверке
                "total_findings": total_findings,
                "critical_findings": total_critical_findings,
                "warning_findings": total_warning_findings,
                "info_findings": total_info_findings,
                "total_pages": total_pages,
                "successful_pages": successful_pages,
                "page_results": page_results,
                "findings": all_findings,
                "summary": f"Проверка завершена. Обработано {total_pages} страниц. Найдено {total_findings} нарушений.",
                "compliance_percentage": compliance_percentage,
                "recommendations": f"Общие рекомендации: {total_critical_findings} критических нарушений, {total_warning_findings} предупреждений, {total_info_findings} замечаний."
            }
            
            # Сохраняем результат в базу данных
            # ===== СОХРАНЕНИЕ РЕЗУЛЬТАТОВ ПРОВЕРКИ LLM =====
            # Сохраняем общий результат проверки в базу данных
            await self.save_norm_control_result(document_id, combined_result)
            
            return {
                "status": "success",
                "result": combined_result,
                "pages_processed": len(pages)
            }
                    
        except Exception as e:
            logger.error(f"Norm control check error: {e}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return {
                "status": "error",
                "error": str(e)
            }

    async def save_norm_control_result(self, document_id: int, check_result: Dict[str, Any]):
        """Сохранение результата проверки нормоконтроля от LLM в базу данных"""
        # Шаг 6: Сохранение результатов и генерация отчета
        total_findings = check_result.get("total_findings", 0)
        trace_logger.log_report_generation(document_id, total_findings)
        trace_logger.log_step("REPORT_GENERATION", f"Создание отчета с {total_findings} находками", document_id)
        def _save_result(conn):
            with conn.cursor() as cursor:
                cursor.execute("""
                    INSERT INTO norm_control_results 
                    (checkable_document_id, analysis_date, analysis_type, model_used, analysis_status,
                     total_findings, critical_findings, warning_findings, info_findings)
                    VALUES (%s, CURRENT_TIMESTAMP, %s, %s, %s, %s, %s, %s, %s)
                    RETURNING id
                """, (
                    document_id,
                    "norm_control",
                    "llama3.1:8b",
                    check_result.get("overall_status", "uncertain"),
                    check_result.get("total_findings", 0),
                    check_result.get("critical_findings", 0),
                    check_result.get("warning_findings", 0),
                    check_result.get("info_findings", 0)
                ))
                
                result_id = cursor.fetchone()[0]
                logger.info(f"Saved norm control result {result_id} for document {document_id}")
                return result_id
        
        try:
            result_id = self.execute_in_transaction(_save_result)
            
            # ===== СОХРАНЕНИЕ ДЕТАЛЬНЫХ FINDINGS =====
            # Сохраняем детальную информацию о каждом нарушении
            findings = check_result.get("findings", [])
            if findings:
                await self.save_findings_detailed(result_id, findings, document_id)
                logger.info(f"Saved {len(findings)} detailed findings for result {result_id}")
            
            # ===== СОЗДАНИЕ ОТЧЕТА О ПРОВЕРКЕ LLM =====
            # Создаем отчет о проверке на основе результатов LLM
            await self.create_review_report(document_id, result_id, check_result)
            
            return result_id
                
        except Exception as e:
            logger.error(f"Save norm control result error: {e}")
            raise

    async def save_findings_detailed(self, result_id: int, findings: List[Dict[str, Any]], document_id: int):
        """Сохранение детальной информации о каждом нарушении с привязкой к нормативным документам"""
        import json
        
        def _save_findings(conn):
            with conn.cursor() as cursor:
                for finding in findings:
                    # Определяем тип нарушения
                    finding_type = finding.get('type', 'violation')
                    if finding_type == 'critical':
                        finding_type = 'violation'
                        severity_level = 5
                    elif finding_type == 'warning':
                        finding_type = 'warning'
                        severity_level = 3
                    elif finding_type == 'info':
                        finding_type = 'info'
                        severity_level = 1
                    else:
                        severity_level = 2
                    
                    # Определяем категорию на основе кода
                    code = finding.get('code', '')
                    category = self.determine_finding_category(code, finding.get('description', ''))
                    
                    # Ищем связанный нормативный документ (clause_id)
                    clause_id = self.find_related_clause_id(finding, cursor)
                    
                    # Формируем ссылку на место в документе
                    element_reference = {
                        "page_number": finding.get('page_number', 1),
                        "finding_code": code,
                        "location": finding.get('location', 'Не указано'),
                        "element_type": finding.get('element_type', 'text'),
                        "bounding_box": finding.get('bounding_box', None)
                    }
                    
                    cursor.execute("""
                        INSERT INTO findings 
                        (norm_control_result_id, finding_type, severity_level, category,
                         title, description, recommendation, related_clause_id,
                         related_clause_text, element_reference, rule_applied, confidence_score)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    """, (
                        result_id,
                        finding_type,
                        severity_level,
                        category,
                        finding.get('title', finding.get('description', 'Нарушение требований'))[:200],  # Ограничиваем длину
                        finding.get('description', ''),
                        finding.get('recommendation', ''),
                        clause_id,
                        finding.get('clause_text', ''),
                        json.dumps(element_reference),
                        code,  # Используем код как правило
                        finding.get('confidence_score', 1.0)
                    ))
                    
                    logger.debug(f"Saved finding: {code} - {finding.get('description', '')[:50]}...")
        
        try:
            return self.execute_in_transaction(_save_findings)
        except Exception as e:
            logger.error(f"Save findings detailed error: {e}")
            raise

    def determine_finding_category(self, code: str, description: str) -> str:
        """Определение категории нарушения на основе кода и описания"""
        code_upper = code.upper()
        description_lower = description.lower()
        
        # Категории на основе кодов
        if 'NC-' in code_upper:
            return 'compliance'
        elif 'FIRE' in code_upper or 'пожар' in description_lower:
            return 'safety'
        elif 'ENERGY' in code_upper or 'энерг' in description_lower:
            return 'energy_efficiency'
        elif 'STRUCTURE' in code_upper or 'конструк' in description_lower:
            return 'structural'
        elif 'FORMAT' in code_upper or 'формат' in description_lower:
            return 'formatting'
        elif 'TECH' in code_upper or 'технич' in description_lower:
            return 'technical'
        else:
            return 'compliance'

    def find_related_clause_id(self, finding: Dict[str, Any], cursor) -> Optional[int]:
        """Поиск связанного нормативного документа по коду и описанию"""
        try:
            code = finding.get('code', '')
            description = finding.get('description', '')
            
            # Ищем по коду
            if code:
                cursor.execute("""
                    SELECT id FROM document_clauses 
                    WHERE clause_id ILIKE %s OR clause_number ILIKE %s
                    LIMIT 1
                """, (f'%{code}%', f'%{code}%'))
                result = cursor.fetchone()
                if result:
                    return result[0]
            
            # Ищем по ключевым словам в описании
            if description:
                # Извлекаем ключевые слова
                keywords = self.extract_keywords(description)
                if keywords:
                    # Ищем в названиях и текстах нормативных документов
                    for keyword in keywords[:3]:  # Берем первые 3 ключевых слова
                        cursor.execute("""
                            SELECT id FROM document_clauses 
                            WHERE clause_title ILIKE %s OR clause_text ILIKE %s
                            LIMIT 1
                        """, (f'%{keyword}%', f'%{keyword}%'))
                        result = cursor.fetchone()
                        if result:
                            return result[0]
            
            return None
            
        except Exception as e:
            logger.error(f"Error finding related clause: {e}")
            return None

    def extract_keywords(self, text: str) -> List[str]:
        """Извлечение ключевых слов из текста"""
        import re
        
        # Убираем стоп-слова и извлекаем существительные
        stop_words = {'требования', 'нарушение', 'соответствие', 'документ', 'проект', 'строительство'}
        
        # Извлекаем слова на русском языке
        words = re.findall(r'\b[а-яё]{4,}\b', text.lower())
        
        # Фильтруем стоп-слова и возвращаем уникальные
        keywords = [word for word in words if word not in stop_words]
        return list(set(keywords))[:5]  # Возвращаем до 5 ключевых слов

    async def create_review_report(self, document_id: int, result_id: int, check_result: Dict[str, Any]):
        """Создание отчета о проверке на основе результатов LLM"""
        def _create_report(conn):
            with conn.cursor() as cursor:
                cursor.execute("""
                    INSERT INTO review_reports 
                    (checkable_document_id, norm_control_result_id, report_date, review_type,
                     overall_status, reviewer_name, conclusion)
                    VALUES (%s, %s, CURRENT_TIMESTAMP, %s, %s, %s, %s)
                    RETURNING id
                """, (
                    document_id,
                    result_id,
                    "automatic",
                    check_result.get("overall_status", "uncertain"),
                    "AI System",
                    f"Автоматическая проверка: {check_result.get('summary', '')}"
                ))
                
                report_id = cursor.fetchone()[0]
                logger.info(f"Created review report {report_id} for document {document_id}")
                return report_id
        
        try:
            return self.execute_in_transaction(_create_report)
        except Exception as e:
            logger.error(f"Create review report error: {e}")
            raise

# Модели Pydantic
class StatusUpdateRequest(BaseModel):
    status: str
    reviewer: str = None
    notes: str = None

# Глобальный экземпляр парсера
parser = DocumentParser()

@app.post("/upload")
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    category: str = Form("other")
):
    """Загрузка и парсинг документа"""
    print(f"🔍 [DEBUG] DocumentParser: Upload started for file: {file.filename}")
    print(f"🔍 [DEBUG] DocumentParser: File size: {file.size} bytes")
    print(f"🔍 [DEBUG] DocumentParser: Content type: {file.content_type}")
    
    # Проверка размера файла из конфигурации
    if file.size and file.size > MAX_NORMATIVE_DOCUMENT_SIZE:
        print(f"🔍 [DEBUG] DocumentParser: File too large: {file.size} bytes (max: {MAX_NORMATIVE_DOCUMENT_SIZE})")
        raise HTTPException(
            status_code=413, 
            detail=f"File too large. Maximum size is {MAX_NORMATIVE_DOCUMENT_SIZE // (1024*1024)}MB. Your file is {file.size // (1024*1024)}MB"
        )
    
    try:
        # Чтение файла
        print(f"🔍 [DEBUG] DocumentParser: Reading file content...")
        file_content = await file.read()
        file_size = len(file_content)
        print(f"🔍 [DEBUG] DocumentParser: File content read, size: {file_size} bytes")
        
        # Определение типа файла
        print(f"🔍 [DEBUG] DocumentParser: Detecting file type...")
        file_type = parser.detect_file_type(file_content)
        print(f"🔍 [DEBUG] DocumentParser: Detected file type: {file_type}")
        
        # Вычисление хеша
        print(f"🔍 [DEBUG] DocumentParser: Calculating file hash...")
        document_hash = parser.calculate_file_hash(file_content)
        print(f"🔍 [DEBUG] DocumentParser: File hash: {document_hash}")
        
        # Проверка на дубликат
        print(f"🔍 [DEBUG] DocumentParser: Checking for duplicates...")
        with parser.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT id FROM uploaded_documents WHERE document_hash = %s", (document_hash,))
            existing_doc = cursor.fetchone()
            if existing_doc:
                print(f"🔍 [DEBUG] DocumentParser: Document already exists with ID: {existing_doc['id']}")
                raise HTTPException(status_code=400, detail="Document already exists")
        print(f"🔍 [DEBUG] DocumentParser: No duplicates found")
        
        # Сохраняем файл на диск
        print(f"🔍 [DEBUG] DocumentParser: Saving file to disk...")
        file_path = f"/app/uploads/{document_hash}_{file.filename}"
        with open(file_path, "wb") as f:
            f.write(file_content)
        print(f"🔍 [DEBUG] DocumentParser: File saved to: {file_path}")
        
        # Создаем запись в базе данных с статусом "processing"
        print(f"🔍 [DEBUG] DocumentParser: Creating initial database record...")
        document_id = parser.create_initial_document_record(
            file.filename,
            file_type,
            file_size,
            document_hash,
            file_path,
            category
        )
        print(f"🔍 [DEBUG] DocumentParser: Created document record with ID: {document_id}")
        
        # Создаем запись о загрузке без начала обработки
        print(f"🔍 [DEBUG] DocumentParser: Document uploaded successfully, processing will start separately")
        
        result = {
            "document_id": document_id,
            "filename": file.filename,
            "file_type": file_type,
            "file_size": file_size,
            "status": "uploaded",
            "message": "Document uploaded successfully. Ready for processing.",
            "upload_complete": True
        }
        print(f"🔍 [DEBUG] DocumentParser: Upload initiated successfully: {result}")
        return result
        
    except HTTPException as e:
        print(f"🔍 [DEBUG] DocumentParser: HTTPException in upload: {e.status_code} - {e.detail}")
        raise e
    except Exception as e:
        print(f"🔍 [DEBUG] DocumentParser: Unexpected error in upload: {type(e).__name__}: {str(e)}")
        import traceback
        print(f"🔍 [DEBUG] DocumentParser: Traceback: {traceback.format_exc()}")
        logger.error(f"Upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/documents/{document_id}/process")
async def start_document_processing(document_id: int):
    """Запуск обработки загруженного документа"""
    try:
        print(f"🔍 [DEBUG] DocumentParser: Starting processing for document {document_id}")
        
        # Получаем информацию о документе
        with parser.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT id, original_filename, file_type, file_path, processing_status
                FROM uploaded_documents
                WHERE id = %s
            """, (document_id,))
            document = cursor.fetchone()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        if document['processing_status'] != 'uploaded':
            raise HTTPException(status_code=400, detail="Document is not ready for processing")
        
        # Запускаем асинхронную обработку
        print(f"🔍 [DEBUG] DocumentParser: Adding background task for document {document_id}")
        background_tasks = BackgroundTasks()
        background_tasks.add_task(
            process_document_async,
            document_id=document_id,
            file_path=document['file_path'],
            file_type=document['file_type'],
            filename=document['original_filename']
        )
        
        # Обновляем статус на "processing"
        parser.update_document_status(document_id, "processing")
        
        return {
            "document_id": document_id,
            "status": "processing",
            "message": "Document processing started successfully."
        }
        
    except HTTPException as e:
        raise e
    except Exception as e:
        print(f"🔍 [DEBUG] DocumentParser: Error starting processing: {e}")
        logger.error(f"Start processing error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

async def process_document_async(document_id: int, file_path: str, file_type: str, filename: str):
    """Асинхронная обработка документа"""
    print(f"🔍 [DEBUG] DocumentParser: Starting async processing for document {document_id}")
    print(f"🔍 [DEBUG] DocumentParser: File type: {file_type}, filename: {filename}")
    print(f"🔍 [DEBUG] DocumentParser: File path: {file_path}")
    
    try:
        # Обновляем статус на "processing"
        print(f"🔍 [DEBUG] DocumentParser: Updating status to 'processing' for document {document_id}")
        parser.update_document_status(document_id, "processing")
        print(f"🔍 [DEBUG] DocumentParser: Status updated successfully")
        
        # Читаем файл с диска
        print(f"🔍 [DEBUG] DocumentParser: Reading file from disk: {file_path}")
        with open(file_path, "rb") as f:
            file_content = f.read()
        print(f"🔍 [DEBUG] DocumentParser: File content size: {len(file_content)} bytes")
        
        # Парсинг документа
        print(f"🔍 [DEBUG] DocumentParser: Parsing document {document_id}...")
        if file_type == "pdf":
            elements = parser.parse_pdf(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        print(f"🔍 [DEBUG] DocumentParser: Parsing completed for document {document_id}, elements count: {len(elements.document_pages_results)}")
        
        # Сохранение элементов в базу данных
        print(f"🔍 [DEBUG] DocumentParser: Saving elements to database for document {document_id}...")
        parser.save_elements_to_database(document_id, elements)
        
        # Подсчет токенов
        print(f"🔍 [DEBUG] DocumentParser: Calculating tokens for document {document_id}...")
        total_tokens = parser.calculate_document_tokens(elements)
        
        # Обновляем запись в базе данных
        print(f"🔍 [DEBUG] DocumentParser: Updating document record for {document_id}...")
        parser.update_document_completion(document_id, len(elements.document_pages_results), total_tokens)
        
        # Асинхронная индексация в RAG-сервис
        print(f"🔍 [DEBUG] DocumentParser: Starting RAG indexing for document {document_id}...")
        asyncio.create_task(
            parser.index_to_rag_service_async(
                document_id=document_id,
                document_title=filename,
                inspection_result=elements
            )
        )
        
        print(f"🔍 [DEBUG] DocumentParser: Async processing completed for document {document_id}")
        
    except Exception as e:
        print(f"🔍 [DEBUG] DocumentParser: Error in async processing for document {document_id}: {str(e)}")
        import traceback
        print(f"🔍 [DEBUG] DocumentParser: Async processing traceback: {traceback.format_exc()}")
        
        # Обновляем статус на "error"
        try:
            parser.update_document_status(document_id, "error")
        except Exception as update_error:
            print(f"🔍 [DEBUG] DocumentParser: Failed to update error status: {update_error}")
        
        logger.error(f"Async processing error for document {document_id}: {e}")

async def process_checkable_document_async(document_id: int, document_content: str, filename: str):
    """Асинхронная обработка проверяемого документа с постраничной проверкой нормоконтроля"""
    start_time = time.time()
    
    # Шаг 3: Начало асинхронной обработки
    trace_logger.log_step("ASYNC_PROCESSING_START", f"Документ {document_id}, контент: {len(document_content)} символов", document_id)
    
    logger.info(f"🔍 [DEBUG] DocumentParser: Starting async processing for checkable document {document_id}")
    logger.info(f"🔍 [DEBUG] DocumentParser: Async processing parameters:")
    logger.info(f"🔍 [DEBUG] DocumentParser: - Document ID: {document_id}")
    logger.info(f"🔍 [DEBUG] DocumentParser: - Document content size: {len(document_content)} characters")
    logger.info(f"🔍 [DEBUG] DocumentParser: - Filename: {filename}")
    log_memory_usage("start of async processing")
    
    try:
        # Обновляем статус на "processing"
        logger.info(f"🔍 [DEBUG] DocumentParser: Updating document {document_id} status to 'processing'")
        parser.update_checkable_document_status(document_id, "processing")
        logger.info(f"🔍 [DEBUG] DocumentParser: Updated document {document_id} status to 'processing'")
        
        # Разбиваем документ на страницы
        logger.info(f"🔍 [DEBUG] DocumentParser: Splitting document {document_id} into pages")
        pages = parser.split_document_into_pages(document_id)
        logger.info(f"🔍 [DEBUG] DocumentParser: Document {document_id} split into {len(pages)} pages")
        
        if not pages:
            logger.warning(f"🔍 [DEBUG] DocumentParser: No pages found for document {document_id}, using fallback")
            # Fallback: используем старый метод
            result = await parser.perform_norm_control_check(document_id, document_content)
        else:
            # Выполняем постраничную проверку нормоконтроля
            logger.info(f"🔍 [DEBUG] DocumentParser: Starting page-by-page norm control check for document {document_id}")
            all_page_results = []
            
            for i, page_data in enumerate(pages):
                page_number = page_data["page_number"]
                logger.info(f"🔍 [DEBUG] DocumentParser: Processing page {page_number}/{len(pages)} for document {document_id}")
                
                try:
                    # Выполняем проверку для одной страницы
                    page_result = await parser.perform_norm_control_check_for_page(document_id, page_data)
                    page_result["page_number"] = page_number
                    all_page_results.append(page_result)
                    
                    logger.info(f"🔍 [DEBUG] DocumentParser: Page {page_number} processed successfully")
                    
                except Exception as page_error:
                    logger.error(f"🔍 [DEBUG] DocumentParser: Error processing page {page_number}: {page_error}")
                    all_page_results.append({
                        "page_number": page_number,
                        "status": "error",
                        "error": str(page_error)
                    })
            
            # Объединяем результаты всех страниц
            logger.info(f"🔍 [DEBUG] DocumentParser: Combining results from {len(all_page_results)} pages")
            result = parser.combine_page_results(document_id, all_page_results)
        
        logger.info(f"🔍 [DEBUG] DocumentParser: Norm control check completed for document {document_id}")
        
        # Обновляем статус на "completed"
        parser.update_checkable_document_status(document_id, "completed")
        logger.info(f"🔍 [DEBUG] DocumentParser: Updated document {document_id} status to 'completed'")
        
        logger.info(f"🔍 [DEBUG] DocumentParser: Async processing completed successfully for document {document_id}")
        
    except Exception as e:
        logger.error(f"🔍 [DEBUG] DocumentParser: Error in async processing for checkable document {document_id}: {str(e)}")
        import traceback
        logger.error(f"🔍 [DEBUG] DocumentParser: Async processing traceback: {traceback.format_exc()}")
        
        # Обновляем статус на "error"
        try:
            parser.update_checkable_document_status(document_id, "error")
        except Exception as update_error:
            logger.error(f"🔍 [DEBUG] DocumentParser: Failed to update error status: {update_error}")

@app.get("/documents")
async def list_documents():
    """Список загруженных документов"""
    try:
        with parser.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT id, original_filename, file_type, file_size, upload_date, processing_status, token_count, category
                FROM uploaded_documents
                ORDER BY upload_date DESC
            """)
            documents = cursor.fetchall()
            
        return {"documents": [dict(doc) for doc in documents]}
        
    except Exception as e:
        logger.error(f"List documents error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/documents/{document_id}/elements")
async def get_document_elements(document_id: int):
    """Получение элементов документа"""
    try:
        with parser.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT id, element_type, element_content, page_number, confidence_score, created_at
                FROM extracted_elements
                WHERE uploaded_document_id = %s
                ORDER BY page_number, id
            """, (document_id,))
            elements = cursor.fetchall()
            
        return {"elements": [dict(elem) for elem in elements]}
        
    except Exception as e:
        logger.error(f"Get elements error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/documents/{document_id}/process")
async def process_document_manual(document_id: int, background_tasks: BackgroundTasks):
    """Ручной запуск обработки документа"""
    try:
        print(f"🔍 [DEBUG] DocumentParser: Manual processing requested for document {document_id}")
        
        # Получаем информацию о документе
        with parser.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT id, original_filename, file_type, file_size, document_hash, processing_status
                FROM uploaded_documents
                WHERE id = %s
            """, (document_id,))
            document = cursor.fetchone()
            
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
            
        if document["processing_status"] == "completed":
            return {"message": "Document already processed", "status": "completed"}
            
        print(f"🔍 [DEBUG] DocumentParser: Document info: {document}")
        
        # Получаем путь к файлу из базы данных
        with parser.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT file_path FROM uploaded_documents WHERE id = %s
            """, (document_id,))
            result = cursor.fetchone()
            file_path = result.get("file_path") if result else None
            
        if not file_path:
            raise HTTPException(status_code=404, detail="File not found on disk")
            
        # Запускаем асинхронную обработку
        background_tasks.add_task(
            process_document_async,
            document_id=document_id,
            file_path=file_path,
            file_type=document["file_type"],
            filename=document["original_filename"]
        )
        
        return {"message": "Processing started", "document_id": document_id}
        
    except Exception as e:
        logger.error(f"Manual processing error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/documents/{document_id}")
async def delete_normative_document(document_id: int):
    """Удаление нормативного документа"""
    try:
        success = await parser.delete_normative_document(document_id)
        
        if success:
            return {"status": "success", "message": f"Document {document_id} deleted successfully"}
        else:
            raise HTTPException(status_code=404, detail="Document not found")
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Delete normative document error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/checkable-documents/{document_id}")
async def delete_checkable_document(document_id: int):
    """Удаление проверяемого документа"""
    try:
        success = parser.delete_checkable_document(document_id)
        
        if success:
            return {"status": "success", "message": f"Checkable document {document_id} deleted successfully"}
        else:
            raise HTTPException(status_code=404, detail="Document not found")
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Delete checkable document error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload/chat")
async def upload_chat_file(
    file: UploadFile = File(...)
):
    """Загрузка файла для обработки в чате"""
    try:
        # Чтение файла
        file_content = await file.read()
        file_size = len(file_content)
        
        # Проверка размера файла (максимум 10MB для чата)
        if file_size > 10 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum size is 10MB")
        
        # Определение типа файла
        file_type = parser.detect_file_type(file_content)
        
        # Парсинг файла
        elements = parser.parse_file(file_content, file_type)
        
        # Извлечение текстового содержимого
        text_content = "\n\n".join([elem.get("element_content", "") for elem in elements if elem.get("element_content")])
        
        # Ограничиваем размер текста для чата (учитывая лимит Ollama ~4000 токенов)
        # Приблизительно 1 токен = 4 символа, но лучше использовать более консервативный подход
        # 4000 токенов ≈ 12000 символов (3 символа на токен)
        max_chars = 10000  # Консервативный лимит для Ollama
        if len(text_content) > max_chars:
            text_content = text_content[:max_chars] + "\n\n[Содержимое обрезано из-за ограничений размера. Максимум 10000 символов для чата]"
        
        return {
            "success": True,
            "filename": file.filename,
            "file_type": file_type,
            "file_size": file_size,
            "content": text_content,
            "elements_count": len(elements.document_pages_results)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Chat file upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload/checkable")
async def upload_checkable_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...)
):
    """Загрузка проверяемого документа с детальным трейслогом"""
    start_time = time.time()
    
    # Шаг 1: Начало загрузки документа
    trace_logger.log_upload_start(file.filename, file.size or 0)
    trace_logger.log_step("UPLOAD_START", f"Файл: {file.filename}, размер: {file.size} байт")
    
    logger.info(f"🔍 [DEBUG] DocumentParser: upload_checkable_document started for file: {file.filename}")
    logger.info(f"🔍 [DEBUG] DocumentParser: Content-Type: {file.content_type}")
    logger.info(f"🔍 [DEBUG] DocumentParser: File size from UploadFile: {file.size}")
    logger.info(f"🔍 [DEBUG] DocumentParser: Processing mode: ASYNC")
    logger.info(f"🔍 [DEBUG] DocumentParser: Memory limit: {MAX_CHECKABLE_DOCUMENT_SIZE // (1024*1024)}MB")
    
    # Проверка размера файла из конфигурации
    if file.size and file.size > MAX_CHECKABLE_DOCUMENT_SIZE:
        logger.warning(f"🔍 [DEBUG] DocumentParser: File too large: {file.size} bytes (max: {MAX_CHECKABLE_DOCUMENT_SIZE})")
        raise HTTPException(
            status_code=413, 
            detail=f"File too large. Maximum size is {MAX_CHECKABLE_DOCUMENT_SIZE // (1024*1024)}MB. Your file is {file.size // (1024*1024)}MB"
        )
    
    # Проверка доступной памяти перед обработкой
    log_memory_usage("before file processing")
    available_memory = get_available_memory()
    file_size_mb = file.size / (1024 * 1024) if file.size else 0
    
    if available_memory < file_size_mb * 3:  # Нужно минимум 3x памяти для обработки
        logger.warning(f"🔍 [DEBUG] DocumentParser: Insufficient memory for file {file_size_mb:.1f}MB. Available: {available_memory:.1f}MB")
        raise HTTPException(
            status_code=507,  # Insufficient Storage
            detail=f"Insufficient memory for processing. File size: {file_size_mb:.1f}MB, Available memory: {available_memory:.1f}MB"
        )
    
    try:
        # Чтение файла
        logger.info(f"🔍 [DEBUG] DocumentParser: Reading file content...")
        file_content = await file.read()
        file_size = len(file_content)
        logger.info(f"🔍 [DEBUG] DocumentParser: File content read successfully, size: {file_size} bytes")
        logger.info(f"🔍 [DEBUG] DocumentParser: Content preview (first 100 bytes): {file_content[:100]}")
        
        # Определение типа файла
        logger.info(f"🔍 [DEBUG] DocumentParser: Detecting file type...")
        file_type = parser.detect_file_type(file_content)
        logger.info(f"🔍 [DEBUG] DocumentParser: Detected file type: {file_type}")
        
        # Вычисление хеша
        logger.info(f"🔍 [DEBUG] DocumentParser: Calculating file hash...")
        document_hash = parser.calculate_file_hash(file_content)
        logger.info(f"🔍 [DEBUG] DocumentParser: Calculated hash: {document_hash}")
        
        # Проверка на дубликат
        logger.info(f"🔍 [DEBUG] DocumentParser: Checking for duplicate document...")
        with parser.db_conn.cursor() as cursor:
            cursor.execute("SELECT id FROM checkable_documents WHERE document_hash = %s", (document_hash,))
            existing_doc = cursor.fetchone()
            if existing_doc:
                logger.warning(f"🔍 [DEBUG] DocumentParser: Document already exists with hash: {document_hash}, existing ID: {existing_doc[0]}")
                raise HTTPException(status_code=400, detail="Document already exists")
            else:
                logger.info(f"🔍 [DEBUG] DocumentParser: No duplicate found, proceeding with upload")
        
        # Парсинг документа
        logger.info(f"🔍 [DEBUG] DocumentParser: Starting document parsing for type: {file_type}")
        
        if file_type == "pdf":
            logger.info(f"🔍 [DEBUG] DocumentParser: Parsing PDF document...")
            logger.info(f"🔍 [DEBUG] DocumentParser: PDF size: {file_size / (1024*1024):.1f}MB")
            
            # Для больших файлов используем увеличенный таймаут
            import asyncio
            pdf_timeout = 1200 if file_size > 10 * 1024 * 1024 else 600  # 20 мин для файлов > 10MB, 10 мин для остальных
            
            try:
                # Запускаем парсинг с таймаутом
                inspection_result = await asyncio.wait_for(
                    asyncio.to_thread(parser.parse_pdf, file_content),
                    timeout=pdf_timeout
                )
                logger.info(f"🔍 [DEBUG] DocumentParser: PDF parsing completed successfully")
                logger.info(f"🔍 [DEBUG] DocumentParser: Parsed {len(inspection_result.document_pages_results)} pages")
                logger.info(f"🔍 [DEBUG] DocumentParser: Document stats - Total: {inspection_result.document_pages}, Vector: {inspection_result.document_pages_vector}, Scanned: {inspection_result.document_pages_scanned}")
                
            except asyncio.TimeoutError:
                logger.error(f"🔍 [DEBUG] DocumentParser: PDF parsing timeout after {pdf_timeout} seconds")
                raise HTTPException(
                    status_code=408,  # Request Timeout
                    detail=f"PDF parsing timeout after {pdf_timeout} seconds. File is too large or complex."
                )
            except Exception as pdf_error:
                logger.error(f"🔍 [DEBUG] DocumentParser: PDF parsing error: {pdf_error}")
                raise HTTPException(
                    status_code=500,
                    detail=f"PDF parsing failed: {str(pdf_error)}"
                )

        elif file_type == "docx":
            logger.info(f"🔍 [DEBUG] DocumentParser: Parsing DOCX document...")
            # Временно используем старый метод для docx
            elements = parser.parse_docx(file_content)
            logger.info(f"🔍 [DEBUG] DocumentParser: DOCX parsing completed, found {len(elements)} elements")
            inspection_result = DocumentInspectionResult()
            inspection_result.document_pages_results = []
            for i, element in enumerate(elements):
                page_result = DocumentPageInspectionResult()
                page_result.page_number = element.get("page_number", 1)
                page_result.page_text = element.get("element_content", "")
                page_result.page_text_confidence = element.get("confidence_score", 1.0)
                page_result.page_text_method = "docx_parsing"
                page_result.page_text_length = len(element.get("element_content", ""))
                page_result.page_type = "vector"
                inspection_result.document_pages_results.append(page_result)
                logger.debug(f"🔍 [DEBUG] DocumentParser: Processed DOCX element {i+1}, page {page_result.page_number}, text length: {page_result.page_text_length}")
            logger.info(f"🔍 [DEBUG] DocumentParser: DOCX processing completed, created {len(inspection_result.document_pages_results)} page results")
        elif file_type == "dwg":
            logger.info(f"🔍 [DEBUG] DocumentParser: Parsing DWG document...")
            # Временно используем старый метод для dwg
            elements = parser.parse_dwg(file_content)
            logger.info(f"🔍 [DEBUG] DocumentParser: DWG parsing completed, found {len(elements)} elements")
            inspection_result = DocumentInspectionResult()
            inspection_result.document_pages_results = []
            for i, element in enumerate(elements):
                page_result = DocumentPageInspectionResult()
                page_result.page_number = element.get("page_number", 1)
                page_result.page_text = element.get("element_content", "")
                page_result.page_text_confidence = element.get("confidence_score", 1.0)
                page_result.page_text_method = "dwg_parsing"
                page_result.page_text_length = len(element.get("element_content", ""))
                page_result.page_type = "vector"
                inspection_result.document_pages_results.append(page_result)
                logger.debug(f"🔍 [DEBUG] DocumentParser: Processed DWG element {i+1}, page {page_result.page_number}, text length: {page_result.page_text_length}")
            logger.info(f"🔍 [DEBUG] DocumentParser: DWG processing completed, created {len(inspection_result.document_pages_results)} page results")
        elif file_type == "ifc":
            logger.info(f"🔍 [DEBUG] DocumentParser: Parsing IFC document...")
            # Временно используем старый метод для ifc
            elements = parser.parse_ifc(file_content)
            logger.info(f"🔍 [DEBUG] DocumentParser: IFC parsing completed, found {len(elements)} elements")
            inspection_result = DocumentInspectionResult()
            inspection_result.document_pages_results = []
            for i, element in enumerate(elements):
                page_result = DocumentPageInspectionResult()
                page_result.page_number = element.get("page_number", 1)
                page_result.page_text = element.get("element_content", "")
                page_result.page_text_confidence = element.get("confidence_score", 1.0)
                page_result.page_text_method = "ifc_parsing"
                page_result.page_text_length = len(element.get("element_content", ""))
                page_result.page_type = "vector"
                inspection_result.document_pages_results.append(page_result)
                logger.debug(f"🔍 [DEBUG] DocumentParser: Processed IFC element {i+1}, page {page_result.page_number}, text length: {page_result.page_text_length}")
            logger.info(f"🔍 [DEBUG] DocumentParser: IFC processing completed, created {len(inspection_result.document_pages_results)} page results")
        elif file_type == "txt":
            logger.info(f"🔍 [DEBUG] DocumentParser: Parsing TXT document...")
            # Временно используем старый метод для txt
            elements = parser.parse_text(file_content)
            logger.info(f"🔍 [DEBUG] DocumentParser: TXT parsing completed, found {len(elements)} elements")
            inspection_result = DocumentInspectionResult()
            inspection_result.document_pages_results = []
            for i, element in enumerate(elements):
                page_result = DocumentPageInspectionResult()
                page_result.page_number = element.get("page_number", 1)
                page_result.page_text = element.get("element_content", "")
                page_result.page_text_confidence = element.get("confidence_score", 1.0)
                page_result.page_text_method = "txt_parsing"
                page_result.page_text_length = len(element.get("element_content", ""))
                page_result.page_type = "vector"
                inspection_result.document_pages_results.append(page_result)
                logger.debug(f"🔍 [DEBUG] DocumentParser: Processed TXT element {i+1}, page {page_result.page_number}, text length: {page_result.page_text_length}")
            logger.info(f"🔍 [DEBUG] DocumentParser: TXT processing completed, created {len(inspection_result.document_pages_results)} page results")
        else:
            logger.error(f"🔍 [DEBUG] DocumentParser: Unsupported file type: {file_type}")
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")
        
        # Определение категории документа
        logger.info(f"🔍 [DEBUG] DocumentParser: Determining document category...")
        content_text = "\n".join([page_result.page_text for page_result in inspection_result.document_pages_results])
        logger.debug(f"🔍 [DEBUG] DocumentParser: Content text length: {len(content_text)} characters")
        category = parser.determine_document_category(file.filename, content_text)
        logger.info(f"🔍 [DEBUG] DocumentParser: Determined category: {category}")
        
        # Сохранение как проверяемый документ
        logger.info(f"🔍 [DEBUG] DocumentParser: Saving checkable document to database...")
        document_id = parser.save_checkable_document(
            file.filename,
            file.filename,
            file_type,
            file_size,
            document_hash,
            inspection_result,
            category
        )
        logger.info(f"🔍 [DEBUG] DocumentParser: Document saved successfully with ID: {document_id}")
        
        # Запуск асинхронной проверки нормоконтроля
        logger.info(f"🔍 [DEBUG] DocumentParser: Starting async norm control check...")
        document_content = "\n\n".join([page_result.page_text for page_result in inspection_result.document_pages_results])
        background_tasks.add_task(
            process_checkable_document_async,
            document_id=document_id,
            document_content=document_content,
            filename=file.filename
        )
        
        # Подготовка ответа
        response_data = {
            "document_id": document_id,
            "filename": file.filename,
            "file_type": file_type,
            "file_size": file_size,
            "pages_count": len(inspection_result.document_pages_results),
            "category": category,
            "status": "processing",  # Изменено с "completed" на "processing"
            "review_deadline": (datetime.now() + timedelta(days=2)).isoformat(),
            "document_stats": {
                "total_pages": inspection_result.document_pages,
                "vector_pages": inspection_result.document_pages_vector,
                "scanned_pages": inspection_result.document_pages_scanned,
                "unknown_pages": inspection_result.document_pages_unknown,
                "a4_sheets_equivalent": inspection_result.document_pages_total_a4_sheets_equivalent
            },
            "message": "Document uploaded successfully. Norm control check started in background."
        }
        
        # Шаг 2: Успешная загрузка документа
        processing_time = time.time() - start_time
        trace_logger.log_upload_success(file.filename, document_id, processing_time)
        trace_logger.log_step("UPLOAD_SUCCESS", f"Документ загружен за {processing_time:.2f}с")
        
        logger.info(f"🔍 [DEBUG] DocumentParser: Upload completed successfully")
        logger.info(f"🔍 [DEBUG] DocumentParser: Response data: {response_data}")
        
        return response_data
        
    except HTTPException as http_ex:
        logger.warning(f"🔍 [DEBUG] DocumentParser: HTTPException raised: {http_ex.status_code} - {http_ex.detail}")
        raise
    except Exception as e:
        logger.error(f"🔍 [DEBUG] DocumentParser: Upload checkable document error: {e}")
        import traceback
        logger.error(f"🔍 [DEBUG] DocumentParser: Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/checkable-documents")
async def list_checkable_documents():
    """Список проверяемых документов"""
    logger.info(f"🔍 [API] List checkable documents called")
    try:
        # Проверяем состояние shutdown
        if is_shutting_down:
            logger.warning("🔍 [API] Service is shutting down, rejecting checkable documents request")
            raise HTTPException(status_code=503, detail="Service is shutting down")
        
        # Проверяем давление на память
        if check_memory_pressure():
            logger.warning("🔍 [API] High memory pressure detected during checkable documents request")
            cleanup_memory()
        
        logger.info(f"🔍 [API] Request list of checked documents via document parser.")
        documents = parser.get_checkable_documents()
        logger.info(f"🔍 [API] Successfully retrieved {len(documents)} checkable documents")
        return {"documents": documents}
        
    except HTTPException as e:
        # Перебрасываем HTTP исключения как есть
        logger.info(f"🔍 [API] HTTP excaption: {e.status_code} - {e.detail}")
        raise
    except Exception as e:
        logger.error(f"🔍 [API] List checkable documents error: {e}")
        # Возвращаем пустой список вместо ошибки 500
        return {"documents": [], "warning": "Service temporarily unavailable"}

@app.get("/checkable-documents/{document_id}/elements")
async def get_checkable_document_elements(document_id: int):
    """Получение элементов проверяемого документа"""
    try:
        logger.info(f"🔍 [API] Get checkable document elements called")
        with parser.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT id, element_type, element_content, page_number, confidence_score, created_at
                FROM checkable_elements
                WHERE checkable_document_id = %s
                ORDER BY page_number, id
            """, (document_id,))
            elements = cursor.fetchall()
            
        return {"elements": [dict(elem) for elem in elements]}
        
    except Exception as e:
        logger.error(f"Get checkable elements error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/checkable-documents/{document_id}/status")
async def update_checkable_document_status(
    document_id: int,
    request: StatusUpdateRequest
):
    """Обновление статуса проверяемого документа"""
    try:
        logger.info(f"🔍 [API] Update checkable document {document_id} status {request.status}")

        success = parser.update_review_status(document_id, request.status, request.reviewer, request.notes)
        if success:
            return {"status": "success", "message": f"Document {document_id} status updated"}
        else:
            raise HTTPException(status_code=404, detail="Document not found")
            
    except Exception as e:
        logger.error(f"Update status error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/cleanup-expired")
async def cleanup_expired_documents():
    """Очистка просроченных документов"""
    try:
        deleted_count = parser.cleanup_expired_documents()
        return {
            "status": "success",
            "deleted_count": deleted_count,
            "message": f"Deleted {deleted_count} expired documents"
        }
        
    except Exception as e:
        logger.error(f"Cleanup error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/reindex-documents")
async def reindex_documents():
    """Принудительная реиндексация всех документов"""
    try:
        # Получаем все документы
        with parser.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT id, original_filename, token_count
                FROM uploaded_documents
                WHERE processing_status = 'completed'
                ORDER BY upload_date DESC
            """)
            documents = cursor.fetchall()
        
        total_documents = len(documents)
        total_tokens = sum(doc['token_count'] or 0 for doc in documents)
        
        # Получаем элементы для каждого документа и пересчитываем токены
        updated_count = 0
        new_total_tokens = 0
        
        for doc in documents:
            try:
                # Получаем элементы документа в отдельном блоке
                with parser.db_conn.cursor(cursor_factory=RealDictCursor) as element_cursor:
                    element_cursor.execute("""
                        SELECT element_content
                        FROM extracted_elements
                        WHERE uploaded_document_id = %s
                    """, (doc['id'],))
                    elements = element_cursor.fetchall()
                
                # Подсчитываем токены
                total_doc_tokens = 0
                for element in elements:
                    if element['element_content']:
                        total_doc_tokens += parser.count_tokens(element['element_content'])
                
                # Обновляем количество токенов
                with parser.db_conn.cursor() as update_cursor:
                    update_cursor.execute("""
                        UPDATE uploaded_documents
                        SET token_count = %s
                        WHERE id = %s
                    """, (total_doc_tokens, doc['id']))
                
                # Индексируем документ в RAG сервис
                try:
                    # Получаем элементы для индексации
                    with parser.db_conn.cursor(cursor_factory=RealDictCursor) as rag_cursor:
                        rag_cursor.execute("""
                            SELECT element_type, element_content, page_number
                            FROM extracted_elements
                            WHERE uploaded_document_id = %s
                            ORDER BY page_number, id
                        """, (doc['id'],))
                        rag_elements = rag_cursor.fetchall()
                    
                    # Индексируем в RAG сервис
                    await parser.index_to_rag_service(
                        document_id=doc['id'],
                        document_title=doc['original_filename'],
                        elements=[dict(elem) for elem in rag_elements]
                    )
                    logger.info(f"Successfully indexed document {doc['id']} in RAG service")
                except Exception as rag_error:
                    logger.error(f"Failed to index document {doc['id']} in RAG service: {rag_error}")
                
                new_total_tokens += total_doc_tokens
                updated_count += 1
                
            except Exception as e:
                logger.error(f"Error updating tokens for document {doc['id']}: {e}")
                parser.db_conn.rollback()
                continue
        
        parser.db_conn.commit()
        
        return {
            "status": "success",
            "total_documents": total_documents,
            "updated_documents": updated_count,
            "old_total_tokens": total_tokens,
            "new_total_tokens": new_total_tokens,
            "message": f"Reindexed {updated_count} documents with {new_total_tokens} total tokens"
        }
        
    except Exception as e:
        logger.error(f"Reindex error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/documents/{document_id}/tokens")
async def get_document_tokens(document_id: int):
    """Получение информации о токенах документа"""
    try:
        with parser.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT id, original_filename, token_count, file_size, upload_date
                FROM uploaded_documents
                WHERE id = %s
            """, (document_id,))
            document = cursor.fetchone()
            
            if not document:
                raise HTTPException(status_code=404, detail="Document not found")
            
            # Получаем детальную информацию о токенах по элементам
            cursor.execute("""
                SELECT element_type, element_content, page_number
                FROM extracted_elements
                WHERE uploaded_document_id = %s
                ORDER BY page_number, id
            """, (document_id,))
            elements = cursor.fetchall()
            
            # Подсчитываем токены по типам элементов
            token_stats = {
                "total_tokens": document['token_count'] or 0,
                "elements_count": len(elements),
                "by_type": {},
                "by_page": {}
            }
            
            for element in elements:
                element_type = element['element_type']
                page_number = element['page_number'] or 1
                content = element['element_content'] or ""
                
                if content:
                    tokens = parser.count_tokens(content)
                    
                    # По типам
                    if element_type not in token_stats["by_type"]:
                        token_stats["by_type"][element_type] = {"count": 0, "tokens": 0}
                    token_stats["by_type"][element_type]["count"] += 1
                    token_stats["by_type"][element_type]["tokens"] += tokens
                    
                    # По страницам
                    if page_number not in token_stats["by_page"]:
                        token_stats["by_page"][page_number] = {"count": 0, "tokens": 0}
                    token_stats["by_page"][page_number]["count"] += 1
                    token_stats["by_page"][page_number]["tokens"] += tokens
            
            return {
                "document": dict(document),
                "token_statistics": token_stats
            }
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get document tokens error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Модели Pydantic для настроек
class SettingUpdateRequest(BaseModel):
    setting_value: str

class SettingCreateRequest(BaseModel):
    setting_key: str
    setting_value: str
    setting_description: str
    setting_type: str = "text"

@app.get("/settings/prompt-templates")
async def get_prompt_templates():
    """Получение доступных шаблонов промптов"""
    try:
        logger.info("Getting prompt templates...")
        
        # Получаем основной промпт
        normcontrol_prompt = parser.get_system_setting("normcontrol_prompt")
        logger.info(f"Normcontrol prompt: {normcontrol_prompt[:100] if normcontrol_prompt else 'None'}...")
        
        # Получаем шаблон промпта (если есть)
        prompt_template = parser.get_system_setting("normcontrol_prompt_template")
        logger.info(f"Prompt template: {prompt_template[:100] if prompt_template else 'None'}...")
        
        templates = {
            "normcontrol_prompt": normcontrol_prompt,
            "normcontrol_prompt_template": prompt_template,
            "has_custom_template": prompt_template is not None
        }
        
        logger.info("Successfully retrieved prompt templates")
        return {"status": "success", "templates": templates}
    except Exception as e:
        logger.error(f"Get prompt templates error: {e}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/settings/prompt-templates")
async def update_prompt_template(request: Dict[str, Any]):
    """Обновление шаблона промпта"""
    try:
        template_key = request.get("template_key")
        template_value = request.get("template_value")
        template_description = request.get("template_description", "")
        
        if not template_key or not template_value:
            raise HTTPException(status_code=400, detail="Missing template_key or template_value")
        
        success = parser.create_system_setting(
            template_key, 
            template_value, 
            template_description, 
            "prompt_template"
        )
        
        if success:
            return {"status": "success", "message": f"Template {template_key} updated successfully"}
        else:
            raise HTTPException(status_code=500, detail="Failed to update template")
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Update prompt template error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/settings")
async def get_settings():
    """Получение всех настроек системы"""
    try:
        settings = parser.get_system_settings()
        return {"settings": settings}
    except Exception as e:
        logger.error(f"Get settings error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/settings/{setting_key}")
async def get_setting(setting_key: str):
    """Получение конкретной настройки"""
    try:
        setting_value = parser.get_system_setting(setting_key)
        if setting_value is None:
            raise HTTPException(status_code=404, detail="Setting not found")
        return {"setting_key": setting_key, "setting_value": setting_value}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get setting error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/settings/{setting_key}")
async def update_setting(setting_key: str, request: SettingUpdateRequest):
    """Обновление настройки"""
    try:
        success = parser.update_system_setting(setting_key, request.setting_value)
        if success:
            return {"status": "success", "message": f"Setting {setting_key} updated successfully"}
        else:
            raise HTTPException(status_code=404, detail="Setting not found")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Update setting error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/settings")
async def create_setting(request: SettingCreateRequest):
    """Создание новой настройки"""
    try:
        success = parser.create_system_setting(
            request.setting_key,
            request.setting_value,
            request.setting_description,
            request.setting_type
        )
        if success:
            return {"status": "success", "message": f"Setting {request.setting_key} created successfully"}
        else:
            raise HTTPException(status_code=500, detail="Failed to create setting")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Create setting error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/settings/{setting_key}")
async def delete_setting(setting_key: str):
    """Удаление настройки"""
    try:
        success = parser.delete_system_setting(setting_key)
        if success:
            return {"status": "success", "message": f"Setting {setting_key} deleted successfully"}
        else:
            raise HTTPException(status_code=404, detail="Setting not found or cannot be deleted")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Delete setting error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/checkable-documents/{document_id}/report")
async def get_checkable_document_report(document_id: int):
    """Получение отчета о проверке документа"""
    try:
        with parser.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            # Получаем информацию о документе
            cursor.execute("""
                SELECT id, original_filename, file_type, upload_date, review_deadline, review_status
                FROM checkable_documents
                WHERE id = %s
            """, (document_id,))
            document = cursor.fetchone()
            
            if not document:
                raise HTTPException(status_code=404, detail="Document not found")
            
            # Получаем результаты нормоконтроля
            cursor.execute("""
                SELECT id, analysis_date, analysis_status,
                       total_findings, critical_findings, warning_findings, info_findings
                FROM norm_control_results
                WHERE checkable_document_id = %s
                ORDER BY analysis_date DESC
                LIMIT 1
            """, (document_id,))
            norm_result = cursor.fetchone()
            
            # Получаем детальные findings
            findings = parser.get_findings_by_document_id(document_id)
            
            # Получаем отчеты о проверке
            cursor.execute("""
                SELECT id, report_date, overall_status, reviewer_name, conclusion
                FROM review_reports
                WHERE checkable_document_id = %s
                ORDER BY report_date DESC
            """, (document_id,))
            reports = cursor.fetchall()
            
            return {
                "document": dict(document),
                "norm_control_result": dict(norm_result) if norm_result else None,
                "findings": findings,
                "review_reports": [dict(report) for report in reports]
            }
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get document report error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/checkable-documents/{document_id}/findings")
async def get_document_findings(document_id: int):
    """Получение детальных findings для документа"""
    try:
        # Проверяем существование документа
        document = parser.get_checkable_document(document_id)
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Получаем детальные findings
        findings = parser.get_findings_by_document_id(document_id)
        
        return {
            "document_id": document_id,
            "findings": findings,
            "total_findings": len(findings),
            "findings_by_category": {
                category: len([f for f in findings if f.get('category') == category])
                for category in set(f.get('category', 'compliance') for f in findings)
            },
            "findings_by_severity": {
                severity: len([f for f in findings if f.get('severity_level') == severity])
                for severity in set(f.get('severity_level', 1) for f in findings)
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get document findings error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/checkable-documents/{document_id}/check")
async def trigger_norm_control_check(document_id: int):
    """Принудительный запуск проверки нормоконтроля"""
    try:
        # Получаем содержимое документа
        with parser.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT element_content
                FROM checkable_elements
                WHERE checkable_document_id = %s
                ORDER BY page_number, id
            """, (document_id,))
            elements = cursor.fetchall()
        
        if not elements:
            raise HTTPException(status_code=404, detail="Document content not found")
        
        # Объединяем содержимое
        document_content = "\n\n".join([elem["element_content"] for elem in elements])
        
        # ===== ЗАПУСК ПРОВЕРКИ НОРМОКОНТРОЛЯ С ПРИМЕНЕНИЕМ LLM =====
        # Выполняем проверку документа с использованием LLM
        result = await parser.perform_norm_control_check(document_id, document_content)
        
        return {
            "status": "success",
            "message": "Norm control check completed",
            "result": result
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Trigger norm control check error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/checkable-documents/{document_id}/format-statistics")
async def get_document_format_statistics(document_id: int):
    """Получение статистики форматов документа"""
    try:
        stats = parser.get_document_format_statistics(document_id)
        return {"status": "success", "statistics": stats}
    except Exception as e:
        logger.error(f"Get document format statistics error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/documents/stats")
async def get_documents_stats():
    """Получение статистики нормативных документов"""
    try:
        with parser.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            # Общая статистика по документам
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_documents,
                    COUNT(CASE WHEN processing_status = 'completed' THEN 1 END) as completed_documents,
                    COUNT(CASE WHEN processing_status = 'pending' THEN 1 END) as pending_documents,
                    COUNT(CASE WHEN processing_status = 'error' THEN 1 END) as error_documents,
                    COUNT(DISTINCT category) as unique_categories,
                    SUM(token_count) as total_tokens
                FROM uploaded_documents
            """)
            doc_stats = cursor.fetchone()
            
            # Статистика по категориям
            cursor.execute("""
                SELECT 
                    category,
                    COUNT(*) as count,
                    SUM(token_count) as total_tokens
                FROM uploaded_documents
                WHERE category IS NOT NULL AND category != ''
                GROUP BY category
                ORDER BY count DESC
            """)
            categories = cursor.fetchall()
            
            # Статистика по извлеченным элементам
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_elements,
                    COUNT(DISTINCT uploaded_document_id) as documents_with_elements
                FROM extracted_elements
            """)
            elements_stats = cursor.fetchone()
            
            # Вычисляем прогресс индексации
            total_docs = doc_stats["total_documents"] or 0
            completed_docs = doc_stats["completed_documents"] or 0
            indexing_progress = (completed_docs / total_docs * 100) if total_docs > 0 else 0
            
        return {
            "status": "success",
            "statistics": {
                "total_documents": total_docs,
                "indexed_documents": completed_docs,
                "indexing_progress_percent": round(indexing_progress, 1),
                "categories_count": doc_stats["unique_categories"] or 0,
                "total_tokens": doc_stats["total_tokens"] or 0,
                "total_elements": elements_stats["total_elements"] or 0,
                "documents_with_elements": elements_stats["documents_with_elements"] or 0,
                "categories": [dict(cat) for cat in categories]
            }
        }
        
    except Exception as e:
        logger.error(f"Get documents stats error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/metrics")
async def get_metrics():
    """Получение метрик сервиса в формате Prometheus"""
    try:
        db_conn = parser.get_db_connection()
        with db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            # Статистика по загруженным документам
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_documents,
                    COUNT(CASE WHEN processing_status = 'completed' THEN 1 END) as completed_documents,
                    COUNT(CASE WHEN processing_status = 'pending' THEN 1 END) as pending_documents,
                    COUNT(CASE WHEN processing_status = 'error' THEN 1 END) as error_documents,
                    COUNT(CASE WHEN file_type = 'pdf' THEN 1 END) as pdf_documents,
                    COUNT(CASE WHEN file_type = 'docx' THEN 1 END) as docx_documents,
                    COUNT(CASE WHEN file_type = 'dwg' THEN 1 END) as dwg_documents,
                    COUNT(CASE WHEN file_type = 'txt' THEN 1 END) as txt_documents,
                    SUM(file_size) as total_size_bytes,
                    AVG(file_size) as avg_file_size_bytes,
                    SUM(token_count) as total_tokens
                FROM uploaded_documents
            """)
            doc_stats = cursor.fetchone()
            
            # Статистика по проверяемым документам
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_checkable_documents,
                    COUNT(CASE WHEN review_status = 'pending' THEN 1 END) as pending_reviews,
                    COUNT(CASE WHEN review_status = 'completed' THEN 1 END) as completed_reviews,
                    COUNT(CASE WHEN review_status = 'in_progress' THEN 1 END) as in_progress_reviews,
                    COUNT(CASE WHEN review_status = 'overdue' THEN 1 END) as overdue_reviews
                FROM checkable_documents
            """)
            checkable_stats = cursor.fetchone()
            
            # Статистика по извлеченным элементам
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_elements,
                    COUNT(CASE WHEN element_type = 'text' THEN 1 END) as text_elements,
                    COUNT(CASE WHEN element_type = 'table' THEN 1 END) as table_elements,
                    COUNT(CASE WHEN element_type = 'figure' THEN 1 END) as figure_elements,
                    COUNT(CASE WHEN element_type = 'stamp' THEN 1 END) as stamp_elements
                FROM extracted_elements
            """)
            elements_stats = cursor.fetchone()
            
            # Статистика по результатам нормоконтроля
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_norm_control_results,
                    COUNT(CASE WHEN analysis_status = 'completed' THEN 1 END) as completed_checks,
                    COUNT(CASE WHEN analysis_status = 'pending' THEN 1 END) as pending_checks,
                    COUNT(CASE WHEN analysis_status = 'error' THEN 1 END) as error_checks,
                    SUM(total_findings) as total_findings,
                    SUM(critical_findings) as critical_findings,
                    SUM(warning_findings) as warning_findings,
                    SUM(info_findings) as info_findings
                FROM norm_control_results
            """)
            norm_control_stats = cursor.fetchone()
            
            # Статистика по отчетам
            cursor.execute("""
                SELECT COUNT(*) as total_review_reports
                FROM review_reports
            """)
            reports_stats = cursor.fetchone()
            
            # Статистика по времени обработки (последние 24 часа)
            cursor.execute("""
                SELECT 
                    COUNT(*) as documents_last_24h
                FROM uploaded_documents 
                WHERE upload_date >= NOW() - INTERVAL '24 hours'
            """)
            time_stats = cursor.fetchone()
        
        # Формируем метрики в формате Prometheus
        metrics_lines = []
        
        # Метрики документов
        metrics_lines.append(f"# HELP document_parser_documents_total Total number of documents")
        metrics_lines.append(f"# TYPE document_parser_documents_total counter")
        metrics_lines.append(f"document_parser_documents_total {doc_stats['total_documents'] or 0}")
        
        metrics_lines.append(f"# HELP document_parser_documents_completed Total number of completed documents")
        metrics_lines.append(f"# TYPE document_parser_documents_completed counter")
        metrics_lines.append(f"document_parser_documents_completed {doc_stats['completed_documents'] or 0}")
        
        metrics_lines.append(f"# HELP document_parser_documents_pending Total number of pending documents")
        metrics_lines.append(f"# TYPE document_parser_documents_pending counter")
        metrics_lines.append(f"document_parser_documents_pending {doc_stats['pending_documents'] or 0}")
        
        metrics_lines.append(f"# HELP document_parser_documents_error Total number of error documents")
        metrics_lines.append(f"# TYPE document_parser_documents_error counter")
        metrics_lines.append(f"document_parser_documents_error {doc_stats['error_documents'] or 0}")
        
        # Метрики по типам документов
        metrics_lines.append(f"# HELP document_parser_documents_by_type Documents by file type")
        metrics_lines.append(f"# TYPE document_parser_documents_by_type counter")
        metrics_lines.append(f'document_parser_documents_by_type{{type="pdf"}} {doc_stats["pdf_documents"] or 0}')
        metrics_lines.append(f'document_parser_documents_by_type{{type="docx"}} {doc_stats["docx_documents"] or 0}')
        metrics_lines.append(f'document_parser_documents_by_type{{type="dwg"}} {doc_stats["dwg_documents"] or 0}')
        metrics_lines.append(f'document_parser_documents_by_type{{type="txt"}} {doc_stats["txt_documents"] or 0}')
        
        # Метрики размера
        metrics_lines.append(f"# HELP document_parser_total_size_bytes Total size of all documents in bytes")
        metrics_lines.append(f"# TYPE document_parser_total_size_bytes gauge")
        metrics_lines.append(f"document_parser_total_size_bytes {doc_stats['total_size_bytes'] or 0}")
        
        metrics_lines.append(f"# HELP document_parser_avg_file_size_bytes Average file size in bytes")
        metrics_lines.append(f"# TYPE document_parser_avg_file_size_bytes gauge")
        metrics_lines.append(f"document_parser_avg_file_size_bytes {float(doc_stats['avg_file_size_bytes'] or 0)}")
        
        # Метрики токенов
        metrics_lines.append(f"# HELP document_parser_total_tokens Total number of tokens")
        metrics_lines.append(f"# TYPE document_parser_total_tokens counter")
        metrics_lines.append(f"document_parser_total_tokens {doc_stats['total_tokens'] or 0}")
        
        # Метрики проверяемых документов
        metrics_lines.append(f"# HELP document_parser_checkable_documents_total Total number of checkable documents")
        metrics_lines.append(f"# TYPE document_parser_checkable_documents_total counter")
        metrics_lines.append(f"document_parser_checkable_documents_total {checkable_stats['total_checkable_documents'] or 0}")
        
        metrics_lines.append(f"# HELP document_parser_checkable_documents_pending_reviews Pending reviews")
        metrics_lines.append(f"# TYPE document_parser_checkable_documents_pending_reviews counter")
        metrics_lines.append(f"document_parser_checkable_documents_pending_reviews {checkable_stats['pending_reviews'] or 0}")
        
        metrics_lines.append(f"# HELP document_parser_checkable_documents_completed_reviews Completed reviews")
        metrics_lines.append(f"# TYPE document_parser_checkable_documents_completed_reviews counter")
        metrics_lines.append(f"document_parser_checkable_documents_completed_reviews {checkable_stats['completed_reviews'] or 0}")
        
        # Метрики элементов
        metrics_lines.append(f"# HELP document_parser_elements_total Total number of extracted elements")
        metrics_lines.append(f"# TYPE document_parser_elements_total counter")
        metrics_lines.append(f"document_parser_elements_total {elements_stats['total_elements'] or 0}")
        
        metrics_lines.append(f"# HELP document_parser_elements_by_type Elements by type")
        metrics_lines.append(f"# TYPE document_parser_elements_by_type counter")
        metrics_lines.append(f'document_parser_elements_by_type{{type="text"}} {elements_stats["text_elements"] or 0}')
        metrics_lines.append(f'document_parser_elements_by_type{{type="table"}} {elements_stats["table_elements"] or 0}')
        metrics_lines.append(f'document_parser_elements_by_type{{type="figure"}} {elements_stats["figure_elements"] or 0}')
        metrics_lines.append(f'document_parser_elements_by_type{{type="stamp"}} {elements_stats["stamp_elements"] or 0}')
        
        # Метрики нормоконтроля
        metrics_lines.append(f"# HELP document_parser_norm_control_results_total Total norm control results")
        metrics_lines.append(f"# TYPE document_parser_norm_control_results_total counter")
        metrics_lines.append(f"document_parser_norm_control_results_total {norm_control_stats['total_norm_control_results'] or 0}")
        
        metrics_lines.append(f"# HELP document_parser_norm_control_findings_total Total findings")
        metrics_lines.append(f"# TYPE document_parser_norm_control_findings_total counter")
        metrics_lines.append(f"document_parser_norm_control_findings_total {norm_control_stats['total_findings'] or 0}")
        
        metrics_lines.append(f"# HELP document_parser_norm_control_findings_critical Critical findings")
        metrics_lines.append(f"# TYPE document_parser_norm_control_findings_critical counter")
        metrics_lines.append(f"document_parser_norm_control_findings_critical {norm_control_stats['critical_findings'] or 0}")
        
        metrics_lines.append(f"# HELP document_parser_norm_control_findings_warning Warning findings")
        metrics_lines.append(f"# TYPE document_parser_norm_control_findings_warning counter")
        metrics_lines.append(f"document_parser_norm_control_findings_warning {norm_control_stats['warning_findings'] or 0}")
        
        metrics_lines.append(f"# HELP document_parser_norm_control_findings_info Info findings")
        metrics_lines.append(f"# TYPE document_parser_norm_control_findings_info counter")
        metrics_lines.append(f"document_parser_norm_control_findings_info {norm_control_stats['info_findings'] or 0}")
        
        # Метрики отчетов
        metrics_lines.append(f"# HELP document_parser_reports_total Total number of reports")
        metrics_lines.append(f"# TYPE document_parser_reports_total counter")
        metrics_lines.append(f"document_parser_reports_total {reports_stats['total_review_reports'] or 0}")
        
        # Метрики производительности
        metrics_lines.append(f"# HELP document_parser_documents_last_24h Documents processed in last 24 hours")
        metrics_lines.append(f"# TYPE document_parser_documents_last_24h counter")
        metrics_lines.append(f"document_parser_documents_last_24h {time_stats['documents_last_24h'] or 0}")
        
        # Возвращаем метрики в формате Prometheus
        from fastapi.responses import Response
        return Response(
            content="\n".join(metrics_lines),
            media_type="text/plain; version=0.0.4; charset=utf-8"
        )
        
    except Exception as e:
        logger.error(f"Get metrics error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

def safe_text(text: str) -> str:
    """Безопасное отображение текста в PDF с поддержкой кириллицы"""
    if text is None:
        return ""
    
    # Просто возвращаем текст как есть - кириллица будет поддерживаться
    # через использование шрифтов с поддержкой Unicode
    return str(text)

def format_long_filename(filename: str, max_chars_per_line: int = 50) -> str:
    """Форматирование длинного имени файла с переносом строк без обрезания"""
    if len(filename) <= max_chars_per_line:
        return filename
    
    # Разделяем имя файла и расширение
    name, ext = os.path.splitext(filename)
    
    # Если расширение слишком длинное, включаем его в перенос
    if len(ext) > 10:
        name = filename
        ext = ""
    
    # Разбиваем имя файла на части по max_chars_per_line символов
    lines = []
    current_line = ""
    
    for char in name:
        current_line += char
        if len(current_line) >= max_chars_per_line:
            lines.append(current_line)
            current_line = ""
    
    # Добавляем оставшуюся часть
    if current_line:
        lines.append(current_line)
    
    # Добавляем расширение к последней строке
    if lines and ext:
        lines[-1] += ext
    
    # Объединяем строки с переносом
    return "<br/>".join(lines)

def get_russian_font():
    """Получение шрифта с поддержкой кириллицы"""
    try:
        # Пытаемся использовать встроенные шрифты с поддержкой кириллицы
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        import os
        
        # Список возможных путей к шрифтам
        font_paths = [
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
            '/usr/share/fonts/TTF/DejaVuSans.ttf',
            '/usr/share/fonts/TTF/LiberationSans-Regular.ttf',
            '/usr/share/fonts/dejavu/DejaVuSans.ttf',
            '/usr/share/fonts/liberation/LiberationSans-Regular.ttf'
        ]
        
        # Пытаемся найти и зарегистрировать шрифт
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    if 'DejaVu' in font_path:
                        pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))
                        logger.info(f"Registered DejaVuSans font from {font_path}")
                        return 'DejaVuSans'
                    elif 'Liberation' in font_path:
                        pdfmetrics.registerFont(TTFont('LiberationSans', font_path))
                        logger.info(f"Registered LiberationSans font from {font_path}")
                        return 'LiberationSans'
                except Exception as e:
                    logger.warning(f"Failed to register font {font_path}: {e}")
                    continue
        
        # Пытаемся использовать встроенные Unicode шрифты ReportLab
        try:
            # Регистрируем встроенный Unicode шрифт
            pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
            logger.info("Registered STSong-Light Unicode font")
            return 'STSong-Light'
        except Exception as e:
            logger.warning(f"Failed to register Unicode font: {e}")
        
        # Если не удалось найти системные шрифты, используем встроенный
        logger.info("Using built-in Helvetica font")
        return 'Helvetica'
        
    except Exception as e:
        logger.error(f"Error in get_russian_font: {e}")
        return 'Helvetica'

def generate_docx_report_from_template(document: Dict, norm_control_result: Dict, page_results: List[Dict], review_report: Dict) -> bytes:
    """Генерация отчета на основе шаблона DOCX"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        # Загружаем шаблон
        template_path = "/app/report_format/ОТЧЕТ_file_name.docx"
        doc = Document(template_path)
        
        # Заменяем плейсхолдеры в документе
        for paragraph in doc.paragraphs:
            # Заменяем плейсхолдеры
            if "{{DOCUMENT_NAME}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{DOCUMENT_NAME}}", document['original_filename'])
            if "{{DOCUMENT_TYPE}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{DOCUMENT_TYPE}}", document['file_type'].upper())
            if "{{FILE_SIZE}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{FILE_SIZE}}", f"{document['file_size'] / 1024:.1f} KB")
            if "{{UPLOAD_DATE}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{UPLOAD_DATE}}", document['upload_date'].strftime("%d.%m.%Y %H:%M"))
            if "{{CATEGORY}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{CATEGORY}}", document['category'])
            if "{{STATUS}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{STATUS}}", document['processing_status'])
            
            # Заменяем результаты нормоконтроля
            if norm_control_result:
                if "{{TOTAL_FINDINGS}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{TOTAL_FINDINGS}}", str(norm_control_result['total_findings'] or 0))
                if "{{CRITICAL_FINDINGS}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{CRITICAL_FINDINGS}}", str(norm_control_result['critical_findings'] or 0))
                if "{{WARNING_FINDINGS}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{WARNING_FINDINGS}}", str(norm_control_result['warning_findings'] or 0))
                if "{{INFO_FINDINGS}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{INFO_FINDINGS}}", str(norm_control_result['info_findings'] or 0))
                if "{{ANALYSIS_DATE}}" in paragraph.text:
                    analysis_date = norm_control_result['analysis_date'].strftime("%d.%m.%Y %H:%M") if norm_control_result['analysis_date'] else "Не указана"
                    paragraph.text = paragraph.text.replace("{{ANALYSIS_DATE}}", analysis_date)
        
        # Добавляем детальные результаты по страницам
        if page_results:
            doc.add_paragraph("ДЕТАЛЬНЫЕ РЕЗУЛЬТАТЫ ПО СТРАНИЦАМ", style='Heading 2')
            
            for page_result in page_results:
                page_num = page_result.get('page_number', 'N/A')
                status = page_result.get('status', 'N/A')
                findings_count = len(page_result.get('findings', []))
                
                p = doc.add_paragraph()
                p.add_run(f"Страница {page_num}: ").bold = True
                p.add_run(f"Статус: {status}, Найдено замечаний: {findings_count}")
                
                # Добавляем замечания для страницы
                findings = page_result.get('findings', [])
                for finding in findings:
                    finding_p = doc.add_paragraph()
                    finding_p.add_run(f"• {finding.get('type', 'Unknown')}: ").bold = True
                    finding_p.add_run(finding.get('description', 'No description'))
        
        # Сохраняем в буфер
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        logger.error(f"Error generating DOCX report: {e}")
        # Fallback к PDF генерации
        return generate_pdf_report(document, norm_control_result, page_results, review_report)

def extract_project_info_from_filename(filename: str) -> Dict[str, str]:
    """Извлечение информации о проекте из имени файла"""
    import re
    
    project_info = {
        'project_name': 'Не определено',
        'engineering_stage': 'Не определена',
        'document_mark': 'Не определена',
        'revision': 'Не определена',
        'page_count': 'Не определено'
    }
    
    try:
        # Примеры паттернов для извлечения информации
        # Е110-0038-УКК_24.848-РД-01-02.12.032-АР_0_0_RU_IFC.pdf
        
        # Извлекаем марку документа (КЖ, КМ, АС, ТХ и т.д.)
        mark_pattern = r'[А-Я]{2}'
        mark_match = re.search(mark_pattern, filename)
        if mark_match:
            project_info['document_mark'] = mark_match.group(0)
        
        # Извлекаем номер проекта
        project_pattern = r'(\d{4}-\d{4})'
        project_match = re.search(project_pattern, filename)
        if project_match:
            project_info['project_name'] = f"Проект {project_match.group(1)}"
        
        # Извлекаем ревизию
        revision_pattern = r'_(\d+)_(\d+)_'
        revision_match = re.search(revision_pattern, filename)
        if revision_match:
            project_info['revision'] = f"{revision_match.group(1)}.{revision_match.group(2)}"
        
        # Определяем стадию проектирования по марке
        if project_info['document_mark'] in ['КЖ', 'КМ']:
            project_info['engineering_stage'] = 'КЖ/КМ'
        elif project_info['document_mark'] in ['АС', 'ТХ']:
            project_info['engineering_stage'] = 'АС/ТХ'
        elif project_info['document_mark'] in ['КС', 'КП']:
            project_info['engineering_stage'] = 'КС/КП'
        
    except Exception as e:
        logger.error(f"Error extracting project info: {e}")
    
    return project_info

def group_findings_by_pages(page_results: List[Dict]) -> Dict[int, List[Dict]]:
    """Группировка замечаний по страницам"""
    page_findings = {}
    
    for finding in page_results:
        # Извлекаем номер страницы из замечания
        page_num = finding.get('page_number', 1)
        if page_num not in page_findings:
            page_findings[page_num] = []
        page_findings[page_num].append(finding)
    
    return page_findings

def get_severity_text(severity_level: int) -> str:
    """Получение текстового описания важности замечания"""
    severity_map = {
        3: "Критическое",
        2: "Предупреждение", 
        1: "Информационное"
    }
    return severity_map.get(severity_level, "Не определено")

def generate_conclusion_from_findings(norm_control_result: Dict, findings: List[Dict]) -> str:
    """Генерация заключения на основе детальных findings"""
    conclusion_parts = []
    
    total_findings = len(findings)
    critical_findings = sum(1 for f in findings if f.get('severity_level', 1) >= 4)
    warning_findings = sum(1 for f in findings if f.get('severity_level', 1) in [2, 3])
    info_findings = sum(1 for f in findings if f.get('severity_level', 1) == 1)
    
    if total_findings == 0:
        conclusion_parts.append("Документ полностью соответствует нормативным требованиям.")
        conclusion_parts.append("Рекомендуется к принятию без замечаний.")
    else:
        if critical_findings > 0:
            conclusion_parts.append(f"Обнаружено {critical_findings} критических нарушений, требующих обязательного исправления.")
            conclusion_parts.append("Документ не может быть принят в текущем виде.")
        
        if warning_findings > 0:
            conclusion_parts.append(f"Выявлено {warning_findings} предупреждений, которые рекомендуется устранить.")
        
        if info_findings > 0:
            conclusion_parts.append(f"Найдено {info_findings} информационных замечаний для улучшения качества документа.")
        
        # Анализ по категориям
        categories = {}
        for finding in findings:
            category = finding.get('category', 'compliance')
            if category not in categories:
                categories[category] = 0
            categories[category] += 1
        
        if categories:
            conclusion_parts.append("Основные области нарушений:")
            for category, count in sorted(categories.items(), key=lambda x: x[1], reverse=True):
                category_name = {
                    'compliance': 'соответствие нормам',
                    'safety': 'безопасность',
                    'energy_efficiency': 'энергоэффективность',
                    'structural': 'конструктивные решения',
                    'formatting': 'оформление',
                    'technical': 'технические требования'
                }.get(category, category)
                conclusion_parts.append(f"- {category_name}: {count} нарушений")
        
        conclusion_parts.append("Необходимо внести соответствующие исправления и провести повторную проверку.")
    
    return " ".join(conclusion_parts)

def generate_conclusion(norm_control_result: Dict, page_summary: Dict[int, List[Dict]]) -> str:
    """Генерация заключения на основе результатов проверки"""
    if not norm_control_result:
        return "Анализ не завершен или результаты недоступны."
    
    total_findings = norm_control_result.get('total_findings', 0)
    critical_findings = norm_control_result.get('critical_findings', 0)
    warning_findings = norm_control_result.get('warning_findings', 0)
    
    total_pages = len(page_summary)
    pages_with_critical = sum(1 for findings in page_summary.values() 
                             if any(f.get('severity_level') == 3 for f in findings))
    pages_with_warnings = sum(1 for findings in page_summary.values() 
                             if any(f.get('severity_level') in [1, 2] for f in findings))
    
    conclusion_parts = []
    
    if critical_findings > 0:
        conclusion_parts.append(f"Обнаружено {critical_findings} критических замечаний на {pages_with_critical} страницах.")
        conclusion_parts.append("Документ требует исправления критических замечаний перед принятием.")
    elif warning_findings > 0:
        conclusion_parts.append(f"Обнаружено {warning_findings} замечаний на {pages_with_warnings} страницах.")
        conclusion_parts.append("Документ может быть принят с учетом устранения замечаний.")
    else:
        conclusion_parts.append("Критических замечаний и предупреждений не обнаружено.")
        conclusion_parts.append("Документ соответствует нормативным требованиям и рекомендуется к принятию.")
    
    conclusion_parts.append(f"Проверено страниц: {total_pages}")
    conclusion_parts.append(f"Общее количество замечаний: {total_findings}")
    
    return " ".join(conclusion_parts)

def generate_pdf_report_with_findings(document: Dict, norm_control_result: Dict, findings: List[Dict], review_report: Dict) -> bytes:
    """Генерация PDF отчета по результатам нормоконтроля с использованием сохраненных findings"""
    try:
        # Создаем буфер для PDF
        buffer = io.BytesIO()
        
        # Создаем документ
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Получаем шрифт с поддержкой кириллицы
        font_name = get_russian_font()
        
        # Стили с поддержкой кириллицы
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=font_name,
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=font_name,
            fontSize=14,
            spaceAfter=20,
            textColor=colors.darkblue
        )
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontName=font_name,
            fontSize=12,
            spaceAfter=15,
            textColor=colors.darkblue
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=10
        )
        small_style = ParagraphStyle(
            'CustomSmall',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=8
        )
        
        # 1. Заголовок отчета
        filename = document.get('original_filename', 'Неизвестный файл')
        story.append(Paragraph(safe_text("ОТЧЕТ ОБ АВТОМАТИЗИРОВАННОЙ ПРОВЕРКЕ"), title_style))
        story.append(Paragraph(safe_text(f'"{filename}"'), title_style))
        story.append(Spacer(1, 30))
        
        # 2. Информация о проекте и документе
        story.append(Paragraph(safe_text("1. ИНФОРМАЦИЯ О ПРОЕКТЕ И ДОКУМЕНТЕ"), heading_style))
        
        # Извлекаем информацию о проекте из имени файла
        project_info = extract_project_info_from_filename(filename)
        
        project_data = [
            [safe_text("Параметр"), safe_text("Значение")],
            [safe_text("Название проекта"), safe_text(project_info.get('project_name', 'Не определено'))],
            [safe_text("Стадия проектирования"), safe_text(project_info.get('engineering_stage', 'Не определена'))],
            [safe_text("Марка комплекта документации"), safe_text(project_info.get('document_mark', 'Не определена'))],
            [safe_text("Ревизия документации"), safe_text(project_info.get('revision', 'Не определена'))],
            [safe_text("Количество страниц"), safe_text(str(project_info.get('page_count', 'Не определено')))],
            [safe_text("Название файла"), safe_text(format_long_filename(filename))],
            [safe_text("Тип файла"), safe_text(document.get('file_type', '').upper())],
            [safe_text("Размер файла"), safe_text(f"{document.get('file_size', 0) / 1024:.1f} KB")],
            [safe_text("Дата загрузки"), safe_text(document.get('upload_date', '').strftime("%d.%m.%Y %H:%M") if document.get('upload_date') else "Не указана")],
            [safe_text("Дата проверки"), safe_text(norm_control_result.get('analysis_date', '').strftime("%d.%m.%Y %H:%M") if norm_control_result and norm_control_result.get('analysis_date') else "Не указана")]
        ]
        
        project_table = Table(project_data, colWidths=[2.5*inch, 3.5*inch])
        project_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), font_name),  # Применяем шрифт ко всем ячейкам
            ('FONTSIZE', (0, 0), (-1, -1), 10),  # Применяем размер шрифта ко всем ячейкам
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(project_table)
        story.append(Spacer(1, 20))
        
        # 3. Сводная таблица по findings
        story.append(Paragraph(safe_text("2. СВОДНАЯ ТАБЛИЦА НАРУШЕНИЙ"), heading_style))
        
        # Группируем findings по категориям
        findings_by_category = {}
        for finding in findings:
            category = finding.get('category', 'compliance')
            if category not in findings_by_category:
                findings_by_category[category] = []
            findings_by_category[category].append(finding)
        
        summary_headers = [
            safe_text("Категория"), 
            safe_text("Критич."), 
            safe_text("Предупреждения"), 
            safe_text("Инфо"), 
            safe_text("Всего")
        ]
        
        summary_data = [summary_headers]
        total_critical = 0
        total_warnings = 0
        total_info = 0
        
        for category, category_findings in findings_by_category.items():
            critical_count = sum(1 for f in category_findings if f.get('severity_level') >= 4)
            warning_count = sum(1 for f in category_findings if f.get('severity_level') in [2, 3])
            info_count = sum(1 for f in category_findings if f.get('severity_level') == 1)
            total_count = len(category_findings)
            
            total_critical += critical_count
            total_warnings += warning_count
            total_info += info_count
            
            category_name = {
                'compliance': 'Соответствие нормам',
                'safety': 'Безопасность',
                'energy_efficiency': 'Энергоэффективность',
                'structural': 'Конструктивные решения',
                'formatting': 'Оформление',
                'technical': 'Технические требования'
            }.get(category, category)
            
            summary_data.append([
                safe_text(category_name),
                safe_text(str(critical_count)),
                safe_text(str(warning_count)),
                safe_text(str(info_count)),
                safe_text(str(total_count))
            ])
        
        # Добавляем итоговую строку
        summary_data.append([
            safe_text("ИТОГО"),
            safe_text(str(total_critical)),
            safe_text(str(total_warnings)),
            safe_text(str(total_info)),
            safe_text(str(len(findings)))
        ])
        
        summary_table = Table(summary_data, colWidths=[2*inch, 1*inch, 1.5*inch, 1*inch, 1*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, -1), font_name),  # Применяем шрифт ко всем ячейкам
            ('FONTSIZE', (0, 0), (-1, -1), 9),  # Применяем размер шрифта ко всем ячейкам
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -2), colors.beige),
            ('BACKGROUND', (0, -1), (-1, -1), colors.lightblue),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 20))
        
        # 4. Детальная информация по findings
        story.append(Paragraph(safe_text("3. ДЕТАЛЬНАЯ ИНФОРМАЦИЯ ПО НАРУШЕНИЯМ"), heading_style))
        
        # Группируем findings по страницам
        findings_by_page = {}
        for finding in findings:
            element_ref = finding.get('element_reference', {})
            if isinstance(element_ref, str):
                import json
                try:
                    element_ref = json.loads(element_ref)
                except:
                    element_ref = {}
            
            page_number = element_ref.get('page_number', 1)
            if page_number not in findings_by_page:
                findings_by_page[page_number] = []
            findings_by_page[page_number].append(finding)
        
        for page_num, page_findings in sorted(findings_by_page.items()):
            if page_findings:  # Показываем только страницы с нарушениями
                story.append(Paragraph(safe_text(f"Страница {page_num}"), subheading_style))
                
                for finding in page_findings:
                    severity_text = get_severity_text(finding.get('severity_level', 1))
                    normative_doc = finding.get('normative_document_title', 'Не указан')
                    normative_clause = finding.get('normative_clause_number', 'Не указан')
                    
                    # Извлекаем информацию о месте в документе
                    element_ref = finding.get('element_reference', {})
                    if isinstance(element_ref, str):
                        import json
                        try:
                            element_ref = json.loads(element_ref)
                        except:
                            element_ref = {}
                    
                    location = element_ref.get('location', 'Не указано')
                    
                    finding_text = f"""
                    <b>Код:</b> {finding.get('rule_applied', 'Не указан')} | 
                    <b>Важность:</b> {severity_text} | 
                    <b>Категория:</b> {finding.get('category', 'Не указана')}
                    <br/>
                    <b>Нормативный документ:</b> {normative_doc} (пункт {normative_clause})
                    <br/>
                    <b>Место в документе:</b> {location}
                    <br/>
                    <b>Название:</b> {finding.get('title', 'Не указано')}
                    <br/>
                    <b>Описание:</b> {finding.get('description', 'Не указано')}
                    <br/>
                    <b>Рекомендация:</b> {finding.get('recommendation', 'Не указано')}
                    """
                    
                    story.append(Paragraph(safe_text(finding_text), normal_style))
                    story.append(Spacer(1, 10))
        
        # 5. Общие результаты
        story.append(Paragraph(safe_text("4. ОБЩИЕ РЕЗУЛЬТАТЫ ПРОВЕРКИ"), heading_style))
        
        if norm_control_result:
            results_data = [
                [safe_text("Параметр"), safe_text("Значение")],
                [safe_text("Общее количество нарушений"), safe_text(str(norm_control_result.get('total_findings', 0)))],
                [safe_text("Критические нарушения"), safe_text(str(norm_control_result.get('critical_findings', 0)))],
                [safe_text("Предупреждения"), safe_text(str(norm_control_result.get('warning_findings', 0)))],
                [safe_text("Информационные замечания"), safe_text(str(norm_control_result.get('info_findings', 0)))],
                [safe_text("Статус анализа"), safe_text(norm_control_result.get('analysis_status', 'Не указан'))]
            ]
            
            results_table = Table(results_data, colWidths=[2.5*inch, 3.5*inch])
            results_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), font_name),  # Применяем шрифт ко всем ячейкам
                ('FONTSIZE', (0, 0), (-1, -1), 10),  # Применяем размер шрифта ко всем ячейкам
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(results_table)
            story.append(Spacer(1, 20))
        
        # 6. Заключение
        story.append(Paragraph(safe_text("5. ЗАКЛЮЧЕНИЕ"), heading_style))
        
        conclusion = generate_conclusion_from_findings(norm_control_result, findings)
        story.append(Paragraph(safe_text(conclusion), normal_style))
        
        # 7. Подпись и дата
        story.append(Spacer(1, 30))
        story.append(Paragraph(safe_text(f"Отчет сгенерирован: {datetime.now().strftime('%d.%m.%Y %H:%M')}"), small_style))
        story.append(Paragraph(safe_text("Система автоматизированной проверки нормоконтроля"), small_style))
        
        # Строим PDF
        doc.build(story)
        
        # Получаем содержимое
        pdf_content = buffer.getvalue()
        buffer.close()
        
        return pdf_content
        
    except Exception as e:
        logger.error(f"PDF generation error: {e}")
        raise

def generate_pdf_report(document: Dict, norm_control_result: Dict, page_results: List[Dict], review_report: Dict) -> bytes:
    """Генерация PDF отчета по результатам нормоконтроля с улучшенной структурой (устаревшая версия)"""
    try:
        # Создаем буфер для PDF
        buffer = io.BytesIO()
        
        # Создаем документ
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Получаем шрифт с поддержкой кириллицы
        font_name = get_russian_font()
        
        # Стили с поддержкой кириллицы
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=font_name,
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=font_name,
            fontSize=14,
            spaceAfter=20,
            textColor=colors.darkblue
        )
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontName=font_name,
            fontSize=12,
            spaceAfter=15,
            textColor=colors.darkblue
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=10
        )
        small_style = ParagraphStyle(
            'CustomSmall',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=8
        )
        
        # 1. Заголовок отчета
        filename = document.get('original_filename', 'Неизвестный файл')
        story.append(Paragraph(safe_text("ОТЧЕТ ОБ АВТОМАТИЗИРОВАННОЙ ПРОВЕРКЕ"), title_style))
        story.append(Paragraph(safe_text(f'"{filename}"'), title_style))
        story.append(Spacer(1, 30))
        
        # 2. Информация о проекте и документе
        story.append(Paragraph(safe_text("1. ИНФОРМАЦИЯ О ПРОЕКТЕ И ДОКУМЕНТЕ"), heading_style))
        
        # Извлекаем информацию о проекте из имени файла
        project_info = extract_project_info_from_filename(filename)
        
        project_data = [
            [safe_text("Параметр"), safe_text("Значение")],
            [safe_text("Название проекта"), safe_text(project_info.get('project_name', 'Не определено'))],
            [safe_text("Стадия проектирования"), safe_text(project_info.get('engineering_stage', 'Не определена'))],
            [safe_text("Марка комплекта документации"), safe_text(project_info.get('document_mark', 'Не определена'))],
            [safe_text("Ревизия документации"), safe_text(project_info.get('revision', 'Не определена'))],
            [safe_text("Количество страниц"), safe_text(str(project_info.get('page_count', 'Не определено')))],
            [safe_text("Название файла"), safe_text(format_long_filename(filename))],
            [safe_text("Тип файла"), safe_text(document.get('file_type', '').upper())],
            [safe_text("Размер файла"), safe_text(f"{document.get('file_size', 0) / 1024:.1f} KB")],
            [safe_text("Дата загрузки"), safe_text(document.get('upload_date', '').strftime("%d.%m.%Y %H:%M") if document.get('upload_date') else "Не указана")],
            [safe_text("Дата проверки"), safe_text(norm_control_result.get('analysis_date', '').strftime("%d.%m.%Y %H:%M") if norm_control_result and norm_control_result.get('analysis_date') else "Не указана")]
        ]
        
        project_table = Table(project_data, colWidths=[2.5*inch, 3.5*inch])
        project_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), font_name),  # Применяем шрифт ко всем ячейкам
            ('FONTSIZE', (0, 0), (-1, -1), 10),  # Применяем размер шрифта ко всем ячейкам
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(project_table)
        story.append(Spacer(1, 20))
        
        # 3. Сводная таблица по страницам
        story.append(Paragraph(safe_text("2. СВОДНАЯ ТАБЛИЦА ПО СТРАНИЦАМ"), heading_style))
        
        # Группируем результаты по страницам
        page_summary = group_findings_by_pages(page_results)
        
        summary_headers = [
            safe_text("№ стр."), 
            safe_text("Проверена"), 
            safe_text("Критич."), 
            safe_text("Замечания"), 
            safe_text("Вывод"), 
            safe_text("Статус")
        ]
        
        summary_data = [summary_headers]
        total_critical = 0
        total_warnings = 0
        total_pages = len(page_summary)
        pages_approved = 0
        pages_rejected = 0
        
        for page_num, page_findings in sorted(page_summary.items()):
            critical_count = sum(1 for f in page_findings if f.get('severity_level') == 3)
            warning_count = sum(1 for f in page_findings if f.get('severity_level') in [1, 2])
            total_critical += critical_count
            total_warnings += warning_count
            
            # Определяем статус страницы
            if critical_count > 0:
                status = safe_text("Отклонена")
                page_status = safe_text("На исправление")
                pages_rejected += 1
            elif warning_count > 0:
                status = safe_text("Условно принята")
                page_status = safe_text("Требует доработки")
                pages_approved += 1
            else:
                status = safe_text("Принята")
                page_status = safe_text("Одобрена")
                pages_approved += 1
            
            summary_data.append([
                safe_text(str(page_num)),
                safe_text("Да"),
                safe_text(str(critical_count)),
                safe_text(str(warning_count)),
                safe_text(status),
                safe_text(page_status)
            ])
        
        # Добавляем итоговую строку
        summary_data.append([
            safe_text("ИТОГО"),
            safe_text(str(total_pages)),
            safe_text(str(total_critical)),
            safe_text(str(total_warnings)),
            safe_text(f"Принято: {pages_approved}, Отклонено: {pages_rejected}"),
            safe_text("" if total_critical == 0 else "Требует исправления")
        ])
        
        summary_table = Table(summary_data, colWidths=[0.8*inch, 1*inch, 0.8*inch, 1*inch, 1.5*inch, 1.5*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, -1), font_name),  # Применяем шрифт ко всем ячейкам
            ('FONTSIZE', (0, 0), (-1, -1), 9),  # Применяем размер шрифта ко всем ячейкам
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -2), colors.beige),
            ('BACKGROUND', (0, -1), (-1, -1), colors.lightblue),
            ('FONTNAME', (0, -1), (-1, -1), font_name),
            ('FONTSIZE', (0, -1), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 20))
        
        # 4. Детальная информация по страницам
        story.append(Paragraph(safe_text("3. ДЕТАЛЬНАЯ ИНФОРМАЦИЯ ПО СТРАНИЦАМ"), heading_style))
        
        for page_num, page_findings in sorted(page_summary.items()):
            if page_findings:  # Показываем только страницы с замечаниями
                story.append(Paragraph(safe_text(f"Страница {page_num}"), subheading_style))
                
                for finding in page_findings:
                    severity_text = get_severity_text(finding.get('severity_level', 1))
                    clause_id = finding.get('clause_id', 'Не указан')
                    location = finding.get('location', 'Не указано')
                    
                    finding_text = f"""
                    <b>Тип:</b> {finding.get('finding_type', 'Не указан')} | 
                    <b>Важность:</b> {severity_text} | 
                    <b>Clause ID:</b> {clause_id} | 
                    <b>Место:</b> {location}
                    <br/>
                    <b>Название:</b> {finding.get('title', 'Не указано')}
                    <br/>
                    <b>Описание:</b> {finding.get('description', 'Не указано')}
                    <br/>
                    <b>Рекомендация:</b> {finding.get('recommendation', 'Не указано')}
                    """
                    
                    story.append(Paragraph(safe_text(finding_text), normal_style))
                    story.append(Spacer(1, 10))
        
        # 5. Общие результаты
        story.append(Paragraph(safe_text("4. ОБЩИЕ РЕЗУЛЬТАТЫ ПРОВЕРКИ"), heading_style))
        
        if norm_control_result:
            results_data = [
                [safe_text("Параметр"), safe_text("Значение")],
                [safe_text("Общее количество замечаний"), safe_text(str(norm_control_result.get('total_findings', 0)))],
                [safe_text("Критические замечания"), safe_text(str(norm_control_result.get('critical_findings', 0)))],
                [safe_text("Предупреждения"), safe_text(str(norm_control_result.get('warning_findings', 0)))],
                [safe_text("Информационные замечания"), safe_text(str(norm_control_result.get('info_findings', 0)))],
                [safe_text("Статус анализа"), safe_text(norm_control_result.get('analysis_status', 'Не указан'))]
            ]
            
            results_table = Table(results_data, colWidths=[2.5*inch, 3.5*inch])
            results_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), font_name),  # Применяем шрифт ко всем ячейкам
                ('FONTSIZE', (0, 0), (-1, -1), 10),  # Применяем размер шрифта ко всем ячейкам
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(results_table)
            story.append(Spacer(1, 20))
        
        # 6. Заключение
        story.append(Paragraph(safe_text("5. ЗАКЛЮЧЕНИЕ"), heading_style))
        
        conclusion = generate_conclusion(norm_control_result, page_summary)
        story.append(Paragraph(safe_text(conclusion), normal_style))
        
        # 7. Подпись и дата
        story.append(Spacer(1, 30))
        story.append(Paragraph(safe_text(f"Отчет сгенерирован: {datetime.now().strftime('%d.%m.%Y %H:%M')}"), small_style))
        story.append(Paragraph(safe_text("Система автоматизированной проверки нормоконтроля"), small_style))
        
        # Строим PDF
        doc.build(story)
        
        # Получаем содержимое
        pdf_content = buffer.getvalue()
        buffer.close()
        
        return pdf_content
        
    except Exception as e:
        logger.error(f"PDF generation error: {e}")
        raise

@app.get("/checkable-documents/{document_id}/download-report")
async def download_report_pdf(document_id: int):
    """Скачивание отчета по проверке нормоконтроля в формате PDF"""
    try:
        # Получаем данные документа
        document = parser.get_checkable_document(document_id)
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Получаем результаты нормоконтроля
        norm_control_result = parser.get_norm_control_result_by_document_id(document_id)
        if not norm_control_result:
            raise HTTPException(status_code=404, detail="Norm control results not found")
        
        # Получаем детальные findings из БД
        findings = parser.get_findings_by_norm_control_id(norm_control_result['id'])
        
        # Получаем результаты по страницам
        page_results = parser.get_page_results_by_document_id(document_id)
        
        # Получаем отчет рецензента
        review_report = parser.get_review_report_by_norm_control_id(norm_control_result['id'])
        
        # Генерируем PDF отчет с использованием сохраненных findings
        report_content = generate_pdf_report_with_findings(document, norm_control_result, findings, review_report)
        
        # Возвращаем PDF файл
        from fastapi.responses import Response
        media_type = "application/pdf"
        filename = f"norm_control_report_{document_id}.pdf"
        
        return Response(
            content=report_content,
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating PDF report for document {document_id}: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@app.get("/ollama/api/tags")
async def get_models():
    """Получение списка доступных моделей Ollama"""
    try:
        import httpx
        
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.get("http://ollama:11434/api/tags")
            
            if response.status_code == 200:
                data = response.json()
                logger.info(f"Retrieved {len(data.get('models', []))} models from Ollama")
                return data
            else:
                logger.error(f"Ollama API error: {response.status_code} - {response.text}")
                return {"models": []}
                
    except Exception as e:
        logger.error(f"Error getting models from Ollama: {e}")
        # Возвращаем базовую модель в случае ошибки
        return {
            "models": [
                {
                    "name": "llama3.1-optimized-v2:latest",
                    "size": 4900000000,
                    "modified_at": "2025-08-24T00:00:00Z"
                }
            ]
        }

@app.post("/ollama/api/generate")
async def generate_response(request: Request):
    """Генерация ответа от LLM через Ollama"""
    try:
        import httpx
        import json
        
        # Получаем данные запроса
        body = await request.json()
        model = body.get("model", "llama3.1:8b")
        prompt = body.get("prompt", "")
        stream = body.get("stream", False)
        
        logger.info(f"Generating response for model: {model}, prompt length: {len(prompt)}")
        
        # Устанавливаем таймаут 2 минуты для чата с ИИ
        timeout = httpx.Timeout(120.0, connect=30.0)
        async with httpx.AsyncClient(timeout=timeout) as client:
            response = await client.post(
                "http://ollama:11434/api/generate",
                json={
                    "model": model,
                    "prompt": prompt,
                    "stream": stream
                }
            )
            
            if response.status_code == 200:
                data = response.json()
                logger.info(f"Generated response successfully, length: {len(data.get('response', ''))}")
                return data
            else:
                logger.error(f"Ollama generation error: {response.status_code} - {response.text}")
                raise HTTPException(status_code=500, detail="LLM generation failed")
                
    except Exception as e:
        logger.error(f"Error generating response: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@app.get("/api/ollama/status")
async def get_ollama_status():
    """Получение статуса Ollama для дашбоарда"""
    try:
        import httpx
        import psutil
        
        # Получаем информацию о моделях
        async with httpx.AsyncClient(timeout=5.0) as client:
            response = await client.get("http://ollama:11434/api/tags")
            models_data = response.json() if response.status_code == 200 else {"models": []}
        
        # Получаем статистику использования ресурсов Ollama
        ollama_process = None
        for proc in psutil.process_iter(['pid', 'name', 'memory_info', 'cpu_percent']):
            if 'ollama' in proc.info['name'].lower():
                ollama_process = proc
                break
        
        memory_usage = "0 GB / 8 GB"
        cpu_usage = "0%"
        
        if ollama_process:
            try:
                memory_mb = ollama_process.info['memory_info'].rss / (1024 * 1024)
                memory_usage = f"{memory_mb:.1f} GB / 8 GB"
                cpu_usage = f"{ollama_process.cpu_percent():.1f}%"
            except:
                pass
        
        return {
            "service_health": "healthy" if response.status_code == 200 else "unhealthy",
            "uptime": "02:15:30",  # TODO: Реализовать расчет uptime
            "last_heartbeat": datetime.now().isoformat(),
            "memory_usage": memory_usage,
            "cpu_usage": cpu_usage,
            "active_connections": len(models_data.get("models", [])),
            "models_count": len(models_data.get("models", []))
        }
        
    except Exception as e:
        logger.error(f"Error getting Ollama status: {e}")
        return {
            "service_health": "unhealthy",
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

@app.get("/api/ollama/performance")
async def get_ollama_performance():
    """Получение метрик производительности Ollama"""
    try:
        # Получаем статистику из базы данных
        db_conn = parser.get_db_connection()
        with db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            # Статистика за последний час
            cursor.execute("""
                SELECT 
                    COUNT(*) as requests_last_hour,
                    AVG(EXTRACT(EPOCH FROM (analysis_date - upload_date))) as avg_response_time
                FROM norm_control_results ncr
                JOIN checkable_documents cd ON ncr.checkable_document_id = cd.id
                WHERE ncr.analysis_date >= NOW() - INTERVAL '1 hour'
            """)
            performance_stats = cursor.fetchone()
            
            # Статистика успешности
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_requests,
                    COUNT(CASE WHEN analysis_status = 'completed' THEN 1 END) as successful_requests
                FROM norm_control_results
                WHERE analysis_date >= NOW() - INTERVAL '1 hour'
            """)
            success_stats = cursor.fetchone()
        
        total_requests = success_stats["total_requests"] or 0
        successful_requests = success_stats["successful_requests"] or 0
        success_rate = (successful_requests / total_requests * 100) if total_requests > 0 else 0
        
        return {
            "requests_last_hour": performance_stats["requests_last_hour"] or 0,
            "average_response_time": round(float(performance_stats["avg_response_time"] or 0), 2),
            "success_rate": round(success_rate, 1),
            "timeout_rate": round(100 - success_rate, 1),
            "tokens_generated": 12500,  # TODO: Реализовать подсчет токенов
            "tokens_per_second": 45.2   # TODO: Реализовать расчет
        }
        
    except Exception as e:
        logger.error(f"Error getting Ollama performance: {e}")
        return {
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

@app.get("/api/normcontrol/analytics")
async def get_normcontrol_analytics():
    """Получение аналитики нормоконтроля для дашбоарда"""
    try:
        db_conn = parser.get_db_connection()
        with db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
            # Общая статистика документов
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_documents,
                    COUNT(CASE WHEN review_status = 'pending' THEN 1 END) as pending_reviews,
                    COUNT(CASE WHEN review_status = 'completed' THEN 1 END) as completed_reviews,
                    COUNT(CASE WHEN review_status = 'in_progress' THEN 1 END) as in_progress_reviews
                FROM checkable_documents
            """)
            doc_stats = cursor.fetchone()
            
            # Статистика находок
            cursor.execute("""
                SELECT 
                    SUM(total_findings) as total_findings,
                    SUM(critical_findings) as critical_findings,
                    SUM(warning_findings) as warning_findings,
                    SUM(info_findings) as info_findings
                FROM norm_control_results
            """)
            findings_stats = cursor.fetchone()
            
            # Производительность
            cursor.execute("""
                SELECT 
                    AVG(EXTRACT(EPOCH FROM (analysis_date - upload_date))) as avg_processing_time,
                    COUNT(*) as documents_processed_today
                FROM norm_control_results ncr
                JOIN checkable_documents cd ON ncr.checkable_document_id = cd.id
                WHERE ncr.analysis_date >= CURRENT_DATE
            """)
            performance_stats = cursor.fetchone()
        
        total_documents = doc_stats["total_documents"] or 0
        completed_reviews = doc_stats["completed_reviews"] or 0
        compliance_rate = (completed_reviews / total_documents * 100) if total_documents > 0 else 0
        
        total_findings = findings_stats["total_findings"] or 0
        avg_findings = (total_findings / total_documents) if total_documents > 0 else 0
        
        return {
            "total_documents": total_documents,
            "compliance_rate": round(compliance_rate, 1),
            "average_findings": round(avg_findings, 2),
            "processing_efficiency": 92.3,  # TODO: Реализовать расчет
            "overview": {
                "pending_reviews": doc_stats["pending_reviews"] or 0,
                "completed_reviews": completed_reviews,
                "in_progress_reviews": doc_stats["in_progress_reviews"] or 0
            },
            "results": {
                "total_findings": total_findings,
                "critical_findings": findings_stats["critical_findings"] or 0,
                "warning_findings": findings_stats["warning_findings"] or 0,
                "info_findings": findings_stats["info_findings"] or 0
            },
            "performance": {
                "average_processing_time": round(float(performance_stats["avg_processing_time"] or 0) / 60, 1),  # в минутах
                "documents_processed_today": performance_stats["documents_processed_today"] or 0
            }
        }
        
    except Exception as e:
        logger.error(f"Error getting normcontrol analytics: {e}")
        return {
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

@app.get("/health")
async def health_check():
    """Проверка здоровья сервиса"""
    health_start_time = datetime.now()
    health_status = {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "checks": {},
        "uptime": str(datetime.now() - startup_time),
        "process_id": os.getpid()
    }
    
    try:
        # Проверка состояния shutdown
        if is_shutting_down:
            health_status["status"] = "shutting_down"
            health_status["checks"]["shutdown"] = "in_progress"
            logger.warning(f"🔍 [HEALTH] Service is shutting down, health check at {health_start_time.strftime('%Y-%m-%d %H:%M:%S.%f')[:-3]}")
        
        # Проверка памяти
        try:
            memory_info = get_memory_usage()
            if "error" not in memory_info:
                health_status["checks"]["memory"] = {
                    "rss_mb": round(memory_info['rss_mb'], 1),
                    "percent": round(memory_info['percent'], 1),
                    "available_mb": round(memory_info['available_mb'], 1)
                }
                if memory_info['percent'] > 90:
                    health_status["status"] = "degraded"
                    health_status["checks"]["memory"]["status"] = "high_usage"
                    logger.warning(f"🔍 [HEALTH] High memory usage detected: {memory_info['percent']:.1f}%")
                else:
                    health_status["checks"]["memory"]["status"] = "ok"
            else:
                health_status["checks"]["memory"] = {"status": "error", "error": memory_info["error"]}
                logger.error(f"🔍 [HEALTH] Memory check error: {memory_info['error']}")
        except Exception as mem_error:
            health_status["checks"]["memory"] = {"status": "error", "error": str(mem_error)}
            logger.error(f"🔍 [HEALTH] Memory check exception: {mem_error}")
        
        # Проверка PostgreSQL
        try:
            db_conn = parser.get_db_connection()
            with db_conn.cursor() as cursor:
                cursor.execute("SELECT 1")
                cursor.fetchone()
            health_status["checks"]["postgresql"] = {"status": "ok"}
        except Exception as pg_error:
            health_status["checks"]["postgresql"] = {"status": "error", "error": str(pg_error)}
            health_status["status"] = "unhealthy"
            logger.error(f"🔍 [HEALTH] PostgreSQL check failed: {pg_error}")
        
        # Проверка Qdrant
        try:
            parser.qdrant_client.get_collections()
            health_status["checks"]["qdrant"] = {"status": "ok"}
        except Exception as qd_error:
            health_status["checks"]["qdrant"] = {"status": "error", "error": str(qd_error)}
            health_status["status"] = "unhealthy"
            logger.error(f"🔍 [HEALTH] Qdrant check failed: {qd_error}")
        
        # Определяем общий статус
        if health_status["status"] == "healthy" and any(check.get("status") == "error" for check in health_status["checks"].values()):
            health_status["status"] = "degraded"
            logger.warning(f"🔍 [HEALTH] Service status degraded due to component errors")
        
        health_end_time = datetime.now()
        health_duration = (health_end_time - health_start_time).total_seconds()
        health_status["check_duration_ms"] = round(health_duration * 1000, 2)
        
        # Логируем результат health check
        if health_status["status"] == "healthy":
            logger.debug(f"🔍 [HEALTH] Health check passed in {health_duration:.3f}s")
        elif health_status["status"] == "degraded":
            logger.warning(f"🔍 [HEALTH] Health check degraded in {health_duration:.3f}s")
        else:
            logger.error(f"🔍 [HEALTH] Health check failed in {health_duration:.3f}s")
        
        if health_status["status"] == "unhealthy":
            return JSONResponse(
                status_code=503,
                content=health_status
            )
        elif health_status["status"] == "degraded":
            return JSONResponse(
                status_code=200,
                content=health_status
            )
        else:
            return health_status
        
    except Exception as e:
        health_end_time = datetime.now()
        health_duration = (health_end_time - health_start_time).total_seconds()
        logger.error(f"🔍 [HEALTH] Health check error after {health_duration:.3f}s: {e}")
        return JSONResponse(
            status_code=503,
            content={
                "status": "unhealthy", 
                "error": str(e),
                "timestamp": datetime.now().isoformat(),
                "check_duration_ms": round(health_duration * 1000, 2)
            }
        )

# Тестовый endpoint для проверки
@app.get("/test-prompt-templates")
async def test_prompt_templates():
    """Тестовый endpoint для проверки промпт-шаблонов"""
    return {"status": "success", "message": "Test endpoint works"}

async def shutdown_event_handler():
    """Обработчик события завершения работы"""
    global is_shutting_down
    shutdown_start_time = datetime.now()
    logger.info(f"🔍 [SHUTDOWN] Starting graceful shutdown at {shutdown_start_time.strftime('%Y-%m-%d %H:%M:%S.%f')[:-3]}")
    logger.info(f"🔍 [SHUTDOWN] Process ID: {os.getpid()}, Uptime: {shutdown_start_time - startup_time}")
    is_shutting_down = True
    
    # Логируем информацию о памяти перед shutdown
    try:
        memory_info = get_memory_usage()
        if "error" not in memory_info:
            logger.info(f"🔍 [SHUTDOWN] Memory usage before shutdown: RSS: {memory_info['rss_mb']:.1f}MB, VMS: {memory_info['vms_mb']:.1f}MB, Percent: {memory_info['percent']:.1f}%")
    except Exception as e:
        logger.warning(f"🔍 [SHUTDOWN] Could not get memory info: {e}")
    
    # Ждем завершения текущих запросов
    logger.info("🔍 [SHUTDOWN] Waiting 5 seconds for current requests to complete...")
    await asyncio.sleep(5)
    
    # Закрываем соединения с базами данных
    try:
        if parser.db_conn and not parser.db_conn.closed:
            parser.db_conn.close()
            logger.info("🔍 [SHUTDOWN] PostgreSQL connection closed successfully")
        else:
            logger.info("🔍 [SHUTDOWN] PostgreSQL connection was already closed")
    except Exception as e:
        logger.error(f"🔍 [SHUTDOWN] Error closing PostgreSQL connection: {e}")
    
    try:
        if parser.qdrant_client:
            parser.qdrant_client.close()
            logger.info("🔍 [SHUTDOWN] Qdrant connection closed successfully")
        else:
            logger.info("🔍 [SHUTDOWN] Qdrant connection was already closed")
    except Exception as e:
        logger.error(f"🔍 [SHUTDOWN] Error closing Qdrant connection: {e}")
    
    shutdown_end_time = datetime.now()
    shutdown_duration = shutdown_end_time - shutdown_start_time
    logger.info(f"🔍 [SHUTDOWN] Graceful shutdown completed at {shutdown_end_time.strftime('%Y-%m-%d %H:%M:%S.%f')[:-3]}")
    logger.info(f"🔍 [SHUTDOWN] Total shutdown duration: {shutdown_duration.total_seconds():.3f} seconds")

# Регистрируем обработчик shutdown
app.add_event_handler("shutdown", shutdown_event_handler)

if __name__ == "__main__":
    import uvicorn
    startup_time = datetime.now()
    logger.info(f"🔍 [STARTUP] Starting Document Parser Service at {startup_time.strftime('%Y-%m-%d %H:%M:%S.%f')[:-3]}")
    logger.info(f"🔍 [STARTUP] Process ID: {os.getpid()}, Parent PID: {os.getppid()}")
    logger.info(f"🔍 [STARTUP] Working directory: {os.getcwd()}")
    logger.info(f"🔍 [STARTUP] Python version: {sys.version}")
    
    # Логируем информацию о системе
    try:
        import platform
        logger.info(f"🔍 [STARTUP] Platform: {platform.platform()}")
        logger.info(f"🔍 [STARTUP] Architecture: {platform.architecture()}")
        logger.info(f"🔍 [STARTUP] Machine: {platform.machine()}")
    except Exception as e:
        logger.warning(f"🔍 [STARTUP] Could not get platform info: {e}")
    
    # Логируем информацию о памяти при запуске
    try:
        memory_info = get_memory_usage()
        if "error" not in memory_info:
            logger.info(f"🔍 [STARTUP] Initial memory usage: RSS: {memory_info['rss_mb']:.1f}MB, VMS: {memory_info['vms_mb']:.1f}MB, Percent: {memory_info['percent']:.1f}%")
    except Exception as e:
        logger.warning(f"🔍 [STARTUP] Could not get initial memory info: {e}")
    
    uvicorn.run(app, host="0.0.0.0", port=8001)
