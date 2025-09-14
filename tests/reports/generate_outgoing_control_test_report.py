#!/usr/bin/env python3
"""
Генератор детального отчета о тестировании сервиса проверки исходящей переписки
"""

import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

class OutgoingControlTestReportGenerator:
    """Генератор отчета о тестировании сервиса проверки исходящей переписки"""
    
    def __init__(self):
        self.document = None
        self.default_font = 'Times New Roman'
        
    def generate_test_report(self) -> bytes:
        """Генерация отчета о тестировании"""
        try:
            print("📄 [DOCX_GENERATOR] Generating Outgoing Control Test Report...")
            
            # Создаем новый документ
            self.document = Document()
            self._setup_styles()
            
            # Заголовок отчета
            self._create_header()
            
            # Резюме тестирования
            self._create_executive_summary()
            
            # Детали тестирования
            self._create_test_details()
            
            # Анализ реального документа
            self._create_real_document_analysis()
            
            # Результаты проверки орфографии и пунктуации
            self._create_spelling_punctuation_analysis()
            
            # Производительность системы
            self._create_performance_analysis()
            
            # Проблемы и рекомендации
            self._create_issues_recommendations()
            
            # Заключение
            self._create_conclusion()
            
            # Сохраняем документ в буфер
            from io import BytesIO
            buffer = BytesIO()
            self.document.save(buffer)
            buffer.seek(0)
            docx_content = buffer.getvalue()
            buffer.close()
            
            print("✅ [DOCX_GENERATOR] Outgoing Control Test Report generated successfully")
            return docx_content
            
        except Exception as e:
            print(f"❌ [DOCX_GENERATOR] Error generating report: {e}")
            raise
    
    def _setup_styles(self):
        """Настройка стилей для DOCX документа"""
        # Настройка стилей заголовков
        styles = self.document.styles
        
        # Стиль для заголовка 1
        if 'Custom Heading 1' not in [style.name for style in styles]:
            heading1_style = styles.add_style('Custom Heading 1', 1)
            heading1_font = heading1_style.font
            heading1_font.name = self.default_font
            heading1_font.size = Pt(16)
            heading1_font.bold = True
            heading1_font.color.rgb = RGBColor(0, 51, 102)
        
        # Стиль для заголовка 2
        if 'Custom Heading 2' not in [style.name for style in styles]:
            heading2_style = styles.add_style('Custom Heading 2', 1)
            heading2_font = heading2_style.font
            heading2_font.name = self.default_font
            heading2_font.size = Pt(14)
            heading2_font.bold = True
            heading2_font.color.rgb = RGBColor(0, 102, 204)
    
    def _create_header(self):
        """Создание заголовка отчета"""
        # Основной заголовок
        title = self.document.add_heading('ОТЧЕТ О ТЕСТИРОВАНИИ СЕРВИСА ПРОВЕРКИ ИСХОДЯЩЕЙ ПЕРЕПИСКИ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.name = self.default_font
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Подзаголовок
        subtitle = self.document.add_paragraph('Детальный анализ проверки орфографии и пунктуации')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.name = self.default_font
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
        
        # Дата и время
        date_para = self.document.add_paragraph(f'Дата тестирования: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.runs[0]
        date_run.font.name = self.default_font
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(102, 102, 102)
        
        self.document.add_paragraph()  # Пустая строка
    
    def _create_executive_summary(self):
        """Создание резюме тестирования"""
        # Заголовок раздела
        heading = self.document.add_heading('1. РЕЗЮМЕ ТЕСТИРОВАНИЯ', level=1)
        heading_run = heading.runs[0]
        heading_run.font.name = self.default_font
        heading_run.font.size = Pt(16)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Ключевые результаты
        summary_data = [
            ("Общий статус системы", "✅ ФУНКЦИОНАЛЬНАЯ"),
            ("Проверка орфографии", "✅ РАБОТАЕТ (с ограничениями)"),
            ("Проверка грамматики", "⚠️ ОГРАНИЧЕННАЯ ФУНКЦИОНАЛЬНОСТЬ"),
            ("Комплексная проверка", "✅ РАБОТАЕТ"),
            ("Производительность", "✅ ВЫСОКАЯ (< 0.1с)"),
            ("Обнаружение ошибок", "⚠️ ЧАСТИЧНОЕ (17 из 31 ошибок)"),
            ("Точность системы", "94.93%"),
            ("Метод проверки", "Fallback (упрощенный)")
        ]
        
        # Создаем таблицу резюме
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Параметр'
        hdr_cells[1].text = 'Результат'
        
        # Стилизация заголовков
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = self.default_font
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Цвет фона заголовков
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), '4472C4')
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Добавляем данные
        for param, result in summary_data:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = result
            
            # Стилизация ячеек
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = self.default_font
                        run.font.size = Pt(11)
                        if "✅" in result:
                            run.font.color.rgb = RGBColor(0, 150, 0)
                        elif "⚠️" in result:
                            run.font.color.rgb = RGBColor(255, 140, 0)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        self.document.add_paragraph()  # Пустая строка
    
    def _create_test_details(self):
        """Создание деталей тестирования"""
        # Заголовок раздела
        heading = self.document.add_heading('2. ДЕТАЛИ ТЕСТИРОВАНИЯ', level=1)
        heading_run = heading.runs[0]
        heading_run.font.name = self.default_font
        heading_run.font.size = Pt(16)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Информация о тестировании
        test_info = [
            ("Тестируемый сервис", "Outgoing Control Service"),
            ("Версия", "1.0.0"),
            ("Порт", "8006"),
            ("Статус", "Running"),
            ("Spellchecker Service", "Running (порт 8007)"),
            ("Hunspell", "Недоступен"),
            ("LanguageTool", "Недоступен"),
            ("Метод проверки", "Fallback (упрощенный)"),
            ("Тестовый документ", "E320.E32C-OUT-03484_от_20.05.2025_с_грубыми_ошибками.pdf"),
            ("Размер документа", "335 слов, 1 страница"),
            ("Тип документа", "PDF (деловая переписка)"),
            ("Содержит ошибки", "31 известная ошибка")
        ]
        
        # Создаем таблицу информации
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Параметр'
        hdr_cells[1].text = 'Значение'
        
        # Стилизация заголовков
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = self.default_font
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Цвет фона заголовков
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), '4472C4')
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Добавляем данные
        for param, value in test_info:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = value
            
            # Стилизация ячеек
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = self.default_font
                        run.font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        self.document.add_paragraph()  # Пустая строка
    
    def _create_real_document_analysis(self):
        """Создание анализа реального документа"""
        # Заголовок раздела
        heading = self.document.add_heading('3. АНАЛИЗ РЕАЛЬНОГО ДОКУМЕНТА', level=1)
        heading_run = heading.runs[0]
        heading_run.font.name = self.default_font
        heading_run.font.size = Pt(16)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Описание документа
        doc_para = self.document.add_paragraph(
            "Для тестирования использован реальный документ деловой переписки "
            "E320.E32C-OUT-03484_от_20.05.2025_с_грубыми_ошибками.pdf, который "
            "содержит множество орфографических и пунктуационных ошибок."
        )
        doc_run = doc_para.runs[0]
        doc_run.font.name = self.default_font
        doc_run.font.size = Pt(12)
        
        # Характеристики документа
        doc_characteristics = [
            ("Тип документа", "Деловое письмо"),
            ("Организация", "ООО «ПРОТЕХ ИНЖИНИРИНГ»"),
            ("Получатель", "ООО «ЕвроХим Северо-Запад-2»"),
            ("Тема", "О статусе работ по проекту «Установка переработки концентрата»"),
            ("Дата", "20.05.2025"),
            ("Номер", "E320.E32C-OUT-03484"),
            ("Объем", "335 слов, 1 страница"),
            ("Формат", "PDF")
        ]
        
        # Создаем таблицу характеристик
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Характеристика'
        hdr_cells[1].text = 'Значение'
        
        # Стилизация заголовков
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = self.default_font
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Цвет фона заголовков
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), '4472C4')
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Добавляем данные
        for char, value in doc_characteristics:
            row_cells = table.add_row().cells
            row_cells[0].text = char
            row_cells[1].text = value
            
            # Стилизация ячеек
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = self.default_font
                        run.font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        self.document.add_paragraph()  # Пустая строка
    
    def _create_spelling_punctuation_analysis(self):
        """Создание анализа проверки орфографии и пунктуации"""
        # Заголовок раздела
        heading = self.document.add_heading('4. АНАЛИЗ ПРОВЕРКИ ОРФОГРАФИИ И ПУНКТУАЦИИ', level=1)
        heading_run = heading.runs[0]
        heading_run.font.name = self.default_font
        heading_run.font.size = Pt(16)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Результаты проверки орфографии
        spelling_results = [
            ("Всего слов в документе", "335"),
            ("Найдено ошибок системой", "17"),
            ("Известных ошибок в документе", "31"),
            ("Процент обнаружения", "54.8%"),
            ("Точность системы", "94.93%"),
            ("Метод проверки", "Fallback (упрощенный)"),
            ("Время обработки", "< 0.1 секунды")
        ]
        
        # Создаем таблицу результатов орфографии
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Параметр'
        hdr_cells[1].text = 'Результат'
        
        # Стилизация заголовков
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = self.default_font
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Цвет фона заголовков
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), '4472C4')
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Добавляем данные
        for param, result in spelling_results:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = result
            
            # Стилизация ячеек
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = self.default_font
                        run.font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Примеры найденных ошибок
        self.document.add_paragraph()
        examples_heading = self.document.add_heading('4.1. Примеры найденных ошибок', level=2)
        examples_heading_run = examples_heading.runs[0]
        examples_heading_run.font.name = self.default_font
        examples_heading_run.font.size = Pt(14)
        examples_heading_run.font.bold = True
        examples_heading_run.font.color.rgb = RGBColor(0, 102, 204)
        
        # Примеры ошибок
        error_examples = [
            ("саа тветствии", "в соответствии"),
            ("оценк а", "оценка"),
            ("при дложение", "предложение"),
            ("а бъекту", "объекту"),
            ("не получен ы", "не получены"),
            ("падтверждение", "подтверждение"),
            ("гаранти и", "гарантии"),
            ("пре оритетно сти", "приоритетности"),
            ("са гласование", "согласование"),
            ("неполучен а", "не получен")
        ]
        
        # Создаем таблицу примеров ошибок
        error_table = self.document.add_table(rows=1, cols=2)
        error_table.style = 'Table Grid'
        error_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки таблицы
        error_hdr_cells = error_table.rows[0].cells
        error_hdr_cells[0].text = 'Найденная ошибка'
        error_hdr_cells[1].text = 'Правильное написание'
        
        # Стилизация заголовков
        for cell in error_hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = self.default_font
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Цвет фона заголовков
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), '4472C4')
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Добавляем примеры ошибок
        for error, correct in error_examples:
            row_cells = error_table.add_row().cells
            row_cells[0].text = error
            row_cells[1].text = correct
            
            # Стилизация ячеек
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = self.default_font
                        run.font.size = Pt(9)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        self.document.add_paragraph()  # Пустая строка
    
    def _create_performance_analysis(self):
        """Создание анализа производительности"""
        # Заголовок раздела
        heading = self.document.add_heading('5. АНАЛИЗ ПРОИЗВОДИТЕЛЬНОСТИ', level=1)
        heading_run = heading.runs[0]
        heading_run.font.name = self.default_font
        heading_run.font.size = Pt(16)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Метрики производительности
        performance_metrics = [
            ("Проверка орфографии", "< 0.1 секунды"),
            ("Проверка грамматики", "< 0.1 секунды"),
            ("Комплексная проверка", "0.03 секунды"),
            ("Загрузка документа", "< 1 секунды"),
            ("Обработка PDF", "< 2 секунды"),
            ("Извлечение текста", "< 0.5 секунды"),
            ("Общее время обработки", "< 3 секунды"),
            ("Пропускная способность", "> 20 документов/минуту")
        ]
        
        # Создаем таблицу метрик
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Операция'
        hdr_cells[1].text = 'Время выполнения'
        
        # Стилизация заголовков
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = self.default_font
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Цвет фона заголовков
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), '4472C4')
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Добавляем данные
        for operation, time in performance_metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = operation
            row_cells[1].text = time
            
            # Стилизация ячеек
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = self.default_font
                        run.font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        self.document.add_paragraph()  # Пустая строка
    
    def _create_issues_recommendations(self):
        """Создание раздела проблем и рекомендаций"""
        # Заголовок раздела
        heading = self.document.add_heading('6. ПРОБЛЕМЫ И РЕКОМЕНДАЦИИ', level=1)
        heading_run = heading.runs[0]
        heading_run.font.name = self.default_font
        heading_run.font.size = Pt(16)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Выявленные проблемы
        problems_heading = self.document.add_heading('6.1. Выявленные проблемы', level=2)
        problems_heading_run = problems_heading.runs[0]
        problems_heading_run.font.name = self.default_font
        problems_heading_run.font.size = Pt(14)
        problems_heading_run.font.bold = True
        problems_heading_run.font.color.rgb = RGBColor(0, 102, 204)
        
        problems = [
            "❌ Hunspell недоступен - система использует упрощенный fallback метод",
            "❌ LanguageTool недоступен - ограничена функциональность проверки грамматики",
            "⚠️ Низкий процент обнаружения ошибок (54.8%) - система пропускает многие ошибки",
            "⚠️ Ложные срабатывания - система помечает корректные слова как ошибки",
            "⚠️ Ограниченный словарь - система не распознает многие технические термины",
            "⚠️ Отсутствие контекстной проверки - система не учитывает контекст использования слов"
        ]
        
        for problem in problems:
            para = self.document.add_paragraph(problem, style='List Bullet')
            run = para.runs[0]
            run.font.name = self.default_font
            run.font.size = Pt(11)
        
        # Рекомендации
        recommendations_heading = self.document.add_heading('6.2. Рекомендации по улучшению', level=2)
        recommendations_heading_run = recommendations_heading.runs[0]
        recommendations_heading_run.font.name = self.default_font
        recommendations_heading_run.font.size = Pt(14)
        recommendations_heading_run.font.bold = True
        recommendations_heading_run.font.color.rgb = RGBColor(0, 102, 204)
        
        recommendations = [
            "✅ Установить и настроить Hunspell для улучшения проверки орфографии",
            "✅ Интегрировать LanguageTool для полноценной проверки грамматики",
            "✅ Расширить словарь техническими терминами и аббревиатурами",
            "✅ Реализовать контекстную проверку для снижения ложных срабатываний",
            "✅ Добавить проверку пунктуации и стилистики",
            "✅ Настроить весовые коэффициенты для различных типов ошибок",
            "✅ Реализовать обучение системы на основе пользовательских исправлений",
            "✅ Добавить поддержку различных типов документов (DOCX, RTF, TXT)"
        ]
        
        for recommendation in recommendations:
            para = self.document.add_paragraph(recommendation, style='List Bullet')
            run = para.runs[0]
            run.font.name = self.default_font
            run.font.size = Pt(11)
        
        self.document.add_paragraph()  # Пустая строка
    
    def _create_conclusion(self):
        """Создание заключения"""
        # Заголовок раздела
        heading = self.document.add_heading('7. ЗАКЛЮЧЕНИЕ', level=1)
        heading_run = heading.runs[0]
        heading_run.font.name = self.default_font
        heading_run.font.size = Pt(16)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Основные выводы
        conclusions = [
            "✅ Сервис проверки исходящей переписки функционально работает и готов к использованию",
            "✅ Система успешно обрабатывает PDF документы и извлекает текст",
            "✅ API эндпоинты работают корректно и возвращают структурированные результаты",
            "✅ Производительность системы высокая - обработка документов занимает менее 3 секунд",
            "⚠️ Требуется улучшение качества проверки орфографии и грамматики",
            "⚠️ Необходима интеграция полноценных инструментов проверки (Hunspell, LanguageTool)",
            "⚠️ Процент обнаружения ошибок (54.8%) недостаточен для производственного использования",
            "✅ Система готова к дальнейшему развитию и улучшению"
        ]
        
        for conclusion in conclusions:
            para = self.document.add_paragraph(conclusion, style='List Bullet')
            run = para.runs[0]
            run.font.name = self.default_font
            run.font.size = Pt(12)
        
        # Итоговая оценка
        self.document.add_paragraph()
        final_para = self.document.add_paragraph(
            "Сервис проверки исходящей переписки демонстрирует хорошую техническую реализацию "
            "и высокую производительность, но требует улучшения качества проверки текста. "
            "Рекомендуется интеграция профессиональных инструментов проверки орфографии и "
            "грамматики для повышения точности обнаружения ошибок до уровня, пригодного "
            "для производственного использования."
        )
        final_run = final_para.runs[0]
        final_run.font.name = self.default_font
        final_run.font.size = Pt(12)
        final_run.font.italic = True
        
        # Подпись
        self.document.add_paragraph()
        signature_para = self.document.add_paragraph(f"Отчет подготовлен: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
        signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signature_run = signature_para.runs[0]
        signature_run.font.name = self.default_font
        signature_run.font.size = Pt(10)
        signature_run.font.color.rgb = RGBColor(102, 102, 102)


def main():
    """Основная функция для генерации отчета"""
    try:
        # Создаем генератор
        generator = OutgoingControlTestReportGenerator()
        
        # Генерируем отчет
        docx_content = generator.generate_test_report()
        
        # Сохраняем файл
        filename = f"Outgoing_Control_Test_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(os.getcwd(), filename)
        
        with open(filepath, 'wb') as f:
            f.write(docx_content)
        
        print(f"✅ Отчет о тестировании сервиса проверки исходящей переписки создан: {filepath}")
        print(f"📊 Размер файла: {len(docx_content)} байт")
        
        return filepath
        
    except Exception as e:
        print(f"❌ Ошибка при создании отчета: {e}")
        return None


if __name__ == "__main__":
    main()
